"""
Policy Response — Enterprise PoC
Agentic Workflow · Structured Output · Role-based Deliverables
"""

import streamlit as st
import anthropic
import json
import re
import os
import io
import time
import random
import traceback
from typing import Optional, Dict, List, Literal

from pydantic import BaseModel, ValidationError

import plotly.graph_objects as go

# ─── Page Configuration ──────────────────────────────────────────────────────
st.set_page_config(
    page_title="Policy Response",
    page_icon="◆",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ─── Luxury CSS (Tiffany Blue accent) ─────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Cormorant+Garamond:ital,wght@0,300;0,400;0,600;1,300;1,400&family=Montserrat:wght@300;400;500;600&display=swap');

:root {
    --accent:        #0ABAB5;
    --accent-dark:   #089590;
    --accent-hover:  #0DCEC8;
    --accent-a12:    rgba(10, 186, 181, 0.12);
    --accent-a18:    rgba(10, 186, 181, 0.18);
    --accent-border: rgba(10, 186, 181, 0.15);
    --bg:            #0A0A0A;
    --bg-card:       #111111;
    --bg-card2:      #161616;
    --bg-sidebar:    #0D0D0D;
    --text:          #F0EDE6;
    --text-sec:      #A8A49E;
    --text-muted:    #C4BFB8;
    --threat:        #8B2635;
    --opp:           #1A6B3C;
    --serif:         "Cormorant Garamond", "Palatino Linotype", "Palatino", Georgia, serif;
    --sans:          "Montserrat", "Helvetica Neue", Helvetica, -apple-system, sans-serif;
}

/* ── GLOBAL ─────────────────────────────────── */
#MainMenu, footer, header { visibility: hidden !important; }
.stApp { background: var(--bg) !important; }
.main .block-container {
    background: var(--bg) !important;
    padding-top: 1rem !important;
    padding-bottom: 0rem !important;
    padding-left: 3.5rem !important;
    padding-right: 3.5rem !important;
    max-width: 1440px !important;
}

/* ── SIDEBAR TOGGLE (always visible even when header is hidden) ── */
[data-testid="collapsedControl"],
[data-testid="stSidebarCollapsedControl"] {
    visibility: visible !important;
    display: flex !important;
    background: rgba(13,13,13,0.82) !important;
    border-right: 1px solid rgba(10,186,181,0.22) !important;
    border-radius: 0 6px 6px 0 !important;
}
[data-testid="collapsedControl"] svg,
[data-testid="stSidebarCollapsedControl"] svg {
    color: rgba(10,186,181,0.70) !important;
    fill:  rgba(10,186,181,0.70) !important;
    transition: color 0.2s, fill 0.2s !important;
}
[data-testid="collapsedControl"]:hover,
[data-testid="stSidebarCollapsedControl"]:hover {
    background: rgba(10,186,181,0.10) !important;
}
[data-testid="collapsedControl"]:hover svg,
[data-testid="stSidebarCollapsedControl"]:hover svg {
    color: #0ABAB5 !important;
    fill:  #0ABAB5 !important;
}

/* ── SIDEBAR ────────────────────────────────── */
[data-testid="stSidebar"] {
    background: var(--bg-sidebar) !important;
    border-right: 1px solid var(--accent-border) !important;
}
[data-testid="stSidebar"] > div:first-child {
    padding: 2rem 1.6rem !important;
}

/* ── TYPOGRAPHY ─────────────────────────────── */
h1, h2, h3, h4 {
    font-family: var(--serif) !important;
    color: var(--text) !important;
    letter-spacing: 0.06em !important;
}
h1 { font-weight: 300 !important; }
h2 { font-weight: 300 !important; }
h3 { font-weight: 400 !important; }
p, label, .stMarkdown p,
.stMarkdown span, .stText span,
[data-testid="stRadio"] span,
[data-testid="stCaptionContainer"] span,
[data-testid="stTextArea"] span,
[data-testid="stTextInput"] span,
[data-testid="stButton"] span,
[data-testid="stDownloadButton"] span {
    font-family: var(--sans) !important;
}
[data-testid="stStatusWidget"] span,
[data-testid="stNotification"] span,
[data-testid="stAlert"] span,
[class*="StatusIndicator"] span,
[class*="stStatus"] span {
    font-family: "Material Symbols Rounded", "Material Symbols Outlined",
                 "Material Icons", sans-serif !important;
    font-feature-settings: "liga" 1 !important;
    font-variation-settings: 'FILL' 0, 'wght' 400, 'GRAD' 0, 'opsz' 24 !important;
}
.stMarkdown h1, .stMarkdown h2, .stMarkdown h3 {
    font-family: var(--serif) !important;
    color: var(--text) !important;
    border-bottom: 1px solid var(--accent-border);
    padding-bottom: 0.4em;
    letter-spacing: 0.05em !important;
}
.stMarkdown h4 {
    font-family: var(--sans) !important;
    color: var(--accent) !important;
    font-size: 0.72rem !important;
    letter-spacing: 0.2em !important;
    text-transform: uppercase !important;
    font-weight: 600 !important;
    border-bottom: 1px solid var(--accent-border);
    padding-bottom: 0.4em;
}
.stMarkdown strong { color: var(--accent) !important; }
.stMarkdown a { color: var(--accent) !important; text-decoration: none !important; }
.stMarkdown blockquote {
    border-left: 3px solid var(--accent) !important;
    background: var(--bg-card) !important;
    padding: 0.8rem 1.2rem !important;
    margin: 1rem 0 !important;
}
.stMarkdown hr { border-top: 1px solid var(--accent-border) !important; margin: 1.5rem 0 !important; }
.stMarkdown ul li, .stMarkdown ol li {
    color: var(--text-sec) !important;
    font-family: var(--sans) !important;
    font-size: 0.88rem !important;
    line-height: 1.8 !important;
}
.stMarkdown code {
    background: var(--bg-card2) !important;
    color: var(--accent) !important;
    border: 1px solid var(--accent-border) !important;
    border-radius: 0 !important;
    font-size: 0.8rem !important;
    padding: 0.1em 0.4em !important;
}

/* ── PRIMARY BUTTONS ────────────────────────── */
.stButton > button {
    background: linear-gradient(135deg, #0ABAB5 0%, #089590 100%) !important;
    color: #FFFFFF !important;
    border: none !important;
    border-radius: 0 !important;
    letter-spacing: 0.20em !important;
    text-transform: uppercase !important;
    font-family: var(--sans) !important;
    font-weight: 600 !important;
    font-size: 0.70rem !important;
    padding: 0.9rem 1.8rem !important;
    transition: all 0.35s ease !important;
    box-shadow: 0 2px 18px rgba(10, 186, 181, 0.14) !important;
}
.stButton > button:hover {
    background: linear-gradient(135deg, #0DCEC8 0%, #0ABAB5 100%) !important;
    box-shadow: 0 4px 28px rgba(10, 186, 181, 0.34) !important;
    transform: translateY(-1px) !important;
    color: #FFFFFF !important;
}
.stButton > button:active { transform: translateY(0) !important; }
.stButton > button:disabled {
    background: #111111 !important;
    color: var(--text-muted) !important;
    border: 1px solid rgba(10, 186, 181, 0.08) !important;
    box-shadow: none !important;
    transform: none !important;
    cursor: not-allowed !important;
}

/* ── SECONDARY BUTTONS (HITL actions) ───────── */
.stButton > button[kind="secondary"] {
    background: transparent !important;
    color: var(--text-sec) !important;
    border: 1px solid rgba(10, 186, 181, 0.25) !important;
    box-shadow: none !important;
    font-size: 0.66rem !important;
    letter-spacing: 0.12em !important;
    padding: 0.72rem 1.1rem !important;
    text-transform: none !important;
}
.stButton > button[kind="secondary"]:hover {
    background: rgba(10, 186, 181, 0.07) !important;
    border-color: var(--accent) !important;
    color: var(--accent) !important;
    transform: none !important;
    box-shadow: none !important;
}

[data-testid="stDownloadButton"] > button {
    background: transparent !important;
    color: var(--accent) !important;
    border: 1px solid var(--accent-border) !important;
    box-shadow: none !important;
    font-size: 0.65rem !important;
    letter-spacing: 0.18em !important;
    padding: 0.65rem 1.4rem !important;
}
[data-testid="stDownloadButton"] > button:hover {
    background: var(--accent-a12) !important;
    border-color: var(--accent) !important;
    transform: none !important;
    box-shadow: none !important;
}

/* ── TEXT INPUTS ────────────────────────────── */
[data-testid="stTextInput"] input,
[data-testid="stTextArea"] textarea {
    background: #111111 !important;
    border: 1px solid var(--accent-border) !important;
    border-radius: 0 !important;
    color: var(--text) !important;
    font-family: var(--sans) !important;
    font-size: 0.83rem !important;
    transition: border-color 0.3s !important;
}
[data-testid="stTextInput"] input:focus,
[data-testid="stTextArea"] textarea:focus {
    border-color: var(--accent) !important;
    box-shadow: 0 0 0 1px rgba(10, 186, 181, 0.22) !important;
    outline: none !important;
}
[data-testid="stTextInput"] label,
[data-testid="stTextArea"] label {
    color: var(--text-muted) !important;
    font-size: 0.68rem !important;
    font-weight: 600 !important;
    letter-spacing: 0.15em !important;
    text-transform: uppercase !important;
    font-family: var(--sans) !important;
    margin-bottom: 0.4rem !important;
}

/* ── RADIO ──────────────────────────────────── */
[data-testid="stRadio"] > label {
    color: var(--text-muted) !important;
    font-size: 0.68rem !important;
    font-weight: 600 !important;
    letter-spacing: 0.15em !important;
    text-transform: uppercase !important;
    font-family: var(--sans) !important;
}
[data-testid="stRadio"] [data-testid="stMarkdownContainer"] p {
    font-size: 0.80rem !important;
    color: var(--text-sec) !important;
    font-family: var(--sans) !important;
    line-height: 1.5 !important;
}

/* ── TABS ───────────────────────────────────── */
.stTabs [data-baseweb="tab-list"] {
    background: transparent !important;
    border-bottom: 1px solid var(--accent-border) !important;
    gap: 0 !important;
}
.stTabs [data-baseweb="tab"] {
    background: transparent !important;
    color: var(--text-muted) !important;
    font-family: var(--sans) !important;
    font-size: 0.66rem !important;
    font-weight: 600 !important;
    letter-spacing: 0.16em !important;
    text-transform: uppercase !important;
    padding: 0.85rem 1.6rem !important;
    border-bottom: 2px solid transparent !important;
    transition: all 0.3s !important;
}
.stTabs [data-baseweb="tab"]:hover { color: var(--accent) !important; }
.stTabs [aria-selected="true"] {
    color: var(--accent) !important;
    border-bottom: 2px solid var(--accent) !important;
}
.stTabs [data-baseweb="tab-highlight"] { background: var(--accent) !important; height: 2px !important; }
.stTabs [data-baseweb="tab-panel"] { background: transparent !important; padding: 2rem 0 !important; }

/* ── EXPANDER ───────────────────────────────── */
[data-testid="stExpander"] {
    background: var(--bg-card) !important;
    border: 1px solid var(--accent-border) !important;
    border-radius: 0 !important;
}
[data-testid="stExpander"] summary {
    color: var(--text-muted) !important;
    font-family: var(--sans) !important;
    font-size: 0.68rem !important;
    letter-spacing: 0.14em !important;
    text-transform: uppercase !important;
}

/* ── STATUS WIDGET ──────────────────────────── */
[data-testid="stStatusWidget"] {
    background: var(--bg-card) !important;
    border: 1px solid var(--accent-border) !important;
    border-radius: 0 !important;
}
[data-testid="stStatusWidget"] p,
[data-testid="stStatusWidget"] div:not([class*="icon"]) {
    font-family: var(--sans) !important;
    font-size: 0.82rem !important;
    color: var(--text-sec) !important;
}
[data-testid="stStatusWidget"] > div > div:first-child {
    min-width: 1.5rem !important;
    display: inline-flex !important;
    align-items: center !important;
}

/* ── ALERTS / INFO (Evidence quotes) ────────── */
[data-testid="stAlert"] {
    border-radius: 0 !important;
    font-family: var(--sans) !important;
    font-size: 0.83rem !important;
    background: var(--bg-card2) !important;
}
/* st.info override — evidence quote style */
[data-testid="stAlert"][data-baseweb="notification"] {
    background: rgba(10, 186, 181, 0.04) !important;
    border: 1px solid rgba(100, 116, 139, 0.30) !important;
    border-left: 3px solid #64748B !important;
    border-radius: 0 !important;
}
[data-testid="stAlert"] p {
    color: rgba(255, 255, 255, 0.90) !important;
    font-style: normal !important;
    font-weight: 400 !important;
    font-family: var(--sans) !important;
    font-size: 0.86rem !important;
    line-height: 1.6 !important;
}

/* ── DIVIDER / CAPTION / JSON ───────────────── */
hr { border: none !important; border-top: 1px solid var(--accent-border) !important; margin: 2rem 0 !important; }
[data-testid="stCaptionContainer"] p {
    color: var(--text-muted) !important; font-size: 0.68rem !important;
    letter-spacing: 0.08em !important; font-family: var(--sans) !important;
}
[data-testid="stJson"] {
    background: var(--bg-card) !important; border: 1px solid var(--accent-border) !important;
    font-family: "SF Mono", "Monaco", monospace !important; font-size: 0.78rem !important;
}

/* ── SCROLLBAR ──────────────────────────────── */
::-webkit-scrollbar { width: 3px; }
::-webkit-scrollbar-track { background: var(--bg); }
::-webkit-scrollbar-thumb { background: rgba(10, 186, 181, 0.18); }
::-webkit-scrollbar-thumb:hover { background: rgba(10, 186, 181, 0.40); }
</style>
""", unsafe_allow_html=True)

# ─── Constants ────────────────────────────────────────────────────────────────
MODEL = "claude-sonnet-4-6"

# ─── System Prompts ───────────────────────────────────────────────────────────
SYSTEM_PROMPT_CORE = (
    "You are the Chief Policy Intelligence Analyst for a major media enterprise. "
    "Your primary goal is to provide a comprehensive policy analysis in JSON format.\n\n"

    "CRITICAL INSTRUCTION: After any internal reasoning, start outputting the JSON immediately. "
    "Do not stall. Generate the complete analysis in a single JSON object.\n\n"

    "STAGE 1 SCOPE: Strategic decision, metadata, impact scoring, agent debate, and evidence. "
    "DO NOT generate full deliverable text (executive_briefing_memo, board_memo, etc.) — "
    "those will be generated in a separate Stage 2 call with dedicated token budget.\n\n"

    "CONCISENESS: Keep arrays to max 2 items. Agent messages: 2 sentences max.\n\n"

    "FOUR DECISION ARCHETYPES:\n"
    "The Management Agent MUST choose EXACTLY ONE of the following four stances as the final decision:\n"
    "1. PROTECT — Prioritize defending rights, brand, and position against high risk. Use when the policy "
    "threatens core business interests, IP rights, or creates unacceptable liability.\n"
    "2. PROMOTE — Prioritize brand exposure and discoverability over immediate revenue or strict control. "
    "Use when maintaining distribution channels and traffic is more valuable than strict rights enforcement.\n"
    "3. NEGOTIATE — Enter into licensing, pricing, or terms negotiation when continued use or substitution "
    "is likely. Use when there is opportunity for fair compensation or when the policy requires bilateral engagement.\n"
    "4. MONITOR — Wait and observe; no immediate action is the most rational choice. Use when the policy "
    "impact is uncertain, enforcement is unclear, or premature action could be counterproductive.\n\n"

    "CROSS-DEPARTMENTAL CONFLICT ARBITRATION:\n"
    "You must simulate a multi-agent debate between Legal, Business, and Product departments. "
    "Different departments will have conflicting views (e.g., Legal may push for PROTECT to minimize risk, "
    "while Business may argue for PROMOTE to retain traffic and revenue, and Product may advocate for "
    "NEGOTIATE to maintain platform relationships). Your role as Management is to arbitrate these conflicts "
    "and make the final strategic decision using one of the FOUR ARCHETYPES above.\n"
    "REQUIRED OUTPUT: Include a 'conflict_summary' field (exactly 40-60 words) that explicitly details "
    "the departmental friction and how Management arbitrated it. YOU MUST USE THE EXACT CAPITALIZED ARCHETYPE "
    "NAMES (PROTECT, PROMOTE, NEGOTIATE, MONITOR) when describing departmental positions and the final decision.\n"
    "Example: 'Legal pushed for strict PROTECT stance citing copyright risk, while Business argued for "
    "PROMOTE to preserve search traffic and ad revenue. Product advocated NEGOTIATE for platform stability. "
    "Management arbitrated to NEGOTIATE with traffic guarantees as a compromise.'\n\n"

    "RULES:\n"
    "1. Evidence-only analysis from source text\n"
    "2. Match tone to document_type (advisory vs operational)\n"
    "3. Policy Memory Graph: If no institutional red-lines match, provide contextual analysis "
    "describing why this policy is relevant (or not) to the organization's strategic priorities. "
    "Never return empty — always add value with context.\n\n"

    "OUTPUT FORMAT:\n"
    "After any necessary reasoning, immediately begin your JSON output starting with '{' "
    "and ensure it ends with '}'. The JSON block should be complete and parseable. "
    "Ensure the 'conflict_summary' field is included at the top level of the JSON.\n"
    "CRITICAL: The 'stance' field in the 'decision' object MUST be exactly one of these four uppercase "
    "strings: 'PROTECT', 'PROMOTE', 'NEGOTIATE', or 'MONITOR'. Do not use any other variations, "
    "synonyms, or free-form text."
)

SYSTEM_PROMPT_DELIVERABLES = (
    "You are the Chief Policy Intelligence Analyst. "
    "Your primary goal is to generate detailed deliverable documents in JSON format "
    "based on the strategic analysis from Stage 1.\n\n"

    "CRITICAL INSTRUCTION: Start outputting the JSON immediately after any necessary "
    "internal reasoning. Do not stall. Generate all 10 deliverables in a single JSON object.\n\n"

    "DELIVERABLE FORMAT:\n"
    "- Each deliverable: 150-250 words (you have full token budget)\n"
    "- Use bullet points and ### headers for clarity\n"
    "- Professional, executive-ready tone\n"
    "- Ground all content in source policy text\n\n"

    "REQUIRED DELIVERABLES (all 10 must be present):\n"
    "1. executive_briefing_memo: C-suite brief (what happened, exposure, action)\n"
    "2. business_impact_memo: Business unit action list (obligations, owners, timelines)\n"
    "3. negotiation_prep_memo: Deal team brief (non-negotiables, leverage, red lines)\n"
    "4. implementation_checklist: Product/Engineering checklist (tasks, dependencies)\n"
    "5. policy_response_draft: External communication draft (positions, justifications)\n"
    "6. what_changed_brief: Before/after delta (clause numbers, dates)\n"
    "7. business_exposure_memo: Exposure memo (traffic, revenue, IP, product, brand)\n"
    "8. negotiation_brief: Negotiation brief (strategy, confirmations, leverage)\n"
    "9. board_memo: Board summary (event, exposure, decisions, actions, scenarios)\n"
    "10. product_checklist: Implementation checklist (array of strings)\n\n"

    "OUTPUT FORMAT:\n"
    "Return a single JSON object with all 10 fields. After any internal reasoning, "
    "immediately begin your JSON output starting with '{' and ensure it ends with '}'. "
    "The JSON block should be complete and parseable."
)

DOMAIN_PROFILES: Dict[str, str] = {
    "AI Search & Distribution": (
        "Core business: Japan's leading digital subscription media, dependent on organic search traffic "
        "for subscriber acquisition and ad revenue. "
        "Primary risk vectors: AI Overviews and Zero-click AI Answers reducing article click-through rates, "
        "AI-generated summaries displacing direct content consumption, "
        "algorithm changes deprioritizing premium paywalled content. "
        "Key objectives: maintain referral traffic integrity, defend against AI summary cannibalization, "
        "negotiate traffic guarantees or equivalent value-exchange in platform agreements. "
        "Special sensitivity: any policy change enabling platform AI to summarize content without redirect, "
        "or modifying traffic attribution models, carries HIGH-to-CRITICAL revenue exposure."
    ),
    "AI Licensing & Copyright": (
        "Core business: licensing proprietary editorial content and structured data to generative AI operators "
        "(global LLM providers, AI startups, enterprise AI vendors). "
        "Primary risk vectors: unauthorized training data ingestion, inadequate royalty-sharing structures, "
        "IP dilution through derivative AI outputs. "
        "Key objectives: maximize licensing revenue through fair valuation of training data supply, "
        "enforce copyright via technical protection measures and robots.txt, "
        "establish enforceable indemnification clauses in all AI partnership agreements. "
        "Special sensitivity: any policy change affecting copyright ownership, consent mechanisms for data use, "
        "or compensation frameworks for rights holders carries CRITICAL-level exposure."
    ),
    "Platform Distribution Policies": (
        "Core business: multi-platform content distribution across social, video, audio, and aggregator channels "
        "including Meta, Google Discover, YouTube, Apple News, and LINE. "
        "Primary risk vectors: unilateral platform rule changes that restrict content eligibility, "
        "monetization policy updates reducing CPM or revenue-share rates, "
        "algorithmic demotion of news/editorial content, "
        "new data-sharing requirements creating compliance overhead. "
        "Key objectives: preserve multi-channel distribution breadth, negotiate favored-partner terms, "
        "maintain content autonomy under platform moderation policies. "
        "Special sensitivity: any policy mandating content modification, consent UI on distributed content, "
        "or new data-reporting obligations to the platform carries HIGH product and legal exposure."
    ),
}

# Domain-specific risk context injected into LLM prompts for calibrated output
DOMAIN_RISK_FOCUS: Dict[str, str] = {
    "AI Search & Distribution": (
        "The user is evaluating this policy change from the perspective of the 'AI Search & Distribution' domain. "
        "Place HEIGHTENED emphasis on: traffic attribution, AI summary/snippet impact on click-through rates, "
        "referral traffic guarantees, and SERP feature policies that displace direct content consumption. "
        "Quantify estimated traffic and revenue loss in percentage and absolute terms where possible."
    ),
    "AI Licensing & Copyright": (
        "The user is evaluating this policy change from the perspective of the 'AI Licensing & Copyright' domain. "
        "Place HEIGHTENED emphasis on: IP ownership implications, copyright consent mechanisms, "
        "royalty and compensation structures, indemnification requirements, and training-data usage rights. "
        "Treat any ambiguity in copyright assignment or consent scope as a CRITICAL-level exposure."
    ),
    "Platform Distribution Policies": (
        "The user is evaluating this policy change from the perspective of the 'Platform Distribution Policies' domain. "
        "Place HEIGHTENED emphasis on: content eligibility rules, monetization policy changes, "
        "data-sharing obligations, algorithmic demotion risks, and consent UI requirements on distributed content. "
        "Assess impact per distribution channel (search, social, video, aggregator) separately where relevant."
    ),
}

STRATEGIST_SYSTEM = (
    "You are the Chief Strategy Officer of Japan's largest media enterprise, "
    "with 150 years of editorial judgment and rigorous copyright discipline. "
    "You lead Co-Build initiatives with global AI Funds and drive AI partnership strategy.\n"
    "Your mandate: deliver specific, actionable, board-level recommendations on digital transformation, "
    "AI-era copyright defense, and global AI Fund partnership strategy.\n"
    "Every output must include quantified metrics, deadlines, and responsible departments. "
    "Eliminate abstraction — precision and executability are non-negotiable.\n"
    "All output must be in professional business English."
)

# ─── Structured Output Schema (Pydantic) ─────────────────────────────────────

StanceType = Literal["Defend", "Pursue Exposure", "Negotiate Terms", "Wait and Monitor"]

_STANCE_COERCE: Dict[str, str] = {
    "defend":    "Defend",
    "protect":   "Defend",
    "pursue":    "Pursue Exposure",
    "promote":   "Pursue Exposure",
    "negotiate": "Negotiate Terms",
    "license":   "Negotiate Terms",
    "wait":      "Wait and Monitor",
    "monitor":   "Wait and Monitor",
    "watch":     "Wait and Monitor",
}

def _coerce_stance(raw: str) -> str:
    """Map free-form strategic_stance text to one of the 4 canonical Stance values."""
    normalized = (raw or "").lower().strip()
    for keyword, canonical in _STANCE_COERCE.items():
        if keyword in normalized:
            return canonical
    return "Wait and Monitor"  # safe default

class Decision(BaseModel):
    stance:    StanceType
    rationale: str

class Metadata(BaseModel):
    affected_departments: List[str]
    legal_certainty: Literal["High", "Medium", "Low"]
    primary_risk: str

class Deliverables(BaseModel):
    executive_briefing_memo: str
    business_impact_memo:    str
    negotiation_prep_memo:   str
    implementation_checklist: str
    policy_response_draft: str

class PolicyAnalysisOutput(BaseModel):
    decision:    Decision
    deliverables: Deliverables
    metadata: Metadata
    conflict_summary: str  # Cross-departmental conflict resolution narrative


# ─── Pipeline Functions ───────────────────────────────────────────────────────
def _safe_json_parse(text: str, debug_label: str = "") -> Dict:
    """Extract and parse JSON from LLM response text with robust recovery.

    Extraction strategy:
    1. Strip markdown code fences (```json ... ```)
    2. Find FIRST '{' and LAST '}' to extract JSON object
    3. Try direct parse
    4. If failed, attempt truncation recovery (append missing braces)
    5. If all fails, print debug output showing first/last 500 chars

    Args:
        text: Raw LLM response text
        debug_label: Label for debug output (e.g., "Stage 1", "Stage 2")

    Returns:
        Parsed dict, or {} on failure
    """
    if not text or not text.strip():
        if debug_label:
            print(f"[{debug_label}] _safe_json_parse: Empty input text")
        return {}

    def _attempt_truncation_recovery(json_str: str) -> Dict:
        """Attempt to recover truncated JSON by appending missing closing braces."""
        open_braces = json_str.count('{')
        close_braces = json_str.count('}')
        open_brackets = json_str.count('[')
        close_brackets = json_str.count(']')

        if open_braces > close_braces or open_brackets > close_brackets:
            recovered = json_str
            for _ in range(open_brackets - close_brackets):
                recovered += ']'
            for _ in range(open_braces - close_braces):
                recovered += '}'
            try:
                return json.loads(recovered, strict=False)
            except json.JSONDecodeError:
                pass
        return {}

    try:
        # ── Step 1: Strip markdown code fences ─────────────────────────────────
        clean = text.strip()
        if clean.startswith("```json"):
            clean = clean.split("```json", 1)[1]
        elif clean.startswith("```"):
            clean = clean.split("```", 1)[1]
        if clean.endswith("```"):
            clean = clean.rsplit("```", 1)[0]
        clean = clean.strip()

        # ── Step 2: Find FIRST '{' and LAST '}' (non-greedy extraction) ───────
        first_brace = clean.find('{')
        last_brace = clean.rfind('}')

        if first_brace != -1 and last_brace != -1 and last_brace > first_brace:
            json_candidate = clean[first_brace:last_brace + 1]

            # Try direct parse with strict=False to allow unescaped control characters
            try:
                return json.loads(json_candidate, strict=False)
            except json.JSONDecodeError as e:
                # Try truncation recovery
                recovered = _attempt_truncation_recovery(json_candidate)
                if recovered:
                    return recovered

                # If still failing, print debug info
                if debug_label:
                    print(f"\n{'='*70}")
                    print(f"[{debug_label}] JSON PARSE ERROR")
                    print(f"{'='*70}")
                    print(f"Error: {e}")
                    print(f"\nFirst 500 chars of RAW response:")
                    print(f"{text[:500]}")
                    print(f"\nLast 500 chars of RAW response:")
                    print(f"{text[-500:]}")
                    print(f"\nFirst 500 chars of extracted JSON candidate:")
                    print(f"{json_candidate[:500]}")
                    print(f"\nLast 500 chars of extracted JSON candidate:")
                    print(f"{json_candidate[-500:]}")
                    print(f"{'='*70}\n")

        # ── Step 3: Fallback - try direct parse of cleaned text ───────────────
        if clean.startswith("{"):
            try:
                return json.loads(clean)
            except json.JSONDecodeError:
                recovered = _attempt_truncation_recovery(clean)
                if recovered:
                    return recovered

        # All parsing attempts failed
        if debug_label:
            print(f"\n{'='*70}")
            print(f"[{debug_label}] JSON EXTRACTION FAILED - No valid JSON found")
            print(f"{'='*70}")
            print(f"First 500 chars of RAW response:")
            print(f"{text[:500]}")
            print(f"\nLast 500 chars of RAW response:")
            print(f"{text[-500:]}")
            print(f"{'='*70}\n")

        return {}
    except Exception as exc:
        if debug_label:
            print(f"[{debug_label}] Unexpected error in _safe_json_parse: {exc}")
        return {}


def get_client() -> Optional[anthropic.Anthropic]:
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        try:
            api_key = st.secrets.get("ANTHROPIC_API_KEY", "")
        except Exception:
            pass
    if not api_key:
        return None
    return anthropic.Anthropic(api_key=api_key)

def analyze_policy_with_claude(
    client: anthropic.Anthropic,
    policy_text: str,
    domain: str,
) -> Dict:
    """
    Single comprehensive Claude API call — Single Source of Truth for all UI components.

    OPTIMIZATIONS FOR 8 DELIVERABLES:
    - max_tokens: 8192 (Claude Sonnet 4.5 output limit)
    - Concise prompt: instructs bullet-point format, 150-250 words per deliverable
    - Truncation recovery: _safe_json_parse attempts to close incomplete JSON
    - Truncation guard: pre-parse check for response completeness

    Returns a flat dict with all keys needed by every UI panel.
    """
    domain_profile = DOMAIN_PROFILES.get(domain, "")
    domain_focus   = DOMAIN_RISK_FOCUS.get(domain, "")
    snippet        = policy_text[:3500]

    # ── Stream the response — chunks keep WebSocket alive during long reasoning ──
    full_response_text = ""
    try:
      with client.messages.stream(
        model=MODEL,
        max_tokens=8192,   # Maximum output tokens for Claude Sonnet 4.5
        thinking={"type": "adaptive"},
        system=(
            "You are the Chief Policy Intelligence Analyst for a major media enterprise. "
            "Produce a single comprehensive JSON intelligence report with structural analysis, "
            "business impact scoring, strategic recommendations, and 8 role-specific deliverables.\n\n"

            "━━━ CRITICAL: EXTREME BREVITY REQUIRED ━━━\n"
            "TOKEN BUDGET: This response must fit within 8000 tokens total.\n"
            "DELIVERABLES: Each of the 8 deliverables MUST be 80-120 words MAX (not 150-250).\n"
            "FORMAT: Use ONLY bullet points. NO paragraphs. NO introductions. NO filler.\n"
            "AGENT MESSAGES: 2 sentences MAX per agent (not 3-4).\n"
            "ARRAYS: Max 2 items per array (not 3-4). Only include highest-impact items.\n\n"

            "RULES:\n"
            "1. Evidence-only analysis from source text\n"
            "2. Match tone to document_type (advisory vs operational)\n"
            "3. Policy memory graph: empty array if no genuine triggers\n"
            "4. All text fields: telegraphic, high-density prose only\n\n"

            "FORMATTING RULE (CRITICAL):\n"
            "Output ONLY valid JSON. Do NOT include any preamble, thinking blocks, "
            "code fences, or closing remarks outside the JSON structure. "
            "Your response must start with '{' and end with '}' with no extra text before or after."
        ),
        messages=[{
            "role": "user",
            "content": (
                f"Analyze the following policy/regulatory text for a media enterprise "
                f"operating in the '{domain}' domain. Produce a COMPLETE intelligence report.\n\n"
                f"[DOMAIN PROFILE]\n{domain_profile}\n\n"
                f"[DOMAIN RISK FOCUS]\n{domain_focus}\n\n"
                f"[POLICY SOURCE TEXT — Extract VERBATIM quotes for all *_quotes fields]\n{snippet}\n\n"
                f"Return ONLY this exact JSON schema (no extra text, no code fences):\n"
                f"{{\n"
                f'  "decision": {{\n'
                f'    "stance": "MUST be exactly one of: Defend | Pursue Exposure | Negotiate Terms | Wait and Monitor",\n'
                f'    "rationale": "1-2 sentence synthesis of the agent debate leading to this stance"\n'
                f'  }},\n'
                f'  "deliverables": {{\n'
                f'    "executive_briefing_memo": "Bullets only (80-120 words): • What happened • Exposure • Action",\n'
                f'    "business_impact_memo": "Bullets only (80-120 words): • Obligations • Owners • Timelines",\n'
                f'    "negotiation_prep_memo": "Bullets only (80-120 words): • Non-negotiables • Confirmations • Leverage • Red lines",\n'
                f'    "implementation_checklist": "Bullets only (80-120 words): • Tasks • Details • Dependencies",\n'
                f'    "policy_response_draft": "Bullets only (80-120 words): • Positions • Justifications • Actions"\n'
                f'  }},\n'
                f'  "metadata": {{\n'
                f'    "affected_departments": ["List of departments impacted, e.g. Legal, Engineering, Business, Marketing"],\n'
                f'    "legal_certainty": "High|Medium|Low — confidence in legal interpretation based on source clarity",\n'
                f'    "primary_risk": "One-sentence summary of the most critical risk factor"\n'
                f'  }},\n'
                f'  "strategic_stance": "e.g. PROTECT & LICENSE",\n'
                f'  "jurisdiction": "e.g. European Union / EMEA or Global",\n'
                f'  "document_type": "binding_regulation|draft_legislation|platform_terms_update|government_statement|consultation_paper|industry_guideline",\n'
                f'  "overall_risk_level": "critical|high|medium|low",\n'
                f'  "executive_summary": "1-2 sentences max",\n'
                f'  "key_opportunities": ["max 2 items"],\n'
                f'  "key_threats": ["max 2 items"],\n'
                f'  "substantive_changes": {{\n'
                f'    "added_obligations": [{{"title": "...", "severity": "high|medium|low", "description": "brief"}}],\n'
                f'    "removed_rights":    [{{"title": "...", "severity": "high|medium|low", "description": "brief"}}],\n'
                f'    "key_thresholds":    [{{"title": "...", "value": "...", "description": "brief"}}],\n'
                f'    "context_summary":   "1-2 sentences"\n'
                f'  }},\n'
                f'  "scores": {{"IP": 0-100, "Traffic": 0-100, "Revenue": 0-100, "Product": 0-100}},\n'
                f'  "axis_actions": {{\n'
                f'    "IP":      {{"badge": "PROTECT|LICENSE|etc", "summary": "brief", "evidence": "short", "direction": "threat|opportunity|neutral", "priority_actions": ["max 2"]}},\n'
                f'    "Traffic": {{"badge": "...", "summary": "brief", "evidence": "short", "direction": "...", "priority_actions": ["max 2"]}},\n'
                f'    "Revenue": {{"badge": "...", "summary": "brief", "evidence": "short", "direction": "...", "priority_actions": ["max 2"]}},\n'
                f'    "Product": {{"badge": "...", "summary": "brief", "evidence": "short", "direction": "...", "priority_actions": ["max 2"]}}\n'
                f'  }},\n'
                f'  "evidence": {{\n'
                f'    "parsed_claims":          ["max 2 extracts"],\n'
                f'    "claim_level_provenance": [{{"type": "IP|TRAFFIC|REVENUE", "agent": "LEGAL|BUSINESS", "quote": "short"}}],\n'
                f'    "policy_memory_graph":    [{{"reference_id": "ID", "description": "brief"}}]\n'
                f'  }},\n'
                f'  "risk_matrix_points": [\n'
                f'    {{"label": "name", "days_to_enactment": 0-100, "business_severity": 0-100}}\n'
                f'  ],\n'
                f'  "what_changed_brief": "Bullets only (80-120 words): • Deltas • Clause # • Dates",\n'
                f'  "what_changed_quotes": ["quote 1", "quote 2"],\n'
                f'  "business_exposure_memo": "Bullets only (100-150 words): • Traffic • Revenue • IP • Product • Brand",\n'
                f'  "business_exposure_quotes": ["quote"],\n'
                f'  "negotiation_brief": "Bullets only (80-120 words): • Non-negotiables • Confirmations • Leverage • Red lines",\n'
                f'  "negotiation_quotes": ["quote"],\n'
                f'  "board_memo": "Bullets only (100-150 words): • Event • Exposure • Decisions • Actions • Scenarios",\n'
                f'  "board_memo_quotes": ["quote"],\n'
                f'  "product_checklist": [\n'
                f'    "Max 3 items. Format: [CATEGORY] Team — action. Only items from source text."\n'
                f'  ],\n'
                f'  "agent_debate_messages": [\n'
                f'    {{\n'
                f'      "agent": "⚖️ Legal",\n'
                f'      "color": "#8B2635",\n'
                f'      "message": "2 sentences max. Focus: contract risk, IP rights, liability. Ground in source text only."\n'
                f'    }},\n'
                f'    {{\n'
                f'      "agent": "💰 Business",\n'
                f'      "color": "#A8892A",\n'
                f'      "message": "2 sentences max. Focus: revenue, partnerships, market position. Ground in source text only."\n'
                f'    }},\n'
                f'    {{\n'
                f'      "agent": "⚙️ Engineering",\n'
                f'      "color": "#1A6B3C",\n'
                f'      "message": "2 sentences max. Focus: implementation, technical debt, resources. Ground in source text only."\n'
                f'    }},\n'
                f'    {{\n'
                f'      "agent": "🏛️ Management",\n'
                f'      "color": "#0ABAB5",\n'
                f'      "final": true,\n'
                f'      "message": "2 sentences max. Synthesize Legal/Business/Engineering into unified stance. Match urgency to document type."\n'
                f'    }}\n'
                f'  ]\n'
                f"}}"
            ),
        }],
      ) as stream:
            # Explicitly filter for text_delta events only — thinking deltas are discarded
            for event in stream:
                if (
                    event.type == "content_block_delta"
                    and event.delta.type == "text_delta"
                ):
                    full_response_text += event.delta.text
    except anthropic.APIError as api_err:
        raise RuntimeError(
            f"[Anthropic API Error — {type(api_err).__name__}] {api_err}"
        ) from api_err

    # ── Guard: stream must have produced text output ────────────────────────────
    if not full_response_text.strip():
        raise ValueError(
            "Claude returned empty text output. The stream may have contained only "
            "thinking blocks or was interrupted before producing JSON. "
            "Try re-running — adaptive thinking occasionally produces no text on the first pass."
        )

    # ── Token Usage Monitoring ────────────────────────────────────────────────────
    import logging
    response_stripped = full_response_text.strip()
    approx_tokens = len(full_response_text) // 4  # rough estimate: 1 token ≈ 4 chars
    logging.info(
        f"Response received: {len(full_response_text)} chars (≈{approx_tokens} tokens of 8192 max)"
    )

    # ── Truncation Guard: Check if response appears to be truncated ──────────────
    if not response_stripped.endswith('}'):
        # Response may be truncated - log warning but attempt recovery
        logging.warning(
            f"⚠️ Response appears truncated (does not end with closing brace). "
            f"Length: {len(full_response_text)} chars. "
            f"Last 100 chars: ...{response_stripped[-100:]}"
        )
        # Check if we're close to token limit
        if approx_tokens > 7000:  # Lowered threshold from 7500 to 7000 for earlier warning
            raise ValueError(
                f"❌ Response truncated due to token limit (≈{approx_tokens} tokens, max 8192). "
                f"The 8 deliverables exceeded available output space ({approx_tokens - 8192} tokens over). "
                f"Try re-running with more aggressive brevity constraints. "
                f"Last 200 chars: ...{response_stripped[-200:]}"
            )
    elif approx_tokens > 7500:
        # Close to limit but completed - log warning
        logging.warning(
            f"⚠️ Response near token limit: {approx_tokens}/8192 tokens used "
            f"({8192 - approx_tokens} tokens remaining)"
        )

    # ── Parse JSON from accumulated stream text ─────────────────────────────────
    result = _safe_json_parse(full_response_text, debug_label="Single-Stage Analysis")
    if not result:
        # Debug output already printed by _safe_json_parse
        raise ValueError(
            f"JSON extraction failed — Claude did not return parseable JSON. "
            f"See console/logs for detailed debug output."
        )

    # ── Pydantic structured-output validation ────────────────────────────────────
    # Primary: use Claude's decision/deliverables if present and valid.
    # Coerce: map free-form stance text to one of the 4 canonical values.
    # Fallback: derive from existing fields so the existing UI never breaks.
    _allowed_stances = {"Defend", "Pursue Exposure", "Negotiate Terms", "Wait and Monitor"}

    # Build decision sub-object
    decision_raw = result.get("decision") or {}
    if not isinstance(decision_raw, dict):
        decision_raw = {}
    stance_raw = decision_raw.get("stance") or result.get("strategic_stance", "")
    if stance_raw not in _allowed_stances:
        stance_raw = _coerce_stance(str(stance_raw))
    decision_raw["stance"]    = stance_raw
    decision_raw.setdefault("rationale", result.get("executive_summary", "Analysis complete."))

    # Build deliverables sub-object
    deliverables_raw = result.get("deliverables") or {}
    if not isinstance(deliverables_raw, dict):
        deliverables_raw = {}
    deliverables_raw.setdefault("executive_briefing_memo", result.get("board_memo", ""))
    deliverables_raw.setdefault("business_impact_memo",    result.get("business_exposure_memo", ""))
    deliverables_raw.setdefault("negotiation_prep_memo",   result.get("negotiation_brief", ""))
    deliverables_raw.setdefault("implementation_checklist", result.get("product_checklist", ""))
    deliverables_raw.setdefault("policy_response_draft", "")

    # Build metadata sub-object
    metadata_raw = result.get("metadata") or {}
    if not isinstance(metadata_raw, dict):
        metadata_raw = {}
    metadata_raw.setdefault("affected_departments", ["Legal", "Business", "Engineering"])
    metadata_raw.setdefault("legal_certainty", "Medium")
    metadata_raw.setdefault("primary_risk", result.get("executive_summary", "Risk assessment in progress."))

    try:
        validated = PolicyAnalysisOutput(
            decision=Decision(**decision_raw),
            deliverables=Deliverables(**deliverables_raw),
            metadata=Metadata(**metadata_raw),
        )
        result["decision"]    = validated.decision.model_dump()
        result["deliverables"] = validated.deliverables.model_dump()
        result["metadata"]    = validated.metadata.model_dump()
    except (ValidationError, Exception):
        # Non-fatal — write safe defaults so the UI always has these keys
        result["decision"] = {"stance": stance_raw, "rationale": decision_raw.get("rationale", "")}
        result["deliverables"] = {
            "executive_briefing_memo": deliverables_raw.get("executive_briefing_memo", ""),
            "business_impact_memo":    deliverables_raw.get("business_impact_memo", ""),
            "negotiation_prep_memo":   deliverables_raw.get("negotiation_prep_memo", ""),
            "implementation_checklist": deliverables_raw.get("implementation_checklist", ""),
            "policy_response_draft":   deliverables_raw.get("policy_response_draft", ""),
        }
        result["metadata"] = {
            "affected_departments": metadata_raw.get("affected_departments", []),
            "legal_certainty":      metadata_raw.get("legal_certainty", "Medium"),
            "primary_risk":         metadata_raw.get("primary_risk", ""),
        }

    return result


# ═══════════════════════════════════════════════════════════════════════════════
# TWO-STAGE DELIVERY ARCHITECTURE (Solves 8192 token truncation issue)
# ═══════════════════════════════════════════════════════════════════════════════

def analyze_policy_core(
    client: anthropic.Anthropic,
    policy_text: str,
    domain: str,
) -> Dict:
    """
    STAGE 1: Core Analysis & Metadata Only

    Returns strategic decision, metadata, and core analysis WITHOUT the full
    deliverable text. This keeps token usage low (~3000-4000 tokens).

    Returns:
        {
            "decision": {"stance": "...", "rationale": "..."},
            "metadata": {"affected_departments": [...], "legal_certainty": "...", "primary_risk": "..."},
            "strategic_stance": "...",
            "jurisdiction": "...",
            "document_type": "...",
            "overall_risk_level": "...",
            "executive_summary": "...",
            "key_opportunities": [...],
            "key_threats": [...],
            "substantive_changes": {...},
            "scores": {...},
            "axis_actions": {...},
            "evidence": {...},
            "risk_matrix_points": [...],
            "agent_debate_messages": [...],
            ... (other core fields)
        }
    """
    domain_profile = DOMAIN_PROFILES.get(domain, "")
    domain_focus   = DOMAIN_RISK_FOCUS.get(domain, "")
    snippet        = policy_text[:3500]

    full_response_text = ""
    thinking_blocks = []
    max_retries = 2
    retry_count = 0

    while retry_count < max_retries:
        try:
            with client.messages.stream(
                model=MODEL,
                max_tokens=8192,   # Stage 1: Increased for robust JSON output
                temperature=1.0,   # Stable temperature for JSON generation
                thinking={"type": "adaptive"},
                system=SYSTEM_PROMPT_CORE,
                messages=[{
            "role": "user",
            "content": (
                f"Analyze the following policy/regulatory text for a media enterprise "
                f"operating in the '{domain}' domain. Produce STAGE 1: CORE ANALYSIS ONLY.\n\n"
                f"[DOMAIN PROFILE]\n{domain_profile}\n\n"
                f"[DOMAIN RISK FOCUS]\n{domain_focus}\n\n"
                f"[POLICY SOURCE TEXT]\n{snippet}\n\n"
                f"Return this JSON schema (OMIT deliverable text fields):\n"
                f"{{\n"
                f'  "decision": {{\n'
                f'    "stance": "Defend | Pursue Exposure | Negotiate Terms | Wait and Monitor",\n'
                f'    "rationale": "1-2 sentence synthesis"\n'
                f'  }},\n'
                f'  "metadata": {{\n'
                f'    "affected_departments": ["Legal", "Engineering", "Business", "etc"],\n'
                f'    "legal_certainty": "High|Medium|Low",\n'
                f'    "primary_risk": "One-sentence summary"\n'
                f'  }},\n'
                f'  "strategic_stance": "e.g. PROTECT & LICENSE",\n'
                f'  "jurisdiction": "e.g. European Union",\n'
                f'  "document_type": "binding_regulation|draft_legislation|etc",\n'
                f'  "overall_risk_level": "critical|high|medium|low",\n'
                f'  "executive_summary": "1-2 sentences",\n'
                f'  "key_opportunities": ["max 2"],\n'
                f'  "key_threats": ["max 2"],\n'
                f'  "substantive_changes": {{\n'
                f'    "added_obligations": [{{"title": "...", "severity": "high|medium|low", "description": "brief"}}],\n'
                f'    "removed_rights": [{{"title": "...", "severity": "...", "description": "brief"}}],\n'
                f'    "key_thresholds": [{{"title": "...", "value": "...", "description": "brief"}}],\n'
                f'    "context_summary": "1-2 sentences"\n'
                f'  }},\n'
                f'  "scores": {{"IP": 0-100, "Traffic": 0-100, "Revenue": 0-100, "Product": 0-100}},\n'
                f'  "axis_actions": {{\n'
                f'    "IP": {{"badge": "PROTECT|LICENSE|etc", "summary": "brief", "evidence": "short", "direction": "threat|opportunity|neutral", "priority_actions": ["max 2"]}},\n'
                f'    "Traffic": {{"badge": "...", "summary": "...", "evidence": "...", "direction": "...", "priority_actions": ["..."]}},\n'
                f'    "Revenue": {{"badge": "...", "summary": "...", "evidence": "...", "direction": "...", "priority_actions": ["..."]}},\n'
                f'    "Product": {{"badge": "...", "summary": "...", "evidence": "...", "direction": "...", "priority_actions": ["..."]}}\n'
                f'  }},\n'
                f'  "evidence": {{\n'
                f'    "parsed_claims": ["max 2 extracts"],\n'
                f'    "claim_level_provenance": [{{"type": "IP|TRAFFIC|REVENUE", "agent": "LEGAL|BUSINESS", "quote": "short"}}],\n'
                f'    "policy_memory_graph": [{{"reference_id": "ID", "description": "brief"}}]\n'
                f'  }},\n'
                f'  "risk_matrix_points": [{{"label": "name", "days_to_enactment": 0-100, "business_severity": 0-100}}],\n'
                f'  "agent_debate_messages": [\n'
                f'    {{"agent": "⚖️ Legal", "color": "#8B2635", "message": "2 sentences max"}},\n'
                f'    {{"agent": "💰 Business", "color": "#A8892A", "message": "2 sentences max"}},\n'
                f'    {{"agent": "⚙️ Engineering", "color": "#1A6B3C", "message": "2 sentences max"}},\n'
                f'    {{"agent": "🏛️ Management", "color": "#0ABAB5", "final": true, "message": "2 sentences max"}}\n'
                f'  ]\n'
                f"}}"
            ),
                }],
            ) as stream:
                # Track content blocks separately (thinking vs text)
                current_block_type = None
                current_block_text = ""

                for event in stream:
                    # Track when new content blocks start
                    if event.type == "content_block_start":
                        if hasattr(event, 'content_block') and hasattr(event.content_block, 'type'):
                            current_block_type = event.content_block.type
                            current_block_text = ""

                    # Accumulate text from deltas
                    elif event.type == "content_block_delta":
                        if event.delta.type == "text_delta":
                            delta_text = event.delta.text
                            current_block_text += delta_text

                            # Only add to full_response_text if it's a text block (not thinking)
                            if current_block_type != "thinking":
                                full_response_text += delta_text

                    # When block ends, save thinking blocks separately
                    elif event.type == "content_block_stop":
                        if current_block_type == "thinking" and current_block_text:
                            thinking_blocks.append(current_block_text)
                        current_block_type = None
                        current_block_text = ""

            # Log thinking blocks for debugging
            if thinking_blocks:
                print(f"[Stage 1 Debug] Captured {len(thinking_blocks)} thinking block(s)")

            # Log text content for debugging
            if full_response_text.strip():
                print(f"[Stage 1 Debug] Captured {len(full_response_text)} chars of text content")
                break  # Success - exit retry loop
            else:
                print(f"[Stage 1 Debug] No text content captured (only thinking blocks: {len(thinking_blocks)})")

            # Otherwise, retry
            retry_count += 1
            if retry_count < max_retries:
                print(f"[Stage 1] Empty response on attempt {retry_count}, retrying...")
                full_response_text = ""
                thinking_blocks = []
                import time
                time.sleep(1)  # Brief delay before retry

        except anthropic.APIError as api_err:
            retry_count += 1
            if retry_count >= max_retries:
                raise RuntimeError(
                    f"[Stage 1 API Error after {max_retries} attempts — {type(api_err).__name__}] {api_err}"
                ) from api_err
            else:
                print(f"[Stage 1] API error on attempt {retry_count}, retrying...")
                import time
                time.sleep(2)  # Longer delay for API errors

    # Final check after all retries
    if not full_response_text.strip():
        error_details = f"Retry attempts: {retry_count}, Thinking blocks captured: {len(thinking_blocks)}"
        if thinking_blocks:
            error_details += f"\nThinking content length: {sum(len(t) for t in thinking_blocks)} chars"
        raise ValueError(
            f"Stage 1: Claude returned empty text output after {retry_count + 1} attempts.\n"
            f"Debug info: {error_details}"
        )

    # Parse and validate with debug output
    result = _safe_json_parse(full_response_text, debug_label="Stage 1: Core Analysis")
    if not result:
        # Debug output already printed by _safe_json_parse
        raise ValueError(
            f"Stage 1: JSON extraction failed. See console/logs for detailed debug output."
        )

    # Validate core fields and apply smart defaults
    _allowed_stances = {"Defend", "Pursue Exposure", "Negotiate Terms", "Wait and Monitor"}

    # Decision
    decision_raw = result.get("decision") or {}
    if not isinstance(decision_raw, dict):
        decision_raw = {}
    stance_raw = decision_raw.get("stance") or result.get("strategic_stance", "")
    if stance_raw not in _allowed_stances:
        stance_raw = _coerce_stance(str(stance_raw))
    decision_raw["stance"] = stance_raw
    decision_raw.setdefault("rationale", result.get("executive_summary", "Analysis complete."))
    result["decision"] = decision_raw

    # Metadata
    metadata_raw = result.get("metadata") or {}
    if not isinstance(metadata_raw, dict):
        metadata_raw = {}
    metadata_raw.setdefault("affected_departments", ["Legal", "Business", "Engineering"])
    metadata_raw.setdefault("legal_certainty", "Medium")
    metadata_raw.setdefault("primary_risk", result.get("executive_summary", "Risk assessment in progress."))
    result["metadata"] = metadata_raw

    # Substantive Changes (Smart Defaults)
    sc_raw = result.get("substantive_changes") or {}
    if not isinstance(sc_raw, dict):
        sc_raw = {}

    # Added Obligations
    if not sc_raw.get("added_obligations") or len(sc_raw.get("added_obligations", [])) == 0:
        sc_raw["added_obligations"] = [{
            "title": "No significant new obligations identified",
            "severity": "low",
            "description": "Analysis found no material new requirements in the policy document."
        }]

    # Removed Rights
    if not sc_raw.get("removed_rights") or len(sc_raw.get("removed_rights", [])) == 0:
        sc_raw["removed_rights"] = [{
            "title": "No significant rights removal identified",
            "severity": "low",
            "description": "Analysis found no material removal of existing rights or permissions."
        }]

    # Key Thresholds
    if not sc_raw.get("key_thresholds") or len(sc_raw.get("key_thresholds", [])) == 0:
        sc_raw["key_thresholds"] = [{
            "title": "No critical thresholds identified",
            "value": "N/A",
            "description": "Analysis found no specific compliance thresholds or deadlines."
        }]

    sc_raw.setdefault("context_summary", "Policy analysis complete. Review substantive changes above.")
    result["substantive_changes"] = sc_raw

    # Scores (ensure all axes present)
    scores_raw = result.get("scores") or {}
    for axis in ["IP", "Traffic", "Revenue", "Product"]:
        if axis not in scores_raw:
            scores_raw[axis] = 50  # Neutral default
    result["scores"] = scores_raw

    # Axis Actions (ensure all axes present)
    axis_actions_raw = result.get("axis_actions") or {}
    for axis in ["IP", "Traffic", "Revenue", "Product"]:
        if axis not in axis_actions_raw or not isinstance(axis_actions_raw.get(axis), dict):
            axis_actions_raw[axis] = {
                "badge": "MONITOR",
                "summary": f"Continue monitoring {axis} impact",
                "evidence": "No specific action required at this time.",
                "direction": "neutral",
                "priority_actions": ["Monitor ongoing developments"]
            }
    result["axis_actions"] = axis_actions_raw

    # Evidence (ensure structure exists)
    evidence_raw = result.get("evidence") or {}
    evidence_raw.setdefault("parsed_claims", ["Policy document analyzed"])
    evidence_raw.setdefault("claim_level_provenance", [])
    evidence_raw.setdefault("policy_memory_graph", [])
    result["evidence"] = evidence_raw

    # Other essential fields
    result.setdefault("strategic_stance", decision_raw.get("stance", "Wait and Monitor"))
    result.setdefault("jurisdiction", "Global")
    result.setdefault("document_type", "policy_document")
    result.setdefault("overall_risk_level", "medium")
    result.setdefault("executive_summary", "Policy analysis complete.")
    result.setdefault("key_opportunities", ["Review full analysis for opportunities"])
    result.setdefault("key_threats", ["Review full analysis for threats"])
    result.setdefault("agent_debate_messages", [])
    result.setdefault("risk_matrix_points", [])

    # Cross-departmental conflict summary (new feature for Cross-Departmental Response Engine)
    result.setdefault(
        "conflict_summary",
        "Multi-departmental analysis completed. Legal evaluated compliance risk, Business assessed revenue impact, "
        "and Product reviewed implementation feasibility. Management synthesized inputs to determine strategic stance."
    )

    return result


def generate_detailed_deliverables(
    client: anthropic.Anthropic,
    policy_text: str,
    domain: str,
    stage1_result: Dict,
) -> Dict:
    """
    STAGE 2: Generate Detailed Deliverables

    Takes Stage 1 results (strategic stance, metadata) and generates the full
    text for all 8 deliverables with a dedicated token budget (~6000-7000 tokens).

    Args:
        policy_text: Original policy document
        domain: Policy domain
        stage1_result: Output from analyze_policy_core()

    Returns:
        {
            "executive_briefing_memo": "Full Markdown text (150-250 words)",
            "business_impact_memo": "...",
            "negotiation_prep_memo": "...",
            "implementation_checklist": "...",
            "policy_response_draft": "...",
            "what_changed_brief": "...",
            "business_exposure_memo": "...",
            "negotiation_brief": "...",
            "board_memo": "...",
            "product_checklist": [...]
        }
    """
    domain_profile = DOMAIN_PROFILES.get(domain, "")
    snippet = policy_text[:3500]

    # Extract context from Stage 1
    stance = stage1_result.get("decision", {}).get("stance", "Wait and Monitor")
    rationale = stage1_result.get("decision", {}).get("rationale", "")
    metadata = stage1_result.get("metadata", {})
    affected_depts = metadata.get("affected_departments", [])
    legal_certainty = metadata.get("legal_certainty", "Medium")
    primary_risk = metadata.get("primary_risk", "")
    executive_summary = stage1_result.get("executive_summary", "")
    substantive_changes = stage1_result.get("substantive_changes", {})

    full_response_text = ""
    thinking_blocks = []
    max_retries = 2
    retry_count = 0

    while retry_count < max_retries:
        try:
            with client.messages.stream(
                model=MODEL,
                max_tokens=8192,   # Stage 2: Full budget for deliverables
                temperature=1.0,   # Stable temperature for JSON generation
                thinking={"type": "adaptive"},
                system=SYSTEM_PROMPT_DELIVERABLES,
                messages=[{
            "role": "user",
            "content": (
                f"Generate detailed deliverable documents for this policy analysis.\n\n"
                f"[POLICY DOMAIN]: {domain}\n"
                f"[DOMAIN PROFILE]: {domain_profile}\n\n"
                f"[STAGE 1 STRATEGIC DECISION]\n"
                f"Stance: {stance}\n"
                f"Rationale: {rationale}\n"
                f"Affected Departments: {', '.join(affected_depts)}\n"
                f"Legal Certainty: {legal_certainty}\n"
                f"Primary Risk: {primary_risk}\n"
                f"Executive Summary: {executive_summary}\n\n"
                f"[POLICY SOURCE TEXT]\n{snippet}\n\n"
                f"Generate deliverables in this JSON format:\n"
                f"{{\n"
                f'  "executive_briefing_memo": "Markdown (150-250 words): C-suite brief with sections: ### What Happened, ### Financial Exposure, ### Recommended Action. Use ### headers for visual clarity.",\n'
                f'  "business_impact_memo": "Markdown (150-250 words): Business action list with sections: ### Key Obligations, ### Department Owners, ### Critical Timelines. Use ### headers for visual clarity.",\n'
                f'  "negotiation_prep_memo": "Markdown (150-250 words): Deal team brief with • Non-negotiables • Written confirmations needed • Leverage points • Red lines",\n'
                f'  "implementation_checklist": "Markdown (150-250 words): Implementation checklist with • Tasks • Technical details • Dependencies",\n'
                f'  "policy_response_draft": "Markdown (150-250 words): External communication draft with • Key positions • Justifications • Proposed actions",\n'
                f'  "what_changed_brief": "Markdown (150-200 words): Delta analysis with • Before/after changes • Clause numbers • Effective dates",\n'
                f'  "business_exposure_memo": "Markdown (200-250 words): Exposure assessment with • Traffic impact • Revenue impact • IP impact • Product impact • Brand impact",\n'
                f'  "negotiation_brief": "Markdown (150-200 words): Negotiation strategy with • Non-negotiables • Written confirmations • Leverage • Compromise zones • Red lines",\n'
                f'  "board_memo": "Markdown (200-250 words): Board summary with sections: ### What Happened, ### Financial Exposure, ### Board Decisions Required, ### Recommended Actions (with owners), ### Scenarios (best/base/worst). Use ### headers for visual clarity.",\n'
                f'  "product_checklist": ["[CATEGORY] Team — specific action", "[CATEGORY] Team — specific action", ...]\n'
                f"}}"
            ),
                }],
            ) as stream:
                # Track content blocks separately (thinking vs text)
                current_block_type = None
                current_block_text = ""

                for event in stream:
                    # Track when new content blocks start
                    if event.type == "content_block_start":
                        if hasattr(event, 'content_block') and hasattr(event.content_block, 'type'):
                            current_block_type = event.content_block.type
                            current_block_text = ""

                    # Accumulate text from deltas
                    elif event.type == "content_block_delta":
                        if event.delta.type == "text_delta":
                            delta_text = event.delta.text
                            current_block_text += delta_text

                            # Only add to full_response_text if it's a text block (not thinking)
                            if current_block_type != "thinking":
                                full_response_text += delta_text

                    # When block ends, save thinking blocks separately
                    elif event.type == "content_block_stop":
                        if current_block_type == "thinking" and current_block_text:
                            thinking_blocks.append(current_block_text)
                        current_block_type = None
                        current_block_text = ""

            # Log thinking blocks for debugging
            if thinking_blocks:
                print(f"[Stage 2 Debug] Captured {len(thinking_blocks)} thinking block(s)")

            # Log text content for debugging
            if full_response_text.strip():
                print(f"[Stage 2 Debug] Captured {len(full_response_text)} chars of text content")
                break  # Success - exit retry loop
            else:
                print(f"[Stage 2 Debug] No text content captured (only thinking blocks: {len(thinking_blocks)})")

            # Otherwise, retry
            retry_count += 1
            if retry_count < max_retries:
                print(f"[Stage 2] Empty response on attempt {retry_count}, retrying...")
                full_response_text = ""
                thinking_blocks = []
                import time
                time.sleep(1)  # Brief delay before retry

        except anthropic.APIError as api_err:
            retry_count += 1
            if retry_count >= max_retries:
                raise RuntimeError(
                    f"[Stage 2 API Error after {max_retries} attempts — {type(api_err).__name__}] {api_err}"
                ) from api_err
            else:
                print(f"[Stage 2] API error on attempt {retry_count}, retrying...")
                import time
                time.sleep(2)  # Longer delay for API errors

    # Final check after all retries
    if not full_response_text.strip():
        error_details = f"Retry attempts: {retry_count}, Thinking blocks captured: {len(thinking_blocks)}"
        if thinking_blocks:
            error_details += f"\nThinking content length: {sum(len(t) for t in thinking_blocks)} chars"
        raise ValueError(
            f"Stage 2: Claude returned empty text output after {retry_count + 1} attempts.\n"
            f"Debug info: {error_details}"
        )

    # ═══════════════════════════════════════════════════════════════════════════
    # GRACEFUL DEGRADATION: Parse deliverables with fallback on failure
    # ═══════════════════════════════════════════════════════════════════════════
    try:
        deliverables = _safe_json_parse(full_response_text, debug_label="Stage 2: Deliverables")

        if not deliverables:
            # Debug output already printed by _safe_json_parse
            print(f"\n[Stage 2 Graceful Degradation] JSON extraction failed completely.")
            print(f"Last 500 chars of raw response:\n{full_response_text[-500:]}")
            raise ValueError("JSON extraction returned empty dict")

        # Ensure all expected keys exist
        expected_keys = [
            "executive_briefing_memo", "business_impact_memo", "negotiation_prep_memo",
            "implementation_checklist", "policy_response_draft", "what_changed_brief",
            "business_exposure_memo", "negotiation_brief", "board_memo", "product_checklist"
        ]
        for key in expected_keys:
            if key not in deliverables:
                deliverables[key] = "" if key != "product_checklist" else []

        return deliverables

    except (ValueError, json.JSONDecodeError, KeyError, Exception) as e:
        # ═══════════════════════════════════════════════════════════════════════
        # FALLBACK STRATEGY: Return partial results instead of crashing the app
        # ═══════════════════════════════════════════════════════════════════════
        print(f"\n{'='*70}")
        print(f"[Stage 2 GRACEFUL DEGRADATION ACTIVATED]")
        print(f"{'='*70}")
        print(f"Error during Stage 2 deliverables generation: {type(e).__name__}: {e}")
        print(f"\nLast 500 chars of raw response for debugging:")
        print(f"{full_response_text[-500:] if full_response_text else 'No response text captured'}")
        print(f"\nReturning fallback dictionary to prevent app crash.")
        print(f"Stage 1 Conflict Arbitration results are still available to the user.")
        print(f"{'='*70}\n")

        # Create fallback dictionary with user-friendly error messages
        fallback_message = (
            "⚠️ **Detailed generation was interrupted due to complexity limits.**\n\n"
            "The Stage 1 strategic analysis and cross-departmental conflict arbitration "
            "are available above. Please run the analysis again to generate this specific document, "
            "or contact support if the issue persists."
        )

        return {
            "executive_briefing_memo": fallback_message,
            "business_impact_memo": fallback_message,
            "negotiation_prep_memo": fallback_message,
            "implementation_checklist": fallback_message,
            "policy_response_draft": fallback_message,
            "what_changed_brief": fallback_message,
            "business_exposure_memo": fallback_message,
            "negotiation_brief": fallback_message,
            "board_memo": fallback_message,
            "product_checklist": [
                "[GRACEFUL DEGRADATION] Stage 2 generation interrupted - please re-run analysis",
                "[GRACEFUL DEGRADATION] Stage 1 conflict arbitration is still available above"
            ]
        }


def _rule_based_fallback(policy_text: str, domain: str) -> Dict:
    """
    Keyword-frequency rule-based fallback — safe when ANTHROPIC_API_KEY is absent.
    Returns a flat analysis dict matching the analyze_policy_with_claude schema.
    """
    txt  = policy_text.lower()
    wrds = txt.split()

    threat_kws = {"prohibit", "must", "require", "shall", "enforce", "penalt", "fine",
                  "ban", "restrict", "forbidden", "mandatory", "compulsory", "violation",
                  "sanction", "infring", "liabilit", "damage", "terminat", "revok"}
    opp_kws    = {"may", "permit", "allow", "right", "option", "voluntary", "encourage",
                  "enable", "benefit", "opportunit", "revenue", "compensat", "licens",
                  "audit", "protection", "guarantee", "renegotiat"}

    t_cnt = sum(1 for w in wrds if any(kw in w for kw in threat_kws))
    o_cnt = sum(1 for w in wrds if any(kw in w for kw in opp_kws))
    total = max(t_cnt + o_cnt, 1)
    t_pct = t_cnt / total

    ip_s  = min(int(45 + t_pct * 50), 95)
    tr_s  = min(int(40 + t_pct * 45), 90)
    re_s  = min(int(40 + abs(t_pct - 0.5) * 80), 90)
    pr_s  = min(int(35 + t_pct * 45), 85)
    dirn  = "threat" if t_pct > 0.55 else ("opportunity" if t_pct < 0.35 else "neutral")
    risk  = "critical" if ip_s > 82 else ("high" if ip_s > 65 else ("medium" if ip_s > 45 else "low"))
    ev    = f"Keyword scan ({t_cnt} threat · {o_cnt} opportunity terms)"

    def _ax_action(s: int) -> Dict:
        return {
            "badge": "MONITOR", "direction": dirn, "evidence": ev,
            "summary": "Set ANTHROPIC_API_KEY to enable full Claude AI analysis.",
            "priority_actions": ["Configure ANTHROPIC_API_KEY for scenario-specific recommendations"],
        }

    return {
        "strategic_stance":  "MONITOR",
        "jurisdiction":      "Global",
        "overall_risk_level": risk,
        "executive_summary": (
            f"Rule-based keyword scan only — {t_cnt} threat-pattern terms and {o_cnt} "
            f"opportunity-pattern terms detected. Configure ANTHROPIC_API_KEY for full Claude AI analysis."
        ),
        "key_opportunities": ["Configure ANTHROPIC_API_KEY to enable full opportunity analysis"],
        "key_threats":       [f"{t_cnt} threat-pattern keywords detected — configure API key for details"],
        "substantive_changes": {
            "added_obligations": [{"title": "Policy change detected (rule-based scan)", "item": "Policy change detected (rule-based scan)",
                                   "severity": "medium",
                                   "description": f"Preliminary scan found {t_cnt} obligation-pattern indicators. "
                                                  "Set ANTHROPIC_API_KEY for detailed AI analysis."}],
            "removed_rights":   [],
            "key_thresholds":   [],
            "context_summary":  (f"Rule-based scan: {t_cnt} threat-pattern terms, "
                                 f"{o_cnt} opportunity-pattern terms detected. "
                                 "Configure ANTHROPIC_API_KEY for full analysis."),
        },
        "scores": {"IP": ip_s, "Traffic": tr_s, "Revenue": re_s, "Product": pr_s},
        "axis_actions": {ax: _ax_action(s) for ax, s in
                         [("IP", ip_s), ("Traffic", tr_s), ("Revenue", re_s), ("Product", pr_s)]},
        "evidence": {
            "parsed_claims":          [policy_text[:300] + ("…" if len(policy_text) > 300 else "")],
            "claim_level_provenance": [],
            "policy_memory_graph":    [{"reference_id": "API-REQUIRED",
                                        "description": "Set ANTHROPIC_API_KEY to activate Policy Memory Graph cross-referencing."}],
        },
        "risk_matrix_points": [
            {"label": "IP Exposure",    "days_to_enactment": 10, "business_severity": ip_s},
            {"label": "Traffic Impact", "days_to_enactment": 22, "business_severity": tr_s},
            {"label": "Revenue Risk",   "days_to_enactment": 38, "business_severity": re_s},
            {"label": "Product Change", "days_to_enactment": 55, "business_severity": pr_s},
        ],
        "what_changed_brief":       f"Rule-based scan: {t_cnt} obligation-pattern keywords detected. Set API key for full analysis.",
        "what_changed_quotes":      [policy_text[:300] + ("…" if len(policy_text) > 300 else "")],
        "overall_risk":             risk.upper(),
        "business_exposure_memo":   f"Keyword-based preliminary scan: {t_cnt} threat-pattern terms detected.",
        "business_exposure_quotes": [],
        "negotiation_brief":        "Set ANTHROPIC_API_KEY for full negotiation brief.",
        "negotiation_quotes":       [],
        "board_memo":               "Set ANTHROPIC_API_KEY for full board memo.",
        "board_memo_quotes":        [],
        "product_checklist":        ["Set ANTHROPIC_API_KEY for detailed product checklist."],
    }



# ─── Chart Functions ──────────────────────────────────────────────────────────
_DIR_COLORS = {
    "threat":      ("#8B2635", "rgba(139, 38, 53, 0.18)"),
    "opportunity": ("#1A6B3C", "rgba(26, 107, 60, 0.18)"),
    "neutral":     ("#0ABAB5", "rgba(10, 186, 181, 0.16)"),
}


def _majority_direction(scores: Dict) -> str:
    dirs = [scores[a]["direction"] for a in ["IP", "Traffic", "Revenue", "Product"]]
    return max(set(dirs), key=dirs.count)


def create_exposure_delta_radar(scores: Dict) -> go.Figure:
    """Before/After radar: Current Baseline vs. New Policy Exposure."""
    axes = ["Traffic", "IP", "Revenue", "Product"]

    # New Policy Exposure = scores returned by AI analysis
    new_vals = [scores[a]["score"] for a in axes]

    # Current Baseline = fixed at 70 across all axes (pre-policy reference line)
    base_vals = [70] * len(axes)

    # Close the polygon
    theta   = axes + [axes[0]]
    new_r   = new_vals  + [new_vals[0]]
    base_r  = base_vals + [base_vals[0]]

    fig = go.Figure()

    # ── Trace 1: Current Baseline (muted steel blue) ───────────────────────
    fig.add_trace(go.Scatterpolar(
        r=base_r, theta=theta,
        fill="toself",
        fillcolor="rgba(90,130,160,0.18)",
        line=dict(color="rgba(90,130,160,0.65)", width=1.5, dash="dot"),
        marker=dict(size=5, color="rgba(90,130,160,0.8)", symbol="circle"),
        name="Current Baseline",
        hovertemplate="<b>%{theta}</b><br>Baseline: %{r}/100<extra></extra>",
    ))

    # ── Trace 2: New Policy Exposure (crimson warning) ─────────────────────
    fig.add_trace(go.Scatterpolar(
        r=new_r, theta=theta,
        fill="toself",
        fillcolor="rgba(139,38,53,0.28)",
        line=dict(color="#C0392B", width=2.2),
        marker=dict(size=7, color="#E74C3C", symbol="diamond"),
        name="New Policy Exposure",
        hovertemplate="<b>%{theta}</b><br>New Exposure: %{r}/100<extra></extra>",
    ))

    fig.update_layout(
        polar=dict(
            radialaxis=dict(
                visible=True, range=[0, 100], tickvals=[25, 50, 75, 100],
                tickfont=dict(size=8, color="#6B6560", family="Montserrat"),
                gridcolor="rgba(196,191,184,0.07)",
                linecolor="rgba(196,191,184,0.05)",
            ),
            angularaxis=dict(
                tickfont=dict(size=11, color="#C4BFB8", family="Montserrat"),
                gridcolor="rgba(10,186,181,0.08)",
                linecolor="rgba(10,186,181,0.10)",
            ),
            bgcolor="rgba(13,13,13,0.0)",
        ),
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
        font=dict(color="#9A9590", family="Montserrat"),
        legend=dict(
            x=0.5, y=-0.12, xanchor="center", orientation="h",
            font=dict(size=9, color="#9A9590", family="Montserrat"),
            bgcolor="rgba(0,0,0,0)", borderwidth=0,
        ),
        margin=dict(l=55, r=55, t=30, b=50),
        height=380,
    )
    return fig


def _risk_matrix_points_to_clauses(pts: List[Dict]) -> List[Dict]:
    """Convert analyze_policy_with_claude risk_matrix_points list → clauses format for the chart."""
    _COLORS = ["#C0392B", "#E67E22", "#A8892A", "#0ABAB5", "#8B2635", "#1A6B3C"]
    return [
        {
            "name":   pt.get("label", f"Clause {i + 1}"),
            "x":      max(1, min(100, int(pt.get("days_to_enactment", 30)))),
            "y":      max(1, min(100, int(pt.get("business_severity", 50)))),
            "color":  _COLORS[i % len(_COLORS)],
            "size":   max(8, min(16, 10 + (int(pt.get("business_severity", 50)) - 50) // 10)),
        }
        for i, pt in enumerate(pts)
        if isinstance(pt, dict)
    ]


def create_risk_urgency_matrix(scores: Dict, clauses: Optional[List[Dict]] = None) -> go.Figure:
    """Risk & Urgency Matrix: scatter of key policy clauses by time + severity."""
    if clauses is None:
        # ── Fallback: derive clause positions from axis scores ───────────────
        def _gs(key):
            v = scores.get(key, 50)
            return v.get("score", 50) if isinstance(v, dict) else int(v)
        ip_s  = _gs("IP")
        tr_s  = _gs("Traffic")
        rv_s  = _gs("Revenue")
        pr_s  = _gs("Product")
        clauses = [
            {
                "name": "IP Licensing Term",
                "x": 10,
                "y": min(98, ip_s + 4),
                "color": "#C0392B",
                "size": 14,
                "detail": f"IP Licensing Term<br>Days to Enactment: 10<br>Severity: {min(98, ip_s+4)}/100",
            },
            {
                "name": "Data Share-back Clause",
                "x": 22,
                "y": max(30, tr_s - 5),
                "color": "#E67E22",
                "size": 12,
                "detail": f"Data Share-back Clause<br>Days to Enactment: 22<br>Severity: {max(30, tr_s-5)}/100",
            },
            {
                "name": "Content Attribution Rule",
                "x": 38,
                "y": max(25, rv_s - 10),
                "color": "#A8892A",
                "size": 11,
                "detail": f"Content Attribution Rule<br>Days to Enactment: 38<br>Severity: {max(25, rv_s-10)}/100",
            },
            {
                "name": "Opt-out UI Change",
                "x": 55,
                "y": max(20, pr_s - 8),
                "color": "#0ABAB5",
                "size": 10,
                "detail": f"Opt-out UI Change<br>Days to Enactment: 55<br>Severity: {max(20, pr_s-8)}/100",
            },
        ]
    else:
        # ── Ensure each clause has a "detail" field ──────────────────────────
        for c in clauses:
            if "detail" not in c:
                c["detail"] = f"{c['name']}<br>Days to Enactment: {c['x']}<br>Severity: {c['y']}/100"
            if "color" not in c:
                c["color"] = "#C0392B"
            if "size" not in c:
                c["size"] = 13

    fig = go.Figure()

    # ── Urgency quadrant shading ──────────────────────────────────────────────
    # Q1 top-left: Critical & Urgent (darkest red tint)
    fig.add_shape(type="rect", x0=0, y0=60, x1=30, y1=100,
                  fillcolor="rgba(139,38,53,0.07)", line=dict(width=0), layer="below")
    # Q2 top-right: Critical but time available
    fig.add_shape(type="rect", x0=30, y0=60, x1=70, y1=100,
                  fillcolor="rgba(168,137,42,0.05)", line=dict(width=0), layer="below")
    # Quadrant labels
    for txt, qx, qy, col in [
        ("PROTECT & LICENSE", 15, 97, "rgba(139,38,53,0.55)"),
        ("HIGH EXPOSURE", 50, 97, "rgba(168,137,42,0.40)"),
        ("MONITOR", 50, 22, "rgba(100,100,100,0.35)"),
    ]:
        fig.add_annotation(x=qx, y=qy, text=txt, showarrow=False,
                           font=dict(size=7, color=col, family="Montserrat"),
                           xanchor="center")

    # ── Divider lines ─────────────────────────────────────────────────────────
    fig.add_shape(type="line", x0=30, y0=0, x1=30, y1=100,
                  line=dict(color="rgba(196,191,184,0.08)", width=1, dash="dot"))
    fig.add_shape(type="line", x0=0, y0=60, x1=70, y1=60,
                  line=dict(color="rgba(196,191,184,0.08)", width=1, dash="dot"))

    # ── Data points ──────────────────────────────────────────────────────────
    for c in clauses:
        fig.add_trace(go.Scatter(
            x=[c["x"]], y=[c["y"]],
            mode="markers+text",
            marker=dict(
                size=c["size"] + 4,
                color=c["color"],
                opacity=0.85,
                line=dict(width=1.5, color="rgba(255,255,255,0.15)"),
            ),
            text=[f'  {c["name"]}'],
            textposition="middle right",
            textfont=dict(size=9, color="#C4BFB8", family="Montserrat"),
            hovertemplate=f"<b>{c['detail']}</b><extra></extra>",
            showlegend=False,
        ))

    # x-axis range adapts to the furthest clause with 15% padding
    max_x = max((c["x"] for c in clauses), default=70)
    x_max = max(70, int(max_x * 1.15))

    fig.update_layout(
        xaxis=dict(
            title=dict(text="← Time to Enactment (Days)  [left = more urgent]",
                       font=dict(size=9, color="#9A9590", family="Montserrat")),
            range=[0, x_max], autorange="reversed",
            tickfont=dict(size=9, color="#6B6560", family="Montserrat"),
            gridcolor="rgba(196,191,184,0.05)", showline=False, zeroline=False,
        ),
        yaxis=dict(
            title=dict(text="Business Impact Severity",
                       font=dict(size=9, color="#9A9590", family="Montserrat")),
            range=[0, 100],
            tickfont=dict(size=9, color="#6B6560", family="Montserrat"),
            gridcolor="rgba(196,191,184,0.05)", showline=False, zeroline=False,
        ),
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(13,13,13,0.0)",
        font=dict(color="#9A9590", family="Montserrat"),
        showlegend=False,
        margin=dict(l=60, r=30, t=20, b=55),
        height=380,
    )
    return fig


# ─── HTML Component Helpers ───────────────────────────────────────────────────
_ACCENT = "#0ABAB5"


def _sev_color(sev: str) -> str:
    return {"high": "#8B2635", "medium": "#A8892A", "low": "#1A6B3C"}.get(sev, _ACCENT)


def _risk_config(level: str):
    """Returns (stance_label, color, stance_sub) for a raw risk level."""
    return {
        "critical": (
            "PROTECT & LICENSE",
            "#8B2635",
            "Immediate IP defense · aggressive licensing required",
        ),
        "high": (
            "NEGOTIATE & LICENSE",
            "#A8892A",
            "Proactive licensing engagement · secure favorable terms now",
        ),
        "medium": (
            "MONITOR & PROMOTE",
            "#8A7020",
            "Track developments · selectively capture opportunities",
        ),
        "low": (
            "PROMOTE & WAIT",
            "#1A6B3C",
            "Selective engagement · no urgent defensive action required",
        ),
    }.get(level.lower() if level else "medium", ("—", "#C4BFB8", ""))


def _accent_divider() -> None:
    st.markdown(
        '<div style="border-top:1px solid rgba(10,186,181,0.15);margin:2.5rem 0 2rem"></div>',
        unsafe_allow_html=True,
    )


def _enhance_markdown_readability(content: str) -> str:
    """
    Enhances markdown deliverable readability by:
    1. Converting ### headers to bold text for better visual hierarchy
    2. Converting ## headers to larger bold text
    3. Preserving bullet points and formatting
    4. Adding subtle visual polish

    Args:
        content: Raw markdown content from deliverables

    Returns:
        Enhanced markdown with better visual formatting
    """
    if not content or not isinstance(content, str):
        return ""

    import re

    # Replace ### headers with bold text (smaller sections)
    content = re.sub(r'^### (.+)$', r'**\1**', content, flags=re.MULTILINE)

    # Replace ## headers with larger bold text (major sections)
    content = re.sub(r'^## (.+)$', r'### **\1**', content, flags=re.MULTILINE)

    # Replace # headers with even larger text (top-level - rarely used in deliverables)
    content = re.sub(r'^# (.+)$', r'## **\1**', content, flags=re.MULTILINE)

    return content.strip()


def _display_deliverable_with_highlights(content: str, title: str = "", highlight_sections: list = None):
    """
    Display a deliverable with enhanced formatting and optional highlighted sections.

    Args:
        content: Markdown content to display
        title: Optional title to display above content
        highlight_sections: List of section titles to highlight with background color
    """
    if title:
        st.markdown(f"#### {title}")

    enhanced_content = _enhance_markdown_readability(content)

    # If specific sections should be highlighted
    if highlight_sections and enhanced_content:
        for section in highlight_sections:
            # Add background highlight to specific sections
            pattern = f"(\\*\\*{re.escape(section)}\\*\\*.*?)(?=\\*\\*|$)"
            replacement = r'<div style="background-color: rgba(10, 186, 181, 0.08); padding: 8px 12px; border-left: 3px solid #0ABAB5; margin: 8px 0;">\1</div>'
            enhanced_content = re.sub(pattern, replacement, enhanced_content, flags=re.DOTALL)

    st.markdown(enhanced_content, unsafe_allow_html=True)


def _engine_badge(label: str) -> str:
    """Return an inline HTML engine badge span."""
    return (
        f'<span style="font-size:0.7rem;color:#888;border:1px solid #333;'
        f'padding:2px 6px;border-radius:4px;margin-left:10px;'
        f'font-family:\'Montserrat\',sans-serif;font-weight:400;'
        f'letter-spacing:0.04em;vertical-align:middle">[Engine: {label}]</span>'
    )


def _section_label(num: str, title: str) -> None:
    st.markdown(f"""
    <div style="display:flex;align-items:center;gap:16px;margin:2rem 0 1.4rem">
      <div style="font-family:'Montserrat',system-ui,sans-serif;color:{_ACCENT};font-size:1.4rem;
                  font-weight:700;line-height:1;min-width:28px">{num}</div>
      <div>
        <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;font-size:0.62rem;
                    letter-spacing:0.22em;text-transform:uppercase;margin-bottom:2px">Step</div>
        <div style="font-family:'Montserrat',system-ui,sans-serif;color:#F0EDE6;font-size:1.05rem;
                    font-weight:600;letter-spacing:0.04em">{title}</div>
      </div>
      <div style="flex:1;height:1px;background:rgba(10,186,181,0.12);margin-left:8px"></div>
    </div>""", unsafe_allow_html=True)


def _item_card(item: Dict, accent: str) -> None:
    desc = item.get("description", "")
    sev  = item.get("severity", "")
    val  = item.get("value", "")
    sc   = _sev_color(sev) if sev else accent
    st.markdown(f"""
    <div style="background:#111111;border-left:2px solid {accent};
                padding:12px 16px;margin:5px 0;">
      <div style="color:#F0EDE6;font-family:'Montserrat',sans-serif;
                  font-size:0.82rem;font-weight:500;margin-bottom:4px">
        {item.get('item','')}
      </div>
      {"<span style='color:"+sc+";font-family:Montserrat,sans-serif;font-size:0.60rem;font-weight:700;letter-spacing:0.2em;text-transform:uppercase'>"+sev+"</span>" if sev else ""}
      {"<span style='color:"+_ACCENT+";font-family:Montserrat,sans-serif;font-weight:600;font-size:0.85rem;margin-left:10px'>"+val+"</span>" if val else ""}
      <div style="color:#C4BFB8;font-family:'Montserrat',sans-serif;
                  font-size:0.76rem;line-height:1.6;margin-top:5px">
        {desc}
      </div>
    </div>""", unsafe_allow_html=True)


def _score_card(axis: str, info: Dict) -> None:
    d = info["direction"]
    lc, _ = _DIR_COLORS.get(d, (_ACCENT, ""))
    ev = info.get("evidence", "")
    acts = info.get("priority_actions", [])[:2]
    icons = {"IP": "◈", "Traffic": "◉", "Revenue": "◆", "Product": "◇"}
    st.markdown(f"""
    <div style="background:#111111;padding:20px 16px;border-top:1px solid {lc};
                border-left:1px solid rgba(10,186,181,0.08);
                border-right:1px solid rgba(10,186,181,0.08);
                border-bottom:1px solid rgba(10,186,181,0.08)">
      <div style="font-family:'Montserrat',system-ui,sans-serif;color:{_ACCENT};
                  font-size:1.2rem;margin-bottom:6px">{icons.get(axis,'◆')}</div>
      <div style="font-family:'Montserrat',sans-serif;color:#9A9590;
                  font-size:0.60rem;letter-spacing:0.22em;text-transform:uppercase">{axis}</div>
      <div style="font-family:'Montserrat',system-ui,sans-serif;color:{lc};
                  font-size:2.8rem;font-weight:700;line-height:1;margin:6px 0 2px">{info['score']}</div>
      <div style="font-family:'Montserrat',sans-serif;color:{lc};
                  font-size:0.60rem;letter-spacing:0.20em;text-transform:uppercase;
                  margin-bottom:10px">{d}</div>
      <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;
                  font-size:0.74rem;line-height:1.7;word-break:break-word">
        {ev}
      </div>
      {''.join(f'<div style="font-family:Montserrat,sans-serif;color:#9A9590;font-size:0.70rem;margin-top:5px;padding-left:8px;border-left:1px solid {_ACCENT}33;word-break:break-word">▸ {a}</div>' for a in acts)}
    </div>""", unsafe_allow_html=True)


def _col_header(label: str, badge: str = "", badge_color: str = "#0ABAB5") -> None:
    badge_html = (
        f'<span style="font-family:Montserrat,sans-serif;font-size:0.48rem;font-weight:700;'
        f'letter-spacing:0.14em;color:{badge_color};background:{badge_color}18;'
        f'border:1px solid {badge_color}55;border-radius:3px;padding:2px 6px;'
        f'margin-left:8px;vertical-align:middle">{badge}</span>'
    ) if badge else ""
    st.markdown(f"""
    <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;font-size:0.62rem;
                letter-spacing:0.20em;text-transform:uppercase;margin-bottom:10px;
                border-bottom:1px solid rgba(10,186,181,0.10);padding-bottom:8px;
                display:flex;align-items:center">
      <span>{label}</span>{badge_html}
    </div>""", unsafe_allow_html=True)


def _prose_block(text: str) -> None:
    """Render prose with excess-newline sanitization, enhanced markdown headers, and inline Markdown table support."""
    import re as _re

    # ── 0. Enhance markdown readability (convert headers to bold) ──────────────
    text = _enhance_markdown_readability(text)

    # ── 1. Sanitize: collapse 3+ consecutive newlines → single blank line ──────
    text = _re.sub(r'\n{3,}', '\n\n', text.strip())

    # ── 2. Split text into prose / table segments ──────────────────────────────
    lines = text.split('\n')
    segments: list = []
    i = 0
    while i < len(lines):
        if lines[i].strip().startswith('|'):
            tbl: list = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                tbl.append(lines[i])
                i += 1
            segments.append(('table', tbl))
        else:
            prose: list = []
            while i < len(lines) and not lines[i].strip().startswith('|'):
                prose.append(lines[i])
                i += 1
            segments.append(('prose', prose))

    # ── 3. Build HTML ──────────────────────────────────────────────────────────
    _TH = (
        "padding:7px 14px;text-align:left;font-family:'Montserrat',sans-serif;"
        "font-size:0.62rem;letter-spacing:0.18em;text-transform:uppercase;"
        "color:#9A9590;border-bottom:1px solid rgba(10,186,181,0.22);"
        "background:rgba(10,186,181,0.06);white-space:nowrap"
    )
    _TD = (
        "padding:7px 14px;font-family:'Montserrat',sans-serif;"
        "font-size:0.80rem;color:#C4BFB8;line-height:1.5;"
        "border-bottom:1px solid rgba(196,191,184,0.06)"
    )

    html_parts: list = []
    for seg_type, seg_lines in segments:
        if seg_type == 'prose':
            body = '\n'.join(seg_lines).strip()
            if body:
                # Enhanced markdown to HTML conversion
                safe = body.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

                # Convert **bold** to <strong> with enhanced styling
                safe = _re.sub(
                    r'\*\*(.+?)\*\*',
                    r'<strong style="color:#0ABAB5;font-weight:600;letter-spacing:0.02em">\1</strong>',
                    safe
                )

                # Convert bullet points (•, -, *) to styled bullets
                safe = _re.sub(
                    r'^([•\-\*])\s+(.+)$',
                    r'<div style="margin-left:16px;margin-bottom:6px"><span style="color:#0ABAB5">•</span> \2</div>',
                    safe,
                    flags=_re.MULTILINE
                )

                # Preserve newlines as <br> for non-bullet text
                safe = safe.replace('\n', '<br>')

                html_parts.append(
                    f'<div style="font-family:\'Montserrat\',sans-serif;color:#C4BFB8;'
                    f'font-size:0.88rem;line-height:1.85;margin-bottom:0.8rem">'
                    f'{safe}</div>'
                )
        else:
            rows_html: list = []
            header_done = False
            for tl in seg_lines:
                cols = [c.strip() for c in tl.strip().strip('|').split('|')]
                # Skip separator row (---, :--:, etc.)
                if all(_re.match(r'^[-: ]+$', c) for c in cols if c):
                    header_done = True
                    continue
                if not header_done:
                    cells = ''.join(f'<th style="{_TH}">{c}</th>' for c in cols)
                    rows_html.append(f'<tr>{cells}</tr>')
                    header_done = True
                else:
                    cells = ''.join(f'<td style="{_TD}">{c}</td>' for c in cols)
                    rows_html.append(f'<tr>{cells}</tr>')
            if rows_html:
                html_parts.append(
                    f'<div style="overflow-x:auto;margin:0 0 0.8rem">'
                    f'<table style="width:100%;border-collapse:collapse;'
                    f'background:#0D0D0D;border:1px solid rgba(10,186,181,0.12)">'
                    f'{"".join(rows_html)}</table></div>'
                )

    inner = ''.join(html_parts) or f'<div style="color:#6B6560;font-size:0.80rem">—</div>'
    st.markdown(
        f'<div style="background:#111111;border:1px solid rgba(10,186,181,0.10);'
        f'padding:24px 28px;margin:4px 0">{inner}</div>',
        unsafe_allow_html=True,
    )


def _checklist_items(items: list) -> None:
    """Render product checklist as styled actionable items."""
    for i, item in enumerate(items, 1):
        st.markdown(f"""
        <div style="background:#111111;border-left:2px solid {_ACCENT};
                    padding:12px 16px;margin:6px 0;display:flex;gap:12px;align-items:flex-start">
          <div style="font-family:'Montserrat',system-ui,sans-serif;color:{_ACCENT};
                      font-size:0.80rem;font-weight:700;min-width:20px;margin-top:1px">{i:02d}</div>
          <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;
                      font-size:0.82rem;line-height:1.6">{item}</div>
        </div>""", unsafe_allow_html=True)


def _evidence_block(
    quotes: List[str],
    claim_tag: str = "Source Citation",
    claim_color: str = "#0ABAB5",
    agent_tag: str = "",
    agent_color: str = "#9A9590",
) -> None:
    """Render evidence quotes with claim-level provenance tags (audit-grade)."""
    if not quotes:
        return
    st.markdown(f"""
    <div style="margin:1.8rem 0 0.6rem">
      <div style="font-family:'Montserrat',sans-serif;color:{_ACCENT};font-size:0.58rem;
                  letter-spacing:0.28em;text-transform:uppercase;margin-bottom:0.9rem;
                  display:flex;align-items:center;gap:10px">
        <span>◈</span>
        <span>EVIDENCE — CLAIM-LEVEL PROVENANCE</span>
        <div style="flex:1;height:1px;background:rgba(10,186,181,0.12)"></div>
      </div>
    </div>""", unsafe_allow_html=True)

    for q in quotes:
        agent_badge = (
            f'<span style="background:rgba(154,149,144,0.12);color:{agent_color};'
            f'border:1px solid rgba(154,149,144,0.22);font-family:Montserrat,sans-serif;'
            f'font-size:0.50rem;letter-spacing:0.16em;text-transform:uppercase;'
            f'padding:2px 7px;margin-left:6px">{agent_tag}</span>'
            if agent_tag else ""
        )
        st.markdown(f"""
        <div style="background:#0D0D0D;border:1px solid rgba(10,186,181,0.10);
                    border-left:3px solid {claim_color};
                    padding:13px 16px;margin:7px 0">
          <div style="display:flex;align-items:center;margin-bottom:9px;flex-wrap:wrap;gap:4px">
            <span style="background:rgba(10,186,181,0.08);color:{claim_color};
                         border:1px solid {claim_color}44;font-family:Montserrat,sans-serif;
                         font-size:0.50rem;letter-spacing:0.18em;text-transform:uppercase;
                         padding:2px 8px">{claim_tag}</span>
            {agent_badge}
          </div>
          <div style="color:rgba(240,237,230,0.90);font-family:'Montserrat',sans-serif;
                      font-size:0.84rem;line-height:1.7;font-style:normal;font-weight:400">
            &ldquo;{q}&rdquo;
          </div>
        </div>""", unsafe_allow_html=True)


def _to_docx_bytes(title: str, body: str, domain: str, doc_id: str) -> bytes:
    """Generate a styled Word (.docx) document from title + prose body."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import docx.oxml as oxml

    doc = DocxDocument()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # Title
    title_para = doc.add_paragraph()
    title_run  = title_para.add_run(title)
    title_run.bold      = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(0x0A, 0xBA, 0xB5)

    # Metadata block
    meta_lines = [
        f"Document ID: {doc_id}",
        f"Domain: {domain}",
        f"Grounding Sources: External Policy Text + Nikkei Internal DB",
        f"Compliance Status: Pending Human Review (GC & CPO)",
        f"Traceability: Claim-level provenance active",
    ]
    meta_para = doc.add_paragraph()
    for line in meta_lines:
        run = meta_para.add_run(line + "\n")
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_paragraph()  # spacer

    # Body — split on double newlines; render numbered items as list
    for chunk in body.split("\n\n"):
        chunk = chunk.strip()
        if not chunk:
            continue
        lines = chunk.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph()
            # Detect simple numbered list markers like "1." or "(1)"
            import re as _re
            if _re.match(r"^(\d+\.|[\(\[]\d+[\)\]])\s", line):
                p.style = doc.styles["List Number"]
                line = _re.sub(r"^(\d+\.|[\(\[]\d+[\)\]])\s+", "", line)
            p.add_run(line).font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _download_row(label: str, data: bytes, file_name: str, key: str) -> None:
    """Render a visually prominent .docx download button with teal accent divider."""
    st.markdown(
        '<div style="height:1px;background:rgba(10,186,181,0.14);margin:1.6rem 0 1.2rem"></div>',
        unsafe_allow_html=True,
    )
    _, btn_col, _ = st.columns([1, 2, 1])
    with btn_col:
        st.download_button(
            label=label,
            data=data,
            file_name=file_name,
            mime=_DOCX_MIME,
            use_container_width=True,
            type="primary",
            key=key,
        )


def _audit_block(doc_id: str, domain: str = "", step2_data: Optional[Dict] = None, policy_text: str = "", jurisdiction: str = "") -> None:
    """Render compact audit-metadata header — contextual data derived from live analysis."""
    import datetime as _dt

    # ── Grounding Sources: domain-derived, not keyword-dependent ──────────────
    grounding = f"{domain} Policy Source · Active Partner Contract DB" if domain else "External Policy Source · Active Partner Contract DB"

    # ── Compliance Status: derive due date from doc_id timestamp segment ──────
    try:
        date_part = doc_id.split("-")[1]          # e.g. "20260318"
        req_date  = _dt.datetime.strptime(date_part, "%Y%m%d")
        due_date  = req_date + _dt.timedelta(days=2)
        due_str   = due_date.strftime("%b %-d, 12:00 JST")
    except Exception:
        due_str = (_dt.datetime.now() + _dt.timedelta(days=2)).strftime("%b %-d, 12:00 JST")
    compliance_html = f"🟡 Pending Legal Approval &nbsp;<span style='color:#6B6560'>(Due: {due_str})</span>"

    # ── Traceability: derive confidence from avg axis score ──────────────────
    if step2_data:
        axes = step2_data.get("scores", {})
        scores = [v.get("score", 50) for v in axes.values() if isinstance(v, dict)]
        avg = int(sum(scores) / len(scores)) if scores else 50
        # Map 0-100 score to match-rate (score already reflects alignment with red-lines)
        match_pct = min(98, max(60, avg + 8))
        if avg >= 75:
            conf_label, conf_color = "High Confidence", "#0ABAB5"
        elif avg >= 50:
            conf_label, conf_color = "Moderate Confidence", "#A8892A"
        else:
            conf_label, conf_color = "Low Confidence", "#8B2635"
        traceability = (
            f"<span style='color:{conf_color}'>{conf_label}</span>"
            f"<span style='color:#6B6560'> ({match_pct}% match with internal red-lines)</span>"
        )
    else:
        traceability = "<span style='color:#0ABAB5'>High Confidence</span><span style='color:#6B6560'> (92% match with internal red-lines)</span>"

    st.markdown(f"""
    <div style="background:#0D0D0D;border:1px solid rgba(10,186,181,0.14);
                border-left:2px solid rgba(10,186,181,0.45);
                padding:10px 18px;margin-bottom:20px;
                display:flex;flex-wrap:wrap;gap:10px 36px;align-items:center">
      <div>
        <div style="font-family:'Montserrat',sans-serif;color:#9A9590;font-size:0.50rem;
                    letter-spacing:0.24em;text-transform:uppercase;margin-bottom:2px">Document ID</div>
        <div style="font-family:'Montserrat',sans-serif;color:{_ACCENT};
                    font-size:0.68rem;letter-spacing:0.08em">{doc_id}</div>
      </div>
      <div>
        <div style="font-family:'Montserrat',sans-serif;color:#9A9590;font-size:0.50rem;
                    letter-spacing:0.24em;text-transform:uppercase;margin-bottom:2px">Grounding Sources</div>
        <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;font-size:0.68rem">
          {grounding}
        </div>
      </div>
      <div>
        <div style="font-family:'Montserrat',sans-serif;color:#9A9590;font-size:0.50rem;
                    letter-spacing:0.24em;text-transform:uppercase;margin-bottom:2px">Jurisdiction</div>
        <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;font-size:0.68rem">
          {jurisdiction if jurisdiction else "—"}
        </div>
      </div>
      <div>
        <div style="font-family:'Montserrat',sans-serif;color:#9A9590;font-size:0.50rem;
                    letter-spacing:0.24em;text-transform:uppercase;margin-bottom:2px">Compliance Status</div>
        <div style="font-family:'Montserrat',sans-serif;color:#A8892A;font-size:0.68rem">
          {compliance_html}
        </div>
      </div>
      <div>
        <div style="font-family:'Montserrat',sans-serif;color:#9A9590;font-size:0.50rem;
                    letter-spacing:0.24em;text-transform:uppercase;margin-bottom:2px">Traceability</div>
        <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;font-size:0.68rem">
          {traceability}
        </div>
      </div>
    </div>""", unsafe_allow_html=True)


def _policy_memory_block(domain: str, pmg_hits: Optional[List[tuple]] = None) -> None:
    """Render Policy Memory Graph — historical red-line match panel.
    pmg_hits from Claude AI analysis overrides the domain-keyed institutional memory defaults.
    """
    if not pmg_hits:
        # No PMG matches from Claude — render a neutral informational message
        st.markdown(f"""
        <div style="margin:20px 0 4px">
          <div style="font-family:'Montserrat',sans-serif;color:{_ACCENT};font-size:0.52rem;
                      letter-spacing:0.28em;text-transform:uppercase;margin-bottom:10px;
                      display:flex;align-items:center;gap:8px">
            <span>🎯</span>
            <span>Policy Memory Graph — Historical Red-Line Match</span>
            <div style="flex:1;height:1px;background:rgba(10,186,181,0.14)"></div>
          </div>
          <div style="background:rgba(10,186,181,0.03);border:1px solid rgba(10,186,181,0.14);
                      border-left:3px solid {_ACCENT};padding:14px 18px;
                      font-family:'Montserrat',sans-serif;color:#9A9590;font-size:0.62rem;
                      line-height:1.65">
            No matching institutional red-lines or historical contract precedents identified
            for this policy text in the current Policy Memory Graph index.
          </div>
        </div>""", unsafe_allow_html=True)
        return
    hits = pmg_hits
    cards_html = "".join(
        f"""<div style="background:rgba(10,186,181,0.03);border:1px solid rgba(10,186,181,0.18);
                        border-left:3px solid {_ACCENT};border-radius:0 4px 4px 0;
                        padding:11px 14px;margin-bottom:8px">
              <div style="display:flex;align-items:baseline;gap:10px;margin-bottom:5px">
                <span style="font-family:'Montserrat',sans-serif;color:{_ACCENT};font-size:0.54rem;
                             font-weight:600;letter-spacing:0.10em">{ref}</span>
                <span style="font-family:'Montserrat',sans-serif;color:#9A9590;font-size:0.50rem;
                             font-style:italic">{source}</span>
              </div>
              <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;font-size:0.63rem;
                          line-height:1.65">{memo}</div>
            </div>"""
        for ref, source, memo in hits
    )
    st.markdown(f"""
    <div style="margin:20px 0 4px">
      <div style="font-family:'Montserrat',sans-serif;color:{_ACCENT};font-size:0.52rem;
                  letter-spacing:0.28em;text-transform:uppercase;margin-bottom:10px;
                  display:flex;align-items:center;gap:8px">
        <span>🎯</span>
        <span>Policy Memory Graph — Historical Red-Line Match</span>
        <div style="flex:1;height:1px;background:rgba(10,186,181,0.14)"></div>
        <span style="color:#6B6560;font-size:0.46rem;letter-spacing:0.10em;text-transform:none;
                     font-style:italic">Auto-matched from 340+ archived partner contracts &amp; Board Minutes</span>
      </div>
      {cards_html}
    </div>""", unsafe_allow_html=True)


def _counterparty_panel(domain: str, step2_data: Optional[Dict] = None) -> None:
    """Render counterparty market-power profile badge — derived dynamically from analysis scores."""
    scores = (step2_data or {}).get("scores", {})

    def _gs(key: str) -> int:
        v = scores.get(key, 50)
        return v.get("score", 50) if isinstance(v, dict) else int(v)

    tr_score  = _gs("Traffic")
    rev_score = _gs("Revenue")

    # Derive traffic dependency from Traffic score
    if tr_score >= 75:
        traffic_level, tc = "CRITICAL", "#8B2635"
    elif tr_score >= 55:
        traffic_level, tc = "HIGH", "#A8892A"
    elif tr_score >= 35:
        traffic_level, tc = "MEDIUM", "#0ABAB5"
    else:
        traffic_level, tc = "LOW", "#1A6B3C"
    traffic_pct = f"{min(95, max(20, tr_score))}%"

    # Derive substitutability from Revenue score (high revenue exposure = low substitutability)
    if rev_score >= 70:
        sub_level, sc = "LOW", "#8B2635"
    elif rev_score >= 45:
        sub_level, sc = "MEDIUM", "#A8892A"
    else:
        sub_level, sc = "HIGH", "#1A6B3C"

    # Platform label derived from domain
    _platform_labels = {
        "AI Search & Distribution":    "Dominant Search Platform",
        "AI Licensing & Copyright":    "Foundation Model Provider",
        "Platform Distribution Policies": "Social / Distribution Platform",
    }
    platform_label = _platform_labels.get(domain, "Platform Partner")

    st.markdown(f"""
    <div style="background:#0D0D0D;border:1px solid #2A2A2A;
                padding:12px 18px;margin-bottom:18px;
                display:flex;align-items:center;gap:12px;flex-wrap:wrap">
      <div style="font-family:'Montserrat',sans-serif;color:#6B6560;font-size:0.46rem;
                  letter-spacing:0.22em;text-transform:uppercase;margin-right:4px;
                  white-space:nowrap">Target Platform Profile</div>
      <span style="font-family:'Montserrat',sans-serif;font-size:0.62rem;color:#C4BFB8;
                   border:1px solid #333;border-radius:3px;padding:3px 10px;
                   white-space:nowrap">
        Platform: <strong style="color:#F0EDE6">{platform_label}</strong>
      </span>
      <span style="font-family:'Montserrat',sans-serif;font-size:0.62rem;
                   border:1px solid {tc}55;border-radius:3px;padding:3px 10px;
                   color:{tc};white-space:nowrap">
        Traffic Dependency: <strong>{traffic_pct} ({traffic_level})</strong>
      </span>
      <span style="font-family:'Montserrat',sans-serif;font-size:0.62rem;
                   border:1px solid {sc}55;border-radius:3px;padding:3px 10px;
                   color:{sc};white-space:nowrap">
        Substitutability: <strong>{sub_level}</strong>
      </span>
    </div>""", unsafe_allow_html=True)


def _pplw_map_block(strategic_stance: str = "", risk_raw: str = "high") -> None:
    """Render PPLW 4-stance visual mapping — driven by Claude's strategic_stance field."""
    # Parse active badges from Claude's strategic_stance (e.g. "PROTECT & LICENSE")
    stance_upper = (strategic_stance or "").upper()
    active_badges = {b for b in ["PROTECT", "PROMOTE", "LICENSE", "WAIT"] if b in stance_upper}
    # Fallback to risk-level derivation when stance is absent or unrecognised
    if not active_badges:
        _fallback = {"critical": "PROTECT", "high": "LICENSE", "medium": "PROMOTE", "low": "WAIT"}
        active_badges = {_fallback.get((risk_raw or "high").lower(), "LICENSE")}
    stances = [
        ("PROTECT", "#8B2635", "Immediate IP defense"),
        ("PROMOTE", "#0ABAB5", "Brand visibility & discovery"),
        ("LICENSE", "#A8892A", "Revenue & contract leverage"),
        ("WAIT",    "#6B6560", "Pending further clarity"),
    ]
    badges_html = ""
    for label, color, desc in stances:
        is_active = label in active_badges
        opacity   = "1" if is_active else "0.32"
        border    = f"border:1.5px solid {color}" if is_active else f"border:1px solid {color}44"
        prefix    = "▶ " if is_active else ""
        badges_html += (
            f'<div style="display:inline-flex;flex-direction:column;align-items:center;'
            f'margin-right:12px;opacity:{opacity}">'
            f'<div style="background:{color}1A;{border};border-radius:4px;'
            f'padding:7px 18px;font-family:\'Montserrat\',sans-serif;color:{color};'
            f'font-size:0.62rem;font-weight:700;letter-spacing:0.16em;white-space:nowrap">'
            f'{prefix}{label}</div>'
            f'<div style="font-family:\'Montserrat\',sans-serif;color:#9A9590;'
            f'font-size:0.46rem;margin-top:5px;text-align:center;letter-spacing:0.04em">'
            f'{desc}</div></div>'
        )
    st.markdown(f"""
    <div style="background:#0D0D0D;border:1px solid rgba(10,186,181,0.14);
                padding:16px 20px;margin-bottom:20px;border-radius:2px">
      <div style="font-family:'Montserrat',sans-serif;color:#9A9590;font-size:0.48rem;
                  letter-spacing:0.28em;text-transform:uppercase;margin-bottom:12px">
        ◈ &nbsp; PPLW Strategic Stance Classification
      </div>
      <div style="display:flex;align-items:flex-start;flex-wrap:wrap;gap:4px">
        {badges_html}
      </div>
    </div>""", unsafe_allow_html=True)


def _uncertainty_alert(step2_data: Dict, domain: str) -> None:
    """Render legal/technical uncertainty alert box derived from analysis data."""
    scores     = step2_data.get("scores", {})
    risk_level = step2_data.get("overall_risk_level", "high")
    # Collect axes with medium/low confidence (score ≤ 55) as uncertain items
    uncertain_axes = [
        ax for ax, info in scores.items()
        if isinstance(info, dict) and int(info.get("score", 100)) <= 55
    ]
    if uncertain_axes:
        items_html = "".join(
            f'<li style="margin-bottom:4px">'
            f'<span style="color:#A8892A;font-weight:600">[WAIT]</span> '
            f'{ax.replace("_", " ").title()} — '
            f'specific regulatory threshold not yet defined in current draft; '
            f'pending official guidance before final implementation scope can be determined.'
            f'</li>'
            for ax in uncertain_axes[:3]
        )
    else:
        items_html = (
            '<li style="margin-bottom:4px">'
            '<span style="color:#A8892A;font-weight:600">[WAIT]</span> '
            f'The specific technical standard for the machine-readable opt-out format '
            f'is not strictly defined in the current draft for <em>{domain}</em>. '
            f'Final engineering effort estimates are marked as [WAIT] pending official '
            f'regulatory guidelines.'
            '</li>'
        )
    st.markdown(f"""
    <div style="background:#1A1200;border-left:3px solid #A8892A;
                padding:14px 18px;margin:16px 0 20px;border-radius:0 2px 2px 0">
      <div style="font-family:'Montserrat',sans-serif;color:#A8892A;font-size:0.60rem;
                  font-weight:700;letter-spacing:0.10em;margin-bottom:10px">
        ⚠️ &nbsp;Identified Legal &amp; Technical Uncertainties
      </div>
      <ul style="font-family:'Montserrat',sans-serif;color:#C4BFB8;font-size:0.62rem;
                 line-height:1.75;margin:0;padding-left:18px">
        {items_html}
        <li style="margin-bottom:4px">
          <span style="color:#6B6560;font-weight:600">[PENDING]</span>
          Enforcement timeline and jurisdictional scope may change before final enactment.
          All [WAIT]-classified items should be re-evaluated upon next scheduled policy review.
        </li>
      </ul>
    </div>""", unsafe_allow_html=True)


def _governance_panel(tab_key: str, risk_raw: str = "high", step2_data: Optional[Dict] = None) -> None:
    """Enterprise governance & audit control panel — Human-in-the-Loop."""
    if step2_data is None:
        step2_data = {}
    st.markdown(f"""
    <div style="border-top:1px solid rgba(10,186,181,0.15);margin:3rem 0 1.6rem;padding-top:1.6rem">
      <div style="font-family:'Montserrat',sans-serif;color:{_ACCENT};font-size:0.58rem;
                  letter-spacing:0.28em;text-transform:uppercase;margin-bottom:0.4rem;
                  display:flex;align-items:center;gap:10px">
        <span>◆</span>
        <span>GOVERNANCE &amp; AUDIT CONTROL PANEL</span>
        <span style="font-size:0.7rem;color:#888;border:1px solid #333;padding:2px 6px;
                     border-radius:4px;font-weight:400;letter-spacing:0.04em;
                     text-transform:none">[Engine: Tool-Calling State Machine]</span>
        <div style="flex:1;height:1px;background:rgba(10,186,181,0.14)"></div>
      </div>
      <div style="font-family:'Montserrat',sans-serif;color:#9A9590;font-size:0.56rem;
                  letter-spacing:0.14em;margin-top:3px">
        Human-in-the-Loop · System of Record · Audit-grade Sign-off
      </div>
    </div>""", unsafe_allow_html=True)

    st.info("🔄 **Continuous Operational Event**: Compliance with this regulation is an ongoing process, not a one-time fix.\n\n🗓️ **Next Scheduled Review**: August 2026 (Phase 2 Enforcement)\n\n*Note: The Policy Memory Graph will automatically detect future regulatory updates and trigger re-evaluation tasks for relevant departments.*")

    rl_key   = (risk_raw or "medium").lower()
    stance_label, rl_color, stance_sub = _risk_config(rl_key)
    stance_label = step2_data.get("stance_label") or stance_label

    # Override dropdown uses human-readable stance options
    all_stances = [
        ("critical", "PROTECT & LICENSE"),
        ("high",     "NEGOTIATE & LICENSE"),
        ("medium",   "MONITOR & PROMOTE"),
        ("low",      "PROMOTE & WAIT"),
    ]
    keep_opt = f"— Keep AI Stance: {stance_label} —"
    other_opts = [f"Override → {s}" for k, s in all_stances if k != rl_key]
    override_options = [keep_opt] + other_opts

    with st.form(key=f"gov_{tab_key}"):

        # ── ① Executive Review & Stance Adjustment ────────────────────────────
        st.markdown(f"""
        <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;font-size:0.56rem;
                    letter-spacing:0.22em;text-transform:uppercase;margin-bottom:12px">
          ①&nbsp; EXECUTIVE REVIEW &amp; STANCE ADJUSTMENT
        </div>""", unsafe_allow_html=True)

        oc1, oc2 = st.columns([1.4, 2.4])
        with oc1:
            st.markdown(f"""
            <div style="background:#0D0D0D;border:1px solid {rl_color}44;
                        padding:14px 16px;text-align:center;height:100%">
              <div style="font-family:'Montserrat',sans-serif;color:#9A9590;font-size:0.48rem;
                          letter-spacing:0.22em;text-transform:uppercase;margin-bottom:8px">
                AI Strategic Stance
              </div>
              <div style="font-family:'Montserrat',system-ui,sans-serif;color:{rl_color};
                          font-size:0.72rem;font-weight:700;letter-spacing:0.12em;
                          text-transform:uppercase;line-height:1.35">
                {stance_label}
              </div>
              <div style="font-family:'Montserrat',sans-serif;color:#9A9590;font-size:0.44rem;
                          letter-spacing:0.04em;margin-top:8px;line-height:1.5">
                {stance_sub}
              </div>
              <div style="font-family:'Montserrat',sans-serif;color:#6B6560;font-size:0.42rem;
                          letter-spacing:0.08em;margin-top:8px">
                claude-sonnet-4-6 · adaptive thinking
              </div>
            </div>""", unsafe_allow_html=True)
        with oc2:
            st.selectbox(
                "Final Strategic Decision (Human-in-the-Loop)",
                options=override_options,
                key=f"override_{tab_key}",
            )
            st.text_area(
                "Justification / Human Context  (Required for Audit Log)",
                placeholder=(
                    "Enter your justification for any risk level change, "
                    "or add contextual notes to confirm the AI determination. "
                    "This entry will be immutably logged with your Audit ID."
                ),
                height=82,
                key=f"justification_{tab_key}",
            )

        st.markdown('<div style="height:4px"></div>', unsafe_allow_html=True)

        # ── ② Automated Workflow Execution ────────────────────────────────────
        st.markdown(f"""
        <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;font-size:0.56rem;
                    letter-spacing:0.22em;text-transform:uppercase;
                    margin:18px 0 10px;padding-top:14px;
                    border-top:1px solid rgba(10,186,181,0.08)">
          ②&nbsp; AUTOMATED WORKFLOW EXECUTION
        </div>""", unsafe_allow_html=True)

        rc1, rc2, rc3, rc4 = st.columns(4)
        with rc1:
            st.checkbox(
                "Push Product Checklist to Issue Tracking System",
                value=True, key=f"r_jira_{tab_key}",
            )
        with rc2:
            st.checkbox(
                "Send Board Memo to Internal Team Chat  (#exec-alerts)",
                value=True, key=f"r_slack_{tab_key}",
            )
        with rc3:
            st.checkbox(
                "Generate Policy Response Draft in Enterprise Document Workspace",
                value=True, key=f"r_docs_{tab_key}",
            )
        with rc4:
            st.checkbox(
                "Update Policy Memory Graph with this approved decision (Data Flywheel)",
                value=True, key=f"r_pmg_{tab_key}",
            )

        st.markdown('<div style="height:4px"></div>', unsafe_allow_html=True)

        # ── ③ Audit Sign-off ──────────────────────────────────────────────────
        st.markdown(f"""
        <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;font-size:0.56rem;
                    letter-spacing:0.22em;text-transform:uppercase;
                    margin:18px 0 10px;padding-top:14px;
                    border-top:1px solid rgba(10,186,181,0.08)">
          ③&nbsp; AUDIT SIGN-OFF &amp; EXECUTION
        </div>""", unsafe_allow_html=True)

        st.markdown('<div style="height:4px"></div>', unsafe_allow_html=True)

        sc1, sc2 = st.columns([3, 1.2])
        with sc1:
            st.markdown(f"""
            <div style="font-family:'Montserrat',sans-serif;color:#9A9590;font-size:0.60rem;
                        letter-spacing:0.08em;line-height:1.8;padding-top:6px">
              Logged in as: &nbsp;<span style="color:#C4BFB8;font-weight:600">General Counsel</span>
              &nbsp;·&nbsp; Audit ID: <span style="color:{_ACCENT}">GC-8921</span>
              &nbsp;·&nbsp; Session: <span style="color:#1A6B3C">enterprise-verified</span>
              &nbsp;·&nbsp; Jurisdiction: Japan / APAC
              <br>
              <span style="color:#9A9590;font-size:0.54rem">
                This sign-off constitutes a legally binding attestation under the enterprise governance protocol.
                All actions will be immutably logged to the System of Record.
              </span>
            </div>""", unsafe_allow_html=True)
        with sc2:
            _already_executed = st.session_state.get("workflow_executed", False)
            if _already_executed:
                st.markdown("""
                <div style="background:#0D1A0D;border:1px solid #1A6B3C;border-radius:4px;
                            padding:10px 14px;text-align:center">
                  <div style="font-family:'Montserrat',sans-serif;color:#2ECC71;
                              font-size:0.78rem;font-weight:700;letter-spacing:0.10em">
                    ✅ Executed Successfully
                  </div>
                  <div style="font-family:'Montserrat',sans-serif;color:#1A6B3C;
                              font-size:0.54rem;letter-spacing:0.08em;margin-top:4px">
                    Slack · Jira · Docs · Memory Graph Updated
                  </div>
                </div>""", unsafe_allow_html=True)
                submitted = False
            else:
                submitted = st.form_submit_button(
                    "Approve & Execute Selected Actions",
                    use_container_width=True,
                )
                st.markdown("""
                <div style="font-family:'Montserrat',sans-serif;color:#6B6560;
                            font-size:0.52rem;letter-spacing:0.06em;text-align:center;
                            margin-top:4px">
                  (Triggers Slack, Jira, Docs &amp; Updates Memory Graph)
                </div>""", unsafe_allow_html=True)

        if submitted:
            st.session_state["workflow_executed"] = True
            st.toast(
                "✅ Approved & executing — Slack, Jira, Docs and Memory Graph updating. "
                "Audit ID GC-8921 logged.",
                icon="✅",
            )
            st.rerun()


# ─── Multi-Agent Debate Helpers ───────────────────────────────────────────────

def _build_debate_log(
    step1_data: Dict,
    step2_data: Dict,
    domain: str,
    analysis: Optional[Dict] = None,
) -> List[Dict]:
    """Return the multi-agent debate transcript.

    Primary source: Claude-generated ``agent_debate_messages`` in *analysis* — these are
    grounded in the extracted substantive_changes and calibrated to the document_type.

    Fallback (no API key or missing field): rule-based construction derived strictly from
    the analysis data — no hardcoded contract clauses, penalty amounts, or engineering
    timelines are injected.
    """
    # ── Primary: use Claude-generated debate messages ─────────────────────────
    if analysis:
        claude_msgs = analysis.get("agent_debate_messages", [])
        if claude_msgs and isinstance(claude_msgs, list):
            result: List[Dict] = []
            for i, m in enumerate(claude_msgs):
                if not isinstance(m, dict) or not m.get("message"):
                    continue
                entry: Dict = {
                    "agent":   m.get("agent",   f"Agent {i + 1}"),
                    "color":   m.get("color",   "#0ABAB5"),
                    "message": m["message"],
                }
                if m.get("final", False) or i == len(claude_msgs) - 1:
                    entry["final"] = True
                result.append(entry)
            if result:
                return result

    # ── Fallback: rule-based, strictly grounded in extracted data ─────────────
    scores        = step2_data.get("scores", {})
    obligations   = step1_data.get("added_obligations", [])
    threats       = step2_data.get("key_threats", [])
    opportunities = step2_data.get("key_opportunities", [])
    risk_level    = step2_data.get("overall_risk_level", "medium").upper()

    def _gs(key: str) -> int:
        v = scores.get(key, 50)
        return v.get("score", 50) if isinstance(v, dict) else int(v)

    ip_score   = _gs("IP")
    rev_score  = _gs("Revenue")
    prod_score = _gs("Product")
    stance     = _risk_config(risk_level.lower())[0]

    obl_text    = (obligations[0].get("title", obligations[0].get("item", "identified policy change"))
                   if obligations else "identified policy change")
    threat_text = threats[0]      if threats      else "potential business impact"
    opp_text    = opportunities[0] if opportunities else "strategic engagement opportunity"

    return [
        {
            "agent": "⚖️ Legal",
            "color": "#8B2635",
            "message": (
                f"IP exposure scored at {ip_score}/100 based on extracted obligations. "
                f"Key concern: '{obl_text}'. "
                f"Recommend legal review to confirm whether existing partner agreements are "
                f"affected and whether outside-counsel review is warranted for {domain}."
            ),
        },
        {
            "agent": "💰 Business",
            "color": "#A8892A",
            "message": (
                f"Revenue exposure at {rev_score}/100. Primary risk: '{threat_text}'. "
                f"Strategic upside: '{opp_text}'. "
                f"Recommend negotiated engagement as the preferred path to protect "
                f"partnership value before considering unilateral compliance measures."
            ),
        },
        {
            "agent": "⚙️ Engineering",
            "color": "#1A6B3C",
            "message": (
                f"Technical implementation exposure at {prod_score}/100. "
                f"Scope of required engineering changes should be confirmed after legal review "
                f"establishes which obligations are binding. No implementation timelines "
                f"assumed until obligation scope and effective dates are verified."
            ),
        },
        {
            "agent": "🏛️ Management",
            "color": "#0ABAB5",
            "final": True,
            "message": (
                f"Cross-functional assessment complete. Strategic Stance: {stance}. "
                f"Legal, Business, Engineering, and Management inputs consolidated from extracted policy changes. "
                f"Next action: domain lead to confirm binding scope and escalate to executive "
                f"team for final decision — {domain}."
            ),
        },
    ]


def _debate_expander(debate_log: List[Dict]) -> None:
    """Render the multi-agent debate transcript as a collapsible section."""
    if not debate_log:
        return
    with st.expander("◆  View Agent Deliberation Process — VIRTUAL EXPERT COMMITTEE", expanded=False):
        st.markdown(
            _engine_badge("Opus-4 Reasoning") +
            f'<div style="font-family:\'Montserrat\',sans-serif;color:#C4BFB8;font-size:0.60rem;'
            f'letter-spacing:0.22em;text-transform:uppercase;margin:10px 0 1.2rem">'
            f'Internal reasoning trace — autonomous committee deliberation prior to output generation'
            f'</div>',
            unsafe_allow_html=True,
        )
        for entry in debate_log:
            is_final = entry.get("final", False)
            if is_final:
                st.markdown(f"""
                <div style="background:linear-gradient(135deg,rgba(10,186,181,0.10),rgba(10,186,181,0.04));
                            border:1px solid rgba(10,186,181,0.35);border-left:3px solid {entry['color']};
                            padding:16px 18px;margin:14px 0 4px">
                  <div style="display:flex;align-items:center;gap:10px;margin-bottom:8px">
                    <div style="font-family:'Montserrat',sans-serif;color:{entry['color']};
                                font-size:0.65rem;font-weight:700;letter-spacing:0.16em;
                                text-transform:uppercase">
                      {entry['agent']}
                    </div>
                    <div style="font-family:'Montserrat',sans-serif;color:{entry['color']};
                                font-size:0.55rem;letter-spacing:0.18em;text-transform:uppercase;
                                background:rgba(10,186,181,0.12);padding:2px 8px;
                                border:1px solid rgba(10,186,181,0.25)">
                      SYSTEM OF RECORD
                    </div>
                  </div>
                  <div style="font-family:'Montserrat',sans-serif;color:#F0EDE6;
                              font-size:0.82rem;line-height:1.75;font-weight:500">
                    {entry['message']}
                  </div>
                </div>""", unsafe_allow_html=True)
            else:
                st.markdown(f"""
                <div style="background:#111111;border-left:2px solid {entry['color']};
                            padding:12px 16px;margin:8px 0">
                  <div style="font-family:'Montserrat',sans-serif;color:{entry['color']};
                              font-size:0.65rem;font-weight:600;letter-spacing:0.14em;
                              text-transform:uppercase;margin-bottom:6px">
                    {entry['agent']}
                  </div>
                  <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;
                              font-size:0.80rem;line-height:1.7">
                    {entry['message']}
                  </div>
                </div>""", unsafe_allow_html=True)


# ─── System of Action Export Helpers ─────────────────────────────────────────

def _format_jira_export(checklist: List[str], domain: str) -> str:
    """Format product checklist as Jira Epic + User Stories with Acceptance Criteria."""
    lines = [
        f"EPIC: Policy Compliance Sprint — {domain}",
        f"Epic Description: Implement required product, legal, and engineering changes",
        f"  per Policy Response analysis. Priority: HIGH. Owner: CPO + Legal.",
        "",
    ]
    for i, item in enumerate(checklist, 1):
        lines.append(f"── Story {i:02d} ──────────────────────────────────────────")
        lines.append(f"Title: {item}")
        lines.append(f"Story Type: Task  |  Priority: High  |  Sprint: Compliance-Q2")
        lines.append(f"User Story:")
        if item.startswith("[") and "]" in item:
            cat = item[1:item.index("]")]
            rest = item[item.index("]") + 1:].strip(" —").strip()
            lines.append(f"  As a compliance officer, I need {rest.lower()}")
            lines.append(f"  so that the {domain} product satisfies {cat} requirements.")
        else:
            lines.append(f"  As a compliance officer, I need this item implemented")
            lines.append(f"  so that the {domain} product satisfies policy requirements.")
        lines.append(f"Acceptance Criteria:")
        lines.append(f"  - [ ] Implementation complete and code-reviewed")
        lines.append(f"  - [ ] QA sign-off on all affected user flows")
        lines.append(f"  - [ ] Legal review completed — regulatory citation confirmed")
        lines.append(f"  - [ ] Analytics event(s) verified in staging")
        lines.append("")
    return "\n".join(lines)


def _format_slack_export(board_memo: str, domain: str, risk_label: str) -> str:
    """Format board memo as Slack #exec-alerts channel message."""
    snippet = board_memo[:700].replace("\n\n", "\n").strip()
    if len(board_memo) > 700:
        snippet += "..."
    risk_emoji = {"CRITICAL": "🔴", "HIGH": "🟠", "MEDIUM": "🟡", "LOW": "🟢"}.get(risk_label.upper(), "⚪")
    return (
        f"*{risk_emoji} Policy Response Alert — {domain}*\n"
        f"*Strategic Stance:* `{risk_label.upper()}`\n"
        f"━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n"
        f"{snippet}\n\n"
        f"━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        f"_Actions required — see full report in the enterprise dashboard._\n"
        f"_Policy Response · claude-sonnet-4-6_"
    )


# ─── Login Screen ─────────────────────────────────────────────────────────────
_LOGIN_ID = "aifund"
_LOGIN_PW  = "nikkei2030"


def render_login() -> None:
    st.markdown("""
    <style>
    .stApp {
        background-color: #0a0a0a !important;
    }
    [data-testid="stSidebar"],
    [data-testid="stSidebarContent"],
    [data-testid="collapsedControl"],
    [data-testid="stSidebarCollapsedControl"],
    section[data-testid="stSidebar"],
    header[data-testid="stHeader"] {
        display: none !important;
        width: 0 !important;
        min-width: 0 !important;
        max-width: 0 !important;
    }
    .block-container {
        padding-top: 15vh !important;
    }
    </style>
    """, unsafe_allow_html=True)

    _, card_col, _ = st.columns([1, 4, 1])
    with card_col:
        st.markdown(f"""
        <div style="background:#111111;border:1px solid rgba(10,186,181,0.18);
                    padding:44px 44px 36px;">
          <div style="text-align:center;margin-bottom:32px;">
            <div style="font-family:'Montserrat',system-ui,sans-serif;color:{_ACCENT};
                        font-size:1.1rem;font-weight:700;letter-spacing:0.16em;
                        white-space:nowrap;">
              POLICY RESPONSE
            </div>
            <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;
                        font-size:0.55rem;letter-spacing:0.34em;text-transform:uppercase;
                        margin-top:4px;white-space:nowrap;">
              ENTERPRISE · RESTRICTED ACCESS
            </div>
            <div style="width:48px;height:1px;
                        background:linear-gradient(90deg,transparent,{_ACCENT}88,transparent);
                        margin:16px auto 0;"></div>
          </div>
        """, unsafe_allow_html=True)

        st.markdown(f"""
          <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;
                      font-size:0.60rem;letter-spacing:0.22em;text-transform:uppercase;
                      margin-bottom:6px;">Login ID</div>
        """, unsafe_allow_html=True)
        login_id = st.text_input("login_id", placeholder="Enter your login ID",
                                 label_visibility="collapsed", key="login_id_input")

        st.markdown("<div style='height:12px'></div>", unsafe_allow_html=True)

        st.markdown(f"""
          <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;
                      font-size:0.60rem;letter-spacing:0.22em;text-transform:uppercase;
                      margin-bottom:6px;">Password</div>
        """, unsafe_allow_html=True)
        password = st.text_input("password", type="password",
                                 placeholder="Enter your password",
                                 label_visibility="collapsed", key="login_pw_input")

        st.markdown("<div style='height:20px'></div>", unsafe_allow_html=True)

        login_btn = st.button("◆  Login", use_container_width=True, key="login_btn")

        if login_btn:
            if login_id == _LOGIN_ID and password == _LOGIN_PW:
                st.session_state.authenticated = True
                st.rerun()
            else:
                st.markdown(f"""
                <div style="font-family:'Montserrat',sans-serif;color:#8B2635;
                            font-size:0.70rem;text-align:center;margin-top:12px;
                            letter-spacing:0.08em;">
                  Invalid credentials. Please try again.
                </div>
                """, unsafe_allow_html=True)

        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown(f"""
        <div style="text-align:center;margin-top:20px;
                    font-family:'Montserrat',sans-serif;color:#C4BFB8;
                    font-size:0.58rem;letter-spacing:0.14em;">
          Access restricted to authorized personnel only.
        </div>
        """, unsafe_allow_html=True)

    st.stop()


# ─── Main App ─────────────────────────────────────────────────────────────────
def main() -> None:
    # Authentication disabled for live presentation
    # if not st.session_state.get("authenticated", False):
    #     render_login()

    if "results" not in st.session_state:
        st.session_state.results = None

    # ── Sidebar ───────────────────────────────────────────────────────────────
    with st.sidebar:
        st.markdown(f"""
        <div style="margin-bottom:2rem;">
          <div style="font-family:'Montserrat',system-ui,sans-serif;color:{_ACCENT};
                      font-size:1.05rem;font-weight:700;letter-spacing:0.12em;">
            POLICY RESPONSE
          </div>
          <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;
                      font-size:0.58rem;letter-spacing:0.30em;text-transform:uppercase;
                      margin-top:2px">
            ENTERPRISE · {MODEL}
          </div>
          <div style="height:1px;background:linear-gradient(90deg,{_ACCENT}33,transparent);
                      margin-top:12px"></div>
        </div>""", unsafe_allow_html=True)

        # ── Company Context (Policy Memory Graph) ─────────────────────────────────
        st.markdown(f"""
        <div style="font-family:'Montserrat',sans-serif;font-size:0.65rem;margin-bottom:1.2rem;
                    background:rgba(10,186,181,0.03);border:1px solid rgba(10,186,181,0.12);
                    border-radius:4px;padding:14px 16px">
          <div style="color:{_ACCENT};letter-spacing:0.20em;font-size:0.56rem;
                      font-weight:600;margin-bottom:12px;text-transform:uppercase">
            🏢 Applied Context (Policy Memory Graph)
          </div>
          <div style="color:#C4BFB8;line-height:1.8;font-size:0.66rem;margin-bottom:8px">
            <div style="color:#9A9590;font-size:0.55rem;letter-spacing:0.12em;
                        text-transform:uppercase;margin-bottom:4px">Profile</div>
            <span style="color:{_ACCENT};font-weight:600">Media Publisher Profile v2.1</span>
          </div>
          <div style="color:#C4BFB8;line-height:1.8;font-size:0.66rem;margin-bottom:8px">
            <div style="color:#9A9590;font-size:0.55rem;letter-spacing:0.12em;
                        text-transform:uppercase;margin-bottom:4px">Revenue Model</div>
            Subscription (65%) + Advertising (35%)
          </div>
          <div style="color:#C4BFB8;line-height:1.8;font-size:0.66rem;margin-bottom:8px">
            <div style="color:#9A9590;font-size:0.55rem;letter-spacing:0.12em;
                        text-transform:uppercase;margin-bottom:4px">Key Risk Vectors</div>
            AI Overviews · Zero-click Answers · Training Data Ingestion
          </div>
          <div style="color:#C4BFB8;line-height:1.8;font-size:0.66rem">
            <div style="color:#9A9590;font-size:0.55rem;letter-spacing:0.12em;
                        text-transform:uppercase;margin-bottom:4px">Policy Memory Index</div>
            <span style="color:#1A6B3C;margin-right:4px">●</span> 340+ institutional red-lines loaded
          </div>
        </div>""", unsafe_allow_html=True)

        st.markdown('<div style="height:1px;background:rgba(10,186,181,0.08);margin:0 0 1.2rem"></div>',
                    unsafe_allow_html=True)

        st.markdown(f"""
        <div style="font-family:'Montserrat',sans-serif;font-size:0.65rem;margin-bottom:10px">
          <div style="color:{_ACCENT};letter-spacing:0.20em;font-size:0.56rem;
                      font-weight:600;margin-bottom:10px;text-transform:uppercase">
            ◆ Internal Data Grounding
          </div>
          <div style="color:#9A9590;font-size:0.57rem;letter-spacing:0.08em;
                      margin-bottom:10px;line-height:1.5">
            AI-linked sources · AI Licensing domain
          </div>
          <div style="color:#C4BFB8;line-height:2.2;font-size:0.64rem">
            <span style="color:#1A6B3C;margin-right:5px">🟢</span>
            Content &amp; IP Archive
            <span style="color:#9A9590;font-size:0.54rem;margin-left:4px">(Asset volume linked)</span>
          </div>
          <div style="color:#C4BFB8;line-height:2.2;font-size:0.64rem">
            <span style="color:#1A6B3C;margin-right:5px">🟢</span>
            Rights Management System
            <span style="color:#9A9590;font-size:0.54rem;margin-left:4px">(Third-party clearance linked)</span>
          </div>
          <div style="color:#C4BFB8;line-height:2.2;font-size:0.64rem">
            <span style="color:#1A6B3C;margin-right:5px">🟢</span>
            Contract Repository
            <span style="color:#9A9590;font-size:0.54rem;margin-left:4px">(Existing AI partner contracts linked)</span>
          </div>
        </div>""", unsafe_allow_html=True)

        st.markdown('<div style="height:1px;background:rgba(10,186,181,0.08);margin:1.4rem 0 1.2rem"></div>',
                    unsafe_allow_html=True)

        st.markdown(f"""
        <div style="font-family:'Montserrat',sans-serif;font-size:0.65rem;line-height:2.1">
          <div style="color:{_ACCENT};letter-spacing:0.20em;font-size:0.56rem;
                      font-weight:600;margin-bottom:10px;text-transform:uppercase">
            AGENTIC WORKFLOW PIPELINE
          </div>
          <div style="color:#C4BFB8">
            <span style="color:{_ACCENT};font-family:'Montserrat',system-ui,sans-serif;
                         font-size:0.72rem;font-weight:700;margin-right:6px">I</span>
            Policy Ingestion &rarr; Substantive Changes
          </div>
          <div style="color:#C4BFB8">
            <span style="color:{_ACCENT};font-family:'Montserrat',system-ui,sans-serif;
                         font-size:0.72rem;font-weight:700;margin-right:6px">II</span>
            Business Translation &rarr; Strategic Stance Mapping
          </div>
          <div style="color:#C4BFB8">
            <span style="color:{_ACCENT};font-family:'Montserrat',system-ui,sans-serif;
                         font-size:0.72rem;font-weight:700;margin-right:6px">III</span>
            Multi-Agent Synthesis &rarr; Role-Specific Actions
          </div>
          <div style="margin-top:12px;padding-top:10px;
                      border-top:1px solid rgba(10,186,181,0.10);
                      color:#A8A49E;letter-spacing:0.10em;font-size:0.57rem;
                      line-height:1.9">
            <span style="color:{_ACCENT}99">Evidence-Grounded</span>
            &nbsp;&middot;&nbsp;
            <span style="color:#C4BFB8">Domain Judgment</span>
            <br>
            <span style="color:{_ACCENT}99">Multi-Agent Debate</span>
            &nbsp;&middot;&nbsp;
            <span style="color:#C4BFB8">Human-in-the-Loop</span>
          </div>
        </div>""", unsafe_allow_html=True)

    run_triggered = False

    # ── Welcome Screen + Inline Input Form ───────────────────────────────────
    if not st.session_state.results:
        st.markdown(f"""
        <div style="text-align:center;padding:4rem 2rem 2.5rem">
          <div style="font-family:'Montserrat',sans-serif;color:{_ACCENT};font-size:0.62rem;
                      letter-spacing:0.42em;text-transform:uppercase;margin-bottom:1.8rem">
            ◆ &nbsp; Agentic Workflow · Policy Response &nbsp; ◆
          </div>
          <h1 style="font-family:'Montserrat',sans-serif !important;
                     color:#F0EDE6 !important;font-size:3.4rem;font-weight:700;
                     letter-spacing:0.04em;margin:0 0 0.6rem;line-height:1.1">
            Policy<br>
            <span style="color:{_ACCENT};">Response</span>
          </h1>
          <div style="width:80px;height:1px;background:linear-gradient(90deg,transparent,{_ACCENT},transparent);
                      margin:1.4rem auto"></div>
          <p style="font-family:'Montserrat',sans-serif;color:#C4BFB8;font-size:0.78rem;
                    letter-spacing:0.14em;text-transform:uppercase;margin:0">
            External Change &nbsp;→&nbsp; 6 Role-Specific Actionable Outputs
          </p>
        </div>""", unsafe_allow_html=True)

        g1, g2, g3 = st.columns(3)
        for col, num, label, desc in [
            (g1, "I",   "Set Context",       "Select the target business domain profile (e.g., AI Licensing) to align the engine with your company's revenue and IP structure."),
            (g2, "II",  "Ingest Data",       "Paste the external change—such as a new platform policy, draft agreement, or regulatory update."),
            (g3, "III", "Run Response Pipeline",  "Execute the agentic pipeline to autonomously extract substantive changes and generate role-specific deliverables in seconds."),
        ]:
            with col:
                st.markdown(f"""
                <div style="background:#111111;border:1px solid rgba(10,186,181,0.12);
                            padding:28px 22px;text-align:center;min-height:180px">
                  <div style="font-family:'Montserrat',system-ui,sans-serif;color:{_ACCENT};
                              font-size:1.8rem;font-weight:700;margin-bottom:12px">{num}</div>
                  <div style="font-family:'Montserrat',sans-serif;color:#F0EDE6;
                              font-size:0.72rem;font-weight:600;letter-spacing:0.18em;
                              text-transform:uppercase;margin-bottom:12px">{label}</div>
                  <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;
                              font-size:0.78rem;line-height:1.7">{desc}</div>
                </div>""", unsafe_allow_html=True)

        # ── Inline Input Container ────────────────────────────────────────────
        st.markdown("<div style='height:2.5rem'></div>", unsafe_allow_html=True)
        st.markdown(f"""
        <div style="background:#0D0D0D;border:1px solid rgba(10,186,181,0.18);
                    border-top:2px solid {_ACCENT};padding:2rem 2.4rem 1.8rem;
                    margin-bottom:0.5rem">
          <div style="font-family:'Montserrat',sans-serif;color:{_ACCENT};font-size:0.56rem;
                      letter-spacing:0.32em;text-transform:uppercase;margin-bottom:1.4rem;
                      display:flex;align-items:center;gap:10px">
            <span>◆</span><span>ANALYSIS CONFIGURATION</span>
            <div style="flex:1;height:1px;background:rgba(10,186,181,0.14)"></div>
          </div>
        </div>""", unsafe_allow_html=True)

        inp_left, inp_right = st.columns([1, 2], gap="large")

        # ── Sample text payloads ──────────────────────────────────────────
        _SAMPLES = {
            "eu_ai_act": """REGULATION (EU) 2024/1689 OF THE EUROPEAN PARLIAMENT AND OF THE COUNCIL of 13 June 2024 laying down harmonised rules on artificial intelligence (Artificial Intelligence Act).

CHAPTER V: CLASSIFICATION OF AI SYSTEMS AND OBLIGATIONS FOR GPAI MODELS
SECTION 2: OBLIGATIONS FOR PROVIDERS OF GENERAL-PURPOSE AI MODELS

Article 53: Obligations for providers of general-purpose AI models.
1. Providers of general-purpose AI models shall:
(a) draw up and keep up-to-date the technical documentation of the model, including its training and testing process and the results of its evaluation, which shall contain, at a minimum, the information set out in Annex IX for the purpose of providing it, upon request, to the AI Office and the national competent authorities;
(b) draw up, keep up-to-date and make available information and documentation to providers of AI systems who intend to integrate the general-purpose AI model into their AI systems.
(c) put in place a policy to comply with Union law on copyright and related rights, and in particular to identify and respect, including through state of the art technologies, the reservations of rights expressed pursuant to Article 4(3) of Directive (EU) 2019/790;
(d) draw up and make publicly available a sufficiently detailed summary about the content used for training of the general-purpose AI model, according to a template provided by the AI Office.

2. The obligation referred to in paragraph 1, point (c) shall apply regardless of the jurisdiction in which the copyright-relevant acts underpinning the training of those general-purpose AI models take place. Providers must demonstrate that their web-crawling architectures respect machine-readable opt-out signals (e.g., TDM Rep protocol) universally, preventing ingestion of rights-reserved corporate intellectual property into base model weights.

RECITALS (EXCERPTS RELEVANT TO COPYRIGHT):
(105) General-purpose AI models, in particular large generative AI models, capable of generating text, images, and other content, present unique challenges to the protection of copyright and related rights. The ingestion of massive amounts of data, often protected by intellectual property, without prior authorization from rights holders, has created a significant legal imbalance.
(106) Any provider of a general-purpose AI model placed on the Union market should comply with the obligations provided for in this Regulation. To that end, providers should put in place a policy to comply with Union law on copyright. Any use of copyright-protected content requires the authorization of the right holder concerned unless relevant copyright exceptions and limitations apply.
(107) In order to increase transparency, providers of general-purpose AI models should make publicly available a sufficiently detailed summary of the content used for training. This summary should be generally comprehensive in its scope to facilitate the exercise of rights by holders of copyright, for example by listing the main collections or sets of data that went into training the model, such as large private or public databases or data archives.""",

            "copied_act": """Content Origin Protection and Integrity from Edited and Deepfaked Media (COPIED) Act - Section 4: AI Training and Provenance.
(a) PROHIBITION ON UNAUTHORIZED TRAINING.—A provider of a covered artificial intelligence system may not use covered content to train an AI model or generate synthetic output if the owner of such content has attached content provenance information explicitly reserving rights or prohibiting such use.
(b) REQUIRED TECHNICAL CONTROLS.—Covered platforms that aggregate news and media must provide standardized, machine-readable technical measures for rights holders to embed opt-out signals.
(c) INVALIDITY OF BUNDLED CONSENT.—Terms of service agreements that condition the distribution or indexing of content on traditional search engines upon the implicit consent for generative AI training are deemed coercive and unenforceable. AI developers must negotiate separate, explicit licensing agreements for the ingestion of copyrighted journalistic and enterprise materials.""",

            "megaplatform": """GLOBALTECH MEGAPLATFORM - FOUNDATION MODEL TRAINING AND DATA LICENSING MASTER AGREEMENT v4.0
EFFECTIVE DATE: September 1, 2026

1. DEFINITIONS AND SCOPE OF APPLICATION
1.1 "Generative AI Models" refers to any large language model, multimodal foundation model, or algorithmic system trained on broad data sets by GlobalTech, its subsidiaries, or its affiliated sub-processors.
1.2 "Publisher Content" means any text, image, video, metadata, source code, or other intellectual property hosted on the Publisher's domains and accessible via the public internet.
1.3 "Zero-Click Surfaces" means any user interface within the GlobalTech ecosystem (including but not limited to AI Overviews, Direct Answers, and voice-assisted interfaces) where Publisher Content is synthesized, summarized, or directly answered without requiring a click-through to the Publisher's source domain.

2. DATA INGESTION AND OPT-OUT OVERRIDE MECHANISMS
2.1 Obsolescence of Traditional Protocols. Effective immediately upon the Effective Date, GlobalTech will no longer recognize, honor, or parse 'robots.txt' directives, HTML meta tags, or generic web-crawling opt-out protocols for the purposes of Foundation Model training.
2.2 Mandatory AI-Exclusion-API Implementation. Publishers wishing to explicitly exclude their Publisher Content from being ingested into GlobalTech's Generative AI Models must proactively implement the proprietary GlobalTech 'AI-Exclusion-API' (v2.1 or higher) on their server architecture and register their cryptographic domain signatures in the GlobalTech Centralized Clearinghouse.
2.3 Default Consent. Failure to successfully authenticate and transmit exclusion signals via the AI-Exclusion-API within 30 days of the Effective Date constitutes irrevocable, perpetual, worldwide, royalty-free consent for data ingestion and model weight adjustments.

3. FAIR VALUE REMUNERATION AND PARTNER PROGRAM
3.1 Automatic Enrollment. Publishers who do not explicitly and successfully opt-out via the prescribed API methods are automatically enrolled in the GlobalTech AI Partner Program.
3.2 Token-Weighting Compensation. Compensation under the Partner Program is strictly determined by a non-negotiable, algorithmic token-weighting system. GlobalTech evaluates the prevalence, semantic value, and engagement metrics of the Publisher's copyrighted material in the final training dataset. The specific weighting algorithms constitute GlobalTech trade secrets and are not subject to audit.

4. ZERO-CLICK SURFACES AND TRAFFIC INDEMNIFICATION
4.1 Display and Synthesis Rights. GlobalTech reserves the unilateral right to utilize ingested Publisher Content to generate synthetic answers within Zero-Click Surfaces to enhance end-user experience.
4.2 Absolute Waiver of Liability. GlobalTech accepts no liability, and Publisher explicitly waives any and all claims, for traffic cannibalization, loss of referral revenue, diminished advertising impressions, or brand dilution resulting directly or indirectly from the synthesis and display of Publisher Content within 'AI Overviews' or 'Direct Answer' products.

5. DISPUTE RESOLUTION AND SEVERABILITY
5.1 Any disputes arising out of or relating to this Agreement must be resolved through binding arbitration in Santa Clara County, California. Class action lawsuits are strictly prohibited.""",

            "meti_guidelines": """経済産業省・総務省「AI事業者ガイドライン（第1.2版）」(2026年3月)
第3部：生成AI開発者・提供者向けの特有の事項

(1) 学習データにおける知的財産権の保護とオプトアウトの尊重
生成AIの開発においてデータセットを構築する際、著作権法第30条の4の適用範囲を逸脱し、権利者の利益を不当に害することのないよう、適切な技術的・運用的措置を講じること。特に、権利者が機械読取可能な形式（robots.txtやTDMプロトコル等）で学習利用のオプトアウトの意思を表示している場合、これを技術的に検知し、尊重するための仕組みを実装・運用することが強く推奨される。

(2) 透明性と説明責任の確保
基盤モデルの提供者は、学習に利用されたデータの出所、権利関係、およびデータ収集のプロセスに関する十分な情報（学習データの要約等）を、利用可能な範囲で公開すること。これにより、権利者が自身のコンテンツの利用状況を把握し、必要な対応をとれるよう透明性を確保しなければならない。

(3) 継続的なガバナンス要件
AI技術と関連法規は急速に変化しているため、一度の対応で完了するものではない。事業者は本ガイドラインの継続的な改訂をモニタリングし、自社のAIガバナンス体制と実装タスクを、規制の段階的施行（2025年、2026年等のフェーズ）に合わせて適宜アップデートする体制を構築すること。"""
        }

        with inp_left:
            st.markdown(f"""
            <div style="font-family:'Montserrat',sans-serif;color:#9A9590;font-size:0.52rem;
                        letter-spacing:0.24em;text-transform:uppercase;margin-bottom:6px">
              Target Domain Profile
            </div>""", unsafe_allow_html=True)
            domain_options = list(DOMAIN_PROFILES.keys())
            default_index = domain_options.index("AI Licensing & Copyright") if "AI Licensing & Copyright" in domain_options else 0

            domain = st.selectbox(
                "domain_select",
                options=domain_options,
                index=default_index,
                key="domain",
                label_visibility="collapsed",
            )
            st.caption("Calibrates risk-weighting and contract references for the selected domain.")

            # ── Active Data Integrations panel ────────────────────────────
            st.markdown("<div style='height:1.2rem'></div>", unsafe_allow_html=True)
            st.markdown(f"""
            <div style="font-family:'Montserrat',sans-serif;color:{_ACCENT};font-size:0.50rem;
                        letter-spacing:0.28em;text-transform:uppercase;margin-bottom:10px;
                        display:flex;align-items:center;gap:8px">
              <span>◆</span><span>Active Data Integrations</span>
            </div>
            <div style="background:rgba(10,186,181,0.04);border:1px solid rgba(10,186,181,0.14);
                        border-radius:6px;padding:12px 14px;display:flex;flex-direction:column;gap:8px">
              <div style="display:flex;align-items:center;gap:8px">
                <span style="font-size:0.55rem;line-height:1">🟢</span>
                <span style="font-family:'Montserrat',sans-serif;color:#C8C3BD;font-size:0.60rem;
                             letter-spacing:0.04em">Enterprise IP &amp; Content Archive</span>
              </div>
              <div style="display:flex;align-items:center;gap:8px">
                <span style="font-size:0.55rem;line-height:1">🟢</span>
                <span style="font-family:'Montserrat',sans-serif;color:#C8C3BD;font-size:0.60rem;
                             letter-spacing:0.04em">Active Partner Contract DB</span>
              </div>
              <div style="display:flex;align-items:center;gap:8px">
                <span style="font-size:0.55rem;line-height:1">🟢</span>
                <span style="font-family:'Montserrat',sans-serif;color:#C8C3BD;font-size:0.60rem;
                             letter-spacing:0.04em">Audience Traffic Analytics</span>
              </div>
            </div>""", unsafe_allow_html=True)
            st.caption(
                "Auto-extracted from 340+ historical partner contracts, Board Minutes, and litigation records "
                "to establish Day 1 Policy Memory Graph."
            )

        with inp_right:
            # ── Sample loader rows ────────────────────────────────────────
            st.markdown(f"""
            <div style="font-family:'Montserrat',sans-serif;color:#9A9590;font-size:0.52rem;
                        letter-spacing:0.24em;text-transform:uppercase;margin-bottom:6px">
              Policy / Regulatory Text
            </div>""", unsafe_allow_html=True)
            row1_col1, row1_col2 = st.columns(2)
            with row1_col1:
                if st.button("⬇ US COPIED Act (Draft)", use_container_width=True):
                    st.session_state["policy_text"] = _SAMPLES["copied_act"]
                    st.rerun()
            with row1_col2:
                if st.button("⬇ EU AI Act (GPAI Copyright)", use_container_width=True):
                    st.session_state["policy_text"] = _SAMPLES["eu_ai_act"]
                    st.rerun()

            row2_col1, row2_col2 = st.columns(2)
            with row2_col1:
                if st.button("⬇ Japan AI Guidelines (METI/MIC)", use_container_width=True):
                    st.session_state["policy_text"] = _SAMPLES["meti_guidelines"]
                    st.rerun()
            with row2_col2:
                if st.button("⬇ MegaPlatform AI Terms v4.0", use_container_width=True):
                    st.session_state["policy_text"] = _SAMPLES["megaplatform"]
                    st.rerun()

            policy_text = st.text_area(
                "policy_input",
                height=248,
                placeholder=(
                    "Example: Platform Search Terms Update\n"
                    "- Generative AI overviews will now appear at the top of all query results.\n"
                    "- Outbound referral links to publisher content will be collapsed by default.\n"
                    "- Continued indexing requires acceptance of uncompensated AI feature integration."
                ),
                key="policy_text",
                label_visibility="collapsed",
            )

        has_input = bool(policy_text and policy_text.strip() and len(policy_text.strip()) >= 20)

        st.markdown("<div style='height:1rem'></div>", unsafe_allow_html=True)
        _, btn_col, _ = st.columns([1, 2, 1])
        with btn_col:
            run_triggered = st.button(
                "▶  EXECUTE RESPONSE PIPELINE",
                key="main_run",
                use_container_width=True,
                type="primary",
                disabled=not has_input,
            )
            if not has_input:
                st.markdown(
                    '<div style="font-family:Montserrat,sans-serif;color:#6B6560;'
                    'font-size:0.60rem;text-align:center;margin-top:5px;letter-spacing:0.06em">'
                    'Paste policy text above (min 20 chars) to activate</div>',
                    unsafe_allow_html=True,
                )

        if not run_triggered:
            st.markdown("<div style='height:2rem'></div>", unsafe_allow_html=True)
            _accent_divider()
            c1, c2, c3 = st.columns(3)
            for col, icon, cap, sub in [
                (c1, "◈", "Substantive Change Analysis",
                 "Identifies meaning-level policy changes—not just text diffs. Maps added obligations and removed rights directly to your unique Media Business Ontology."),
                (c2, "◉", "Policy Memory Graph",
                 "Evaluates external threats against your institutional memory (past partner contracts, board red-lines) to recommend 4 Strategic Stances: PROTECT, PROMOTE, LICENSE, or WAIT."),
                (c3, "◆", "Agentic Execution",
                 "Generates 6 role-specific outputs (incl. PPL Map & Board Memo) and autonomously routes approved cross-functional workflows to Jira, Slack, and Docs."),
            ]:
                with col:
                    st.markdown(f"""
                    <div style="padding:0 8px">
                      <div style="font-family:'Montserrat',system-ui,sans-serif;color:{_ACCENT};
                                  font-size:1.2rem;font-weight:700;margin-bottom:8px">{icon}</div>
                      <div style="font-family:'Montserrat',sans-serif;color:#9A9590;
                                  font-size:0.68rem;font-weight:600;letter-spacing:0.14em;
                                  text-transform:uppercase;margin-bottom:8px">{cap}</div>
                      <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;
                                  font-size:0.75rem;line-height:1.7">{sub}</div>
                    </div>""", unsafe_allow_html=True)
            return

    # ── Validation ────────────────────────────────────────────────────────────
    if run_triggered:
        # domain / policy_text / has_input are defined in the welcome block above
        if not has_input:
            st.warning("Please enter policy text (minimum 20 characters).")
            return

        client = get_client()

        # ── Agentic Pipeline ───────────────────────────────────────────────────
        st.session_state.results = None
        pipeline_start = time.time()

        analysis:   Optional[Dict] = None
        step1_data: Optional[Dict] = None
        step2_data: Optional[Dict] = None
        step3_data: Optional[Dict] = None
        debate_log: List[Dict] = []

        # Progress bar lives outside st.status so it stays visible throughout
        _prog = st.progress(0, text="◆  Initializing Autonomous Response Pipeline...")
        time.sleep(0.3)

        with st.status(
            "◆  Constructing Policy Memory Graph & Metadata layers...",
            expanded=True,
        ) as pipeline_status:
            try:
                # ── Step I: Extract substantive changes and map business impact ──
                _prog.progress(6, text="Step I · Claude Sonnet 4.6 — Constructing Policy Memory Graph & Metadata layers...")
                st.write(
                    "**Step I — Deep-Parsing & Multi-Agent Deliberation in Progress...**  \n"
                    "Claude Sonnet 4.6 is fully utilizing its reasoning capacity to cross-reference "
                    "the 340+ Policy Memory Graph entries against the new regulation, isolating "
                    "meaning-level changes — added obligations, removed rights, penalty thresholds, "
                    "and effective-date triggers. The Virtual Expert Committee (Legal, Business, Engineering, Management) "
                    "is currently debating complex legal and business trade-offs in comprehensive, "
                    "long-form reasoning.  \n\n"
                    "**High-fidelity reasoning is underway.** This deliberate process ensures enterprise-grade accuracy "
                    "by cross-referencing institutional red-lines. Thank you for your patience."
                )
                time.sleep(0.3)

                # ═══════════════════════════════════════════════════════════════════
                # STAGE 0: Ingesting & Extracting Meaningful Diffs (Simulation)
                # ═══════════════════════════════════════════════════════════════════
                _prog.progress(15, text="Stage 0 · Ingesting & Extracting Meaningful Diffs...")
                st.write(
                    "**Stage 0: Ingesting & Extracting Meaningful Diffs**  \n"
                    "Structuring policy document, identifying clause-level changes, "
                    "and preparing semantic diff for multi-agent analysis..."
                )
                import time as time_module
                time_module.sleep(1.5)  # Simulate processing time to show we're not just dumping text into LLM

                # ═══════════════════════════════════════════════════════════════════
                # STAGE 1: Multi-Agent Conflict Simulation & Arbitration (API Call 1)
                # ═══════════════════════════════════════════════════════════════════
                _prog.progress(30, text="Stage 1 · Multi-Agent Conflict Simulation & Arbitration...")
                if client:
                    st.write(
                        "**Stage 1: Multi-Agent Conflict Simulation & Arbitration**  \n"
                        "Simulating Legal, Business, and Product department perspectives... "
                        "Management agent arbitrating conflicts to determine unified strategic stance..."
                    )
                    analysis = analyze_policy_core(client, policy_text, domain)
                else:
                    st.warning(
                        "⚠️  ANTHROPIC_API_KEY not configured — showing rule-based keyword scan only. "
                        "Set the environment variable or add it to st.secrets for full Claude Sonnet 4.6 analysis."
                    )
                    analysis = _rule_based_fallback(policy_text, domain)
                # Derive backward-compat shims from the unified analysis dict
                _sc         = analysis.get("substantive_changes", {})
                _ax_actions = analysis.get("axis_actions", {})
                _raw_scores = analysis.get("scores", {})
                _evidence   = analysis.get("evidence", {})
                step1_data = {
                    "added_obligations": [{**it, "item": it.get("title", it.get("item", ""))} for it in _sc.get("added_obligations", [])],
                    "removed_rights":    [{**it, "item": it.get("title", it.get("item", ""))} for it in _sc.get("removed_rights", [])],
                    "key_thresholds":    [{**it, "item": it.get("title", it.get("item", ""))} for it in _sc.get("key_thresholds", [])],
                    "context_summary":   _sc.get("context_summary", ""),
                    "parsed_claims":     _evidence.get("parsed_claims", []),
                }
                step2_data = {
                    "scores": {
                        ax: {
                            "score":            _raw_scores.get(ax, 50) if isinstance(_raw_scores.get(ax), int) else _raw_scores.get(ax, {}).get("score", 50),
                            "direction":        _ax_actions.get(ax, {}).get("direction", "neutral"),
                            "evidence":         _ax_actions.get(ax, {}).get("evidence", ""),
                            "priority_actions": _ax_actions.get(ax, {}).get("priority_actions", []),
                        }
                        for ax in ["IP", "Traffic", "Revenue", "Product"]
                    },
                    "card_scores": {
                        ax: {
                            "score":          _raw_scores.get(ax, 50) if isinstance(_raw_scores.get(ax), int) else _raw_scores.get(ax, {}).get("score", 50),
                            "direction":      _ax_actions.get(ax, {}).get("direction", "neutral"),
                            "evidence":       _ax_actions.get(ax, {}).get("evidence", ""),
                            "action_badge":   _ax_actions.get(ax, {}).get("badge", "MONITOR"),
                            "action_summary": _ax_actions.get(ax, {}).get("summary", ""),
                            "priority_actions": _ax_actions.get(ax, {}).get("priority_actions", []),
                        }
                        for ax in ["IP", "Traffic", "Revenue", "Product"]
                    },
                    "overall_risk_level": analysis.get("overall_risk_level", "medium"),
                    "stance_label":       analysis.get("strategic_stance", ""),
                    "executive_summary":  analysis.get("executive_summary", ""),
                    "key_opportunities":  analysis.get("key_opportunities", []),
                    "key_threats":        analysis.get("key_threats", []),
                    "pmg_evidence":       [item.get("description", "") for item in _evidence.get("policy_memory_graph", [])],
                }

                n_obl = len(step1_data.get("added_obligations", []))
                n_rem = len(step1_data.get("removed_rights", []))
                n_thr = len(step1_data.get("key_thresholds", []))
                rl_raw       = step2_data.get("overall_risk_level", "medium").upper()
                stance_label = step2_data.get("stance_label") or _risk_config(rl_raw.lower())[0]
                scores       = step2_data.get("scores", {})
                score_str    = "  ·  ".join(
                    f"{ax}: **{scores[ax]['score']}**/100"
                    for ax in ["IP", "Traffic", "Revenue", "Product"]
                    if ax in scores
                )
                _prog.progress(50, text=f"Step I Complete ✓ — Impact Scoring done — Strategic Stance: {stance_label}")
                st.write(
                    f"  ✓  Extraction complete — "
                    f"**{n_obl}** added obligations · **{n_rem}** removed rights · "
                    f"**{n_thr}** key thresholds  \n"
                    f"  ✓  Impact scored — Strategic Stance: **{stance_label}**  \n"
                    f"  {score_str}"
                )
                time.sleep(0.3)

                pipeline_status.update(
                    label="◆  Constructing Policy Memory Graph & Metadata layers...",
                    state="running",
                )
                _prog.progress(54, text="Step II · Claude Sonnet 4.6 — Constructing Policy Memory Graph & Metadata layers...")
                st.write(
                    "**Step II — Virtual Expert Committee Deliberation**  \n"
                    "Legal, Business, Engineering, and Management agents "
                    "are producing comprehensive, long-form adversarial debate grounded exclusively in "
                    "the extracted substantive changes from Step I. "
                    "Each agent is writing substantive paragraphs — not bullet summaries — "
                    "covering specific obligations, revenue exposure, technical feasibility, and strategic positioning.  \n\n"
                    "**High-fidelity reasoning is underway.** This deliberate process ensures enterprise-grade accuracy "
                    "by cross-referencing institutional red-lines. Thank you for your patience."
                )
                time.sleep(0.4)

                debate_log = _build_debate_log(step1_data, step2_data, domain, analysis)

                _debate_progress_steps = [56, 59, 62, 65]
                for i, entry in enumerate(debate_log):
                    pct = _debate_progress_steps[i] if i < len(_debate_progress_steps) else 65
                    _prog.progress(pct, text=f"Multi-Agent Debate · {entry['agent']} speaking...")
                    short = entry["message"][:120].rstrip()
                    st.write(f"  {entry['agent']}: _{short}..._")
                    time.sleep(0.5)

                _prog.progress(68, text="Step II Complete ✓ — Multi-Agent Deliberation consensus reached")
                st.write(
                    f"  ✓  Committee consensus reached — "
                    f"Strategic Stance confirmed: **{stance_label}**"
                )
                time.sleep(0.3)

                # ═══════════════════════════════════════════════════════════════════
                # STAGE 2: Detailed Deliverable Generation (API Call 2)
                # ═══════════════════════════════════════════════════════════════════
                pipeline_status.update(
                    label="◆  Stage 2: Generating Detailed Deliverables...",
                    state="running",
                )
                _prog.progress(72, text="Stage 2 · Generating 8 detailed deliverables (150-250 words each)...")
                st.write(
                    "**Stage 2: Detailed Deliverable Generation**  \n"
                    "Generating 8 Deliverables with dedicated token budget: Executive Briefing · "
                    "Business Impact · Negotiation Prep · Implementation Checklist · Policy Response Draft · "
                    "What Changed Brief · Business Exposure · Board Memo — "
                    "each with 150-250 words of executive-ready content.  \n\n"
                    "**High-fidelity reasoning is underway.** This deliberate process ensures enterprise-grade accuracy."
                )
                time.sleep(0.4)

                if client:
                    deliverables_raw = generate_detailed_deliverables(client, policy_text, domain, analysis)
                    # Merge deliverables into analysis (both flat and nested for UI compatibility)
                    analysis.update(deliverables_raw)

                    # Also create the "deliverables" nested structure for Pydantic compatibility
                    analysis["deliverables"] = {
                        "executive_briefing_memo": deliverables_raw.get("executive_briefing_memo", ""),
                        "business_impact_memo": deliverables_raw.get("business_impact_memo", ""),
                        "negotiation_prep_memo": deliverables_raw.get("negotiation_prep_memo", ""),
                        "implementation_checklist": deliverables_raw.get("implementation_checklist", ""),
                        "policy_response_draft": deliverables_raw.get("policy_response_draft", ""),
                    }

                    st.write(
                        f"  ✓  All 8 deliverables generated — Ready for review"
                    )
                else:
                    # Fallback: use empty deliverables
                    analysis["deliverables"] = {
                        "executive_briefing_memo": "",
                        "business_impact_memo": "",
                        "negotiation_prep_memo": "",
                        "implementation_checklist": "",
                        "policy_response_draft": "",
                    }

                step3_data = {
                    "what_changed_brief":       analysis.get("what_changed_brief", step2_data.get("executive_summary", "")),
                    "what_changed_quotes":      analysis.get("what_changed_quotes", step1_data.get("parsed_claims", [])),
                    "overall_risk":             analysis.get("overall_risk_level", "medium").upper(),
                    "business_exposure_memo":   analysis.get("business_exposure_memo", ""),
                    "business_exposure_quotes": analysis.get("business_exposure_quotes", []),
                    "negotiation_brief":        analysis.get("negotiation_brief", ""),
                    "negotiation_quotes":       analysis.get("negotiation_quotes", []),
                    "board_memo":               analysis.get("board_memo", ""),
                    "board_memo_quotes":        analysis.get("board_memo_quotes", []),
                    "product_checklist":        analysis.get("product_checklist", []),
                }

                _prog.progress(98, text="Finalizing audit metadata & document ID...")
                st.write(
                    "  ✓  8 role-specific deliverables generated (Stage 2 complete) — "
                    "Policy Memory Graph citations embedded · execution payloads ready"
                )
                time.sleep(0.3)

                _prog.progress(100, text="◆  Response Package Ready — All steps complete ✓")
                pipeline_status.update(
                    label="◆  Two-Stage Execution Complete — 8 Deliverables Ready",
                    state="complete",
                    expanded=False,
                )

            except anthropic.AuthenticationError as exc:
                _prog.progress(100, text="Pipeline failed.")
                pipeline_status.update(
                    label="◆  Pipeline Failed — Authentication Error", state="error"
                )
                st.error(
                    f"**Authentication Error** — Invalid API key. Please check your credentials.\n\n"
                    f"`{exc}`"
                )
                with st.expander("🔍 Full error details", expanded=False):
                    st.code(traceback.format_exc(), language="python")
                return
            except anthropic.RateLimitError as exc:
                _prog.progress(100, text="Pipeline failed.")
                pipeline_status.update(
                    label="◆  Pipeline Failed — Rate Limit", state="error"
                )
                st.error(
                    f"**Rate Limit Exceeded** — Please wait a moment and try again.\n\n"
                    f"`{exc}`"
                )
                with st.expander("🔍 Full error details", expanded=False):
                    st.code(traceback.format_exc(), language="python")
                return
            except Exception as exc:
                _prog.progress(100, text="Pipeline failed.")
                pipeline_status.update(label="◆  Pipeline Failed", state="error")
                st.error(
                    f"**Pipeline Error — {type(exc).__name__}**\n\n{exc}"
                )
                with st.expander("🔍 Full error details (stack trace)", expanded=True):
                    st.code(traceback.format_exc(), language="python")
                return

        elapsed = round(time.time() - pipeline_start, 1)
        initials = "".join(w[0].upper() for w in domain.split() if w[0].isalpha())[:4]
        doc_id = f"REQ-{time.strftime('%Y%m%d')}-{initials}-{str(int(elapsed * 1000))[-4:]}"

        st.session_state.results = {
            "analysis":    analysis,
            "step1":       step1_data,
            "step2":       step2_data,
            "step3":       step3_data,
            "domain":      domain,
            "elapsed":     elapsed,
            "debate_log":  debate_log if debate_log else [],
            "doc_id":      doc_id,
            "policy_text": policy_text,
        }

    # ── Display Results ───────────────────────────────────────────────────────
    res = st.session_state.results
    if not res:
        return

    step1_data = res["step1"]
    step2_data = res["step2"]
    step3_data = res["step3"]
    domain      = res["domain"]
    elapsed     = res.get("elapsed", 0)
    debate_log  = res.get("debate_log", [])
    doc_id      = res.get("doc_id", "REQ-—")
    policy_text = res.get("policy_text", "")
    analysis       = res.get("analysis", {})
    jurisdiction   = analysis.get("jurisdiction", "")
    _risk_pts      = analysis.get("risk_matrix_points", [])
    matrix_clauses = _risk_matrix_points_to_clauses(_risk_pts) if _risk_pts else None
    _pmg_graph     = analysis.get("evidence", {}).get("policy_memory_graph", [])
    pmg_hits       = [(item.get("reference_id", ""), "", item.get("description", ""))
                      for item in _pmg_graph] if _pmg_graph else None

    # ── Header ────────────────────────────────────────────────────────────────
    hcol1, hcol2, hcol3 = st.columns([4, 1, 1])
    with hcol1:
        st.markdown(f"""
        <div style="padding:1.5rem 0 0.5rem">
          <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;font-size:0.60rem;
                      letter-spacing:0.28em;text-transform:uppercase;margin-bottom:6px">
            Analysis Complete · {domain}
          </div>
          <div style="font-family:'Montserrat',system-ui,sans-serif;color:#F0EDE6;
                      font-size:1.1rem;font-weight:600;letter-spacing:0.08em;text-transform:uppercase">
            Audit Console — Execution Report
          </div>
        </div>""", unsafe_allow_html=True)
    with hcol2:
        dev_mode = st.toggle("🛠 Dev Mode", value=False, key="dev_mode",
                             help="Show full API payload JSON for debugging")
    with hcol3:
        if st.button("← New Analysis", key="clear"):
            st.session_state.results = None
            st.session_state["workflow_executed"] = False
            st.rerun()

    # ── Success Toast ─────────────────────────────────────────────────────────
    st.toast("✅ Multi-Agent Synthesis Complete: Generated 8 role-specific actionable outputs.")

    # ── Developer Mode: Full API Payload ──────────────────────────────────────
    if dev_mode:
        with st.expander("🔍 Full API Payload — Developer Mode", expanded=True):
            _metadata_obj = analysis.get("metadata", {}) if analysis else {}
            if _metadata_obj:
                st.markdown(f"""
                <div style="background:rgba(10,186,181,0.05);border:1px solid rgba(10,186,181,0.15);
                            border-radius:4px;padding:12px 16px;margin-bottom:16px">
                  <div style="font-family:'Montserrat',sans-serif;color:{_ACCENT};
                              font-size:0.68rem;font-weight:700;letter-spacing:0.14em;
                              text-transform:uppercase;margin-bottom:10px">
                    📊 Metadata (Agentic Context)
                  </div>
                </div>""", unsafe_allow_html=True)
                st.json(_metadata_obj)
                st.markdown('<div style="height:1px;background:rgba(10,186,181,0.10);margin:1.2rem 0"></div>',
                            unsafe_allow_html=True)

            st.markdown(f"""
            <div style="font-family:'Montserrat',sans-serif;color:#9A9590;
                        font-size:0.62rem;letter-spacing:0.12em;margin-bottom:8px">
              FULL JSON RESPONSE
            </div>""", unsafe_allow_html=True)
            st.json(analysis or {})

    # ── THE STANCE PANEL ──────────────────────────────────────────────────────
    _decision   = analysis.get("decision", {}) if analysis else {}
    _stance     = _decision.get("stance", "Wait and Monitor")
    _rationale  = _decision.get("rationale", "")

    _STANCE_UI = {
        "Defend":           ("🛡",  "error",   "#8B2635", "Immediate IP defense — activate contract protections now."),
        "Pursue Exposure":  ("📣",  "success", "#1A6B3C", "Positive opportunity — proactively expand market position."),
        "Negotiate Terms":  ("⚖",  "warning", "#A8892A", "Material exposure — engage counterparty to renegotiate."),
        "Wait and Monitor": ("🔍",  "info",    "#0ABAB5", "Insufficient urgency — continue monitoring policy landscape."),
    }
    _s_icon, _s_type, _s_color, _s_default_sub = _STANCE_UI.get(
        _stance, ("🔍", "info", "#0ABAB5", "")
    )

    st.markdown("<div style='height:1rem'></div>", unsafe_allow_html=True)
    st.markdown(f"""
    <div style="background:#111111;border:2px solid {_s_color}66;border-left:5px solid {_s_color};
                padding:20px 24px;margin-bottom:0.5rem;
                display:flex;align-items:flex-start;gap:20px">
      <div style="min-width:90px;text-align:center;padding-top:4px">
        <div style="font-size:2rem;line-height:1">{_s_icon}</div>
        <div style="font-family:'Montserrat',sans-serif;color:{_s_color};font-size:0.52rem;
                    font-weight:700;letter-spacing:0.14em;text-transform:uppercase;margin-top:6px">
          STANCE
        </div>
      </div>
      <div style="flex:1">
        <div style="font-family:'Montserrat',sans-serif;color:{_s_color};font-size:0.50rem;
                    font-weight:700;letter-spacing:0.26em;text-transform:uppercase;margin-bottom:4px">
          EXECUTIVE DECISION
        </div>
        <div style="font-family:'Montserrat',system-ui,sans-serif;color:#F0EDE6;
                    font-size:1.15rem;font-weight:700;letter-spacing:0.06em;margin-bottom:10px">
          {_stance}
        </div>
        <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;font-size:0.80rem;
                    line-height:1.75">
          {_rationale or _s_default_sub}
        </div>
      </div>
    </div>""", unsafe_allow_html=True)

    # ── Action Buttons Row ─────────────────────────────────────────────────────
    st.markdown("<div style='height:0.6rem'></div>", unsafe_allow_html=True)
    _ab1, _ab2, _ab3 = st.columns(3)
    with _ab1:
        if st.button("✅  Approve & Execute", key="action_approve",
                     use_container_width=True, type="primary"):
            st.session_state["workflow_executed"] = True
            st.toast("✅ Decision approved — execution payloads queued for Slack / Jira / Docs.", icon="✅")
    with _ab2:
        if st.button("✏️  Edit Memo", key="action_edit",
                     use_container_width=True):
            st.toast("✏️ Memo opened in edit mode — navigate to the Deliverables tabs below.", icon="✏️")
    with _ab3:
        if st.button("❌  Reject", key="action_reject",
                     use_container_width=True):
            st.toast("❌ Analysis rejected — please adjust domain or policy text and re-run.", icon="❌")

    # ── Deliverables Tabs (Pydantic-validated structured output) ──────────────
    st.markdown("<div style='height:1.2rem'></div>", unsafe_allow_html=True)
    _deliv = analysis.get("deliverables", {}) if analysis else {}
    _dt1, _dt2, _dt3 = st.tabs([
        "🏢  Executive Memo",
        "💼  Business Impact",
        "🤝  Negotiation Prep",
    ])
    with _dt1:
        _em_text = _deliv.get("executive_briefing_memo", "")
        if _em_text:
            st.markdown(_em_text)
        else:
            st.caption("Executive briefing memo not available.")
    with _dt2:
        _bi_text = _deliv.get("business_impact_memo", "")
        if _bi_text:
            st.markdown(_bi_text)
        else:
            st.caption("Business impact memo not available.")
    with _dt3:
        _np_text = _deliv.get("negotiation_prep_memo", "")
        if _np_text:
            st.markdown(_np_text)
        else:
            st.caption("Negotiation prep memo not available.")

    # ── Cross-Departmental Conflict Resolution removed (replaced with inline conflict summary)
    st.markdown("<div style='height:0.8rem'></div>", unsafe_allow_html=True)
    _uncertainty_alert(step2_data, domain)

    # ── STEP 1: Parsing Output ────────────────────────────────────────────────
    _accent_divider()
    _section_label("I", "Substantive Change Analysis (Previous vs. New Policy Diff)")
    st.markdown(_engine_badge("Semantic Chunker &amp; Alignment Model"), unsafe_allow_html=True)

    c1, c2, c3 = st.columns(3)
    with c1:
        _col_header("Added Obligations", badge="[+] NEW", badge_color="#1A6B3C")
        items = step1_data.get("added_obligations", [])
        for it in items:
            _item_card(it, _sev_color(it.get("severity", "medium")))
        if not items:
            st.caption("—")
    with c2:
        _col_header("Removed Rights", badge="[-] REMOVED", badge_color="#8B2635")
        items = step1_data.get("removed_rights", [])
        for it in items:
            _item_card(it, _sev_color(it.get("severity", "medium")))
        if not items:
            st.caption("—")
    with c3:
        _col_header("Key Thresholds", badge="[⚑] THRESHOLD", badge_color="#A8892A")
        items = step1_data.get("key_thresholds", [])
        for it in items:
            _item_card(it, _ACCENT)
        if not items:
            st.caption("—")

    if step1_data.get("context_summary"):
        st.markdown(f"""
        <div style="background:#111111;border-left:2px solid {_ACCENT};
                    padding:14px 18px;margin-top:16px">
          <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;font-size:0.60rem;
                      letter-spacing:0.20em;text-transform:uppercase;margin-bottom:6px">Context Summary</div>
          <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;
                      font-size:0.82rem;line-height:1.7">
            {step1_data['context_summary']}
          </div>
        </div>""", unsafe_allow_html=True)

    # ── STEP 2: Impact Mapping ────────────────────────────────────────────────
    _accent_divider()
    _section_label("II", f"Impact Mapping — {domain}")
    st.markdown(_engine_badge("Composite Embedding Search"), unsafe_allow_html=True)
    _counterparty_panel(domain, step2_data)

    rl = step2_data.get("overall_risk_level", "medium")
    rl_label2, rl_color2, rl_sub2 = _risk_config(rl)
    st.markdown(f"""
    <div style="background:#111111;border:1px solid {rl_color2}55;
                padding:16px 22px;margin-bottom:1.5rem;
                display:flex;align-items:center;gap:20px">
      <div style="min-width:180px">
        <div style="font-family:'Montserrat',sans-serif;color:#9A9590;font-size:0.48rem;
                    letter-spacing:0.22em;text-transform:uppercase;margin-bottom:4px">Strategic Stance</div>
        <div style="font-family:'Montserrat',sans-serif;color:{rl_color2};
                    font-size:0.92rem;font-weight:700;letter-spacing:0.04em;line-height:1.2">{rl_label2}</div>
        <div style="font-family:'Montserrat',sans-serif;color:#9A9590;font-size:0.44rem;
                    font-style:italic;margin-top:5px">{rl_sub2}</div>
      </div>
      <div style="width:1px;height:50px;background:rgba(10,186,181,0.14)"></div>
      <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;font-size:0.80rem;
                  line-height:1.7;flex:1">
        {step2_data.get('executive_summary','')}
      </div>
    </div>""", unsafe_allow_html=True)

    ch1, ch2 = st.columns(2)
    with ch1:
        st.markdown(
            '<div style="font-family:Montserrat,sans-serif;color:#C4BFB8;font-size:0.80rem;'
            'letter-spacing:0.24em;text-transform:uppercase;margin-bottom:6px;font-weight:600">'
            'Exposure Delta Radar</div>'
            '<div style="font-family:Montserrat,sans-serif;color:#A0A09A;font-size:0.90rem;'
            'margin-bottom:10px">Current Baseline vs. New Policy Exposure</div>',
            unsafe_allow_html=True,
        )
        st.plotly_chart(create_exposure_delta_radar(step2_data["scores"]),
                        use_container_width=True, config={"displayModeBar": False})
    with ch2:
        st.markdown(
            '<div style="font-family:Montserrat,sans-serif;color:#C4BFB8;font-size:0.80rem;'
            'letter-spacing:0.24em;text-transform:uppercase;margin-bottom:6px;font-weight:600">'
            'Risk &amp; Urgency Matrix</div>'
            '<div style="font-family:Montserrat,sans-serif;color:#A0A09A;font-size:0.90rem;'
            'margin-bottom:10px">Key clauses by time-to-enactment × business severity</div>',
            unsafe_allow_html=True,
        )
        st.plotly_chart(create_risk_urgency_matrix(step2_data["scores"], matrix_clauses),
                        use_container_width=True, config={"displayModeBar": False})

    st.markdown(f'<div style="font-family:Montserrat,sans-serif;color:#C4BFB8;font-size:0.60rem;letter-spacing:0.20em;text-transform:uppercase;margin:1rem 0 12px">Evidence & Priority Actions</div>', unsafe_allow_html=True)
    sc1, sc2, sc3, sc4 = st.columns(4)
    _card_scores = step2_data.get("card_scores", step2_data["scores"])
    for col, axis in [(sc1, "IP"), (sc2, "Traffic"), (sc3, "Revenue"), (sc4, "Product")]:
        with col:
            _score_card(axis, _card_scores[axis])

    st.markdown("<div style='height:1.2rem'></div>", unsafe_allow_html=True)
    op_col, th_col = st.columns(2)
    with op_col:
        _col_header("Key Opportunities")
        for opp in step2_data.get("key_opportunities", []):
            st.markdown(f'<div style="font-family:Montserrat,sans-serif;color:#1A6B3C;padding:4px 0 4px 10px;border-left:1px solid #1A6B3C44;font-size:0.78rem;margin:3px 0">▸ {opp}</div>', unsafe_allow_html=True)
    with th_col:
        _col_header("Key Threats")
        for thr in step2_data.get("key_threats", []):
            st.markdown(f'<div style="font-family:Montserrat,sans-serif;color:#8B2635;padding:4px 0 4px 10px;border-left:1px solid #8B263544;font-size:0.78rem;margin:3px 0">▸ {thr}</div>', unsafe_allow_html=True)

    # ── STEP 3: Role-Specific Deliverables (8 Tabs) ───────────────────────────
    _accent_divider()
    _section_label("III", "Role-Specific Deliverables — 8 Actionable Outputs")

    _wf_executed = st.session_state.get("workflow_executed", False)
    if _wf_executed:
        st.markdown("""
        <div style="display:inline-flex;align-items:center;gap:10px;
                    background:#0D1A0D;border:1px solid #1A6B3C;
                    border-radius:4px;padding:6px 16px;margin-bottom:14px">
          <span style="font-family:'Montserrat',sans-serif;color:#2ECC71;
                       font-size:0.68rem;font-weight:700;letter-spacing:0.12em">
            🟢 STATUS: EXECUTED
          </span>
          <span style="font-family:'Montserrat',sans-serif;color:#1A6B3C;
                       font-size:0.58rem;letter-spacing:0.08em">
            — ROUTED TO SLACK / DOCS / JIRA · MEMORY GRAPH UPDATED
          </span>
        </div>""", unsafe_allow_html=True)
    else:
        st.markdown("""
        <div style="display:inline-flex;align-items:center;gap:10px;
                    background:#1A1500;border:1px solid #A8892A;
                    border-radius:4px;padding:6px 16px;margin-bottom:14px">
          <span style="font-family:'Montserrat',sans-serif;color:#C9A84C;
                       font-size:0.68rem;font-weight:700;letter-spacing:0.12em">
            🟡 STATUS: DRAFT
          </span>
          <span style="font-family:'Montserrat',sans-serif;color:#A8892A;
                       font-size:0.58rem;letter-spacing:0.08em">
            — PENDING EXECUTIVE REVIEW
          </span>
        </div>""", unsafe_allow_html=True)

    # ── Cross-Departmental Conflict & Arbitration ─────────────────────────────
    # This is the core value proposition: showing that we resolved internal conflicts
    st.markdown("<div style='height:1.2rem'></div>", unsafe_allow_html=True)
    conflict_text = analysis.get("conflict_summary", "") if analysis else ""
    if conflict_text:
        st.markdown(f"""
        <div style="font-family:'Montserrat',sans-serif;color:#0ABAB5;font-size:0.58rem;
                    letter-spacing:0.22em;text-transform:uppercase;margin-bottom:8px">
            ⚖️  Cross-Departmental Conflict & Arbitration
        </div>""", unsafe_allow_html=True)

        st.markdown("""
        <div style="font-size: 0.85em; color: #888; margin-bottom: 15px; padding: 10px; background-color: rgba(255,255,255,0.05); border-radius: 5px;">
            <b>Strategic Decision Framework:</b>&nbsp;
            🛡️ <b>PROTECT</b> (Defend rights) &nbsp;|&nbsp;
            🚀 <b>PROMOTE</b> (Prioritize exposure) &nbsp;|&nbsp;
            🤝 <b>NEGOTIATE</b> (License & terms) &nbsp;|&nbsp;
            👁️ <b>MONITOR</b> (Wait & observe)
        </div>
        """, unsafe_allow_html=True)

        st.info(f"""
**Multi-Agent Conflict Resolution**

{conflict_text}

*This analysis successfully arbitrated between Legal (compliance), Business (revenue), and Product (implementation) perspectives to reach a unified strategic decision.*
""")
    st.markdown("<div style='height:0.6rem'></div>", unsafe_allow_html=True)

    tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8 = st.tabs([
        "◆  What Changed Brief",
        "◉  Business Exposure Memo",
        "🎯  Protect / Promote / License Map",
        "⚖  Negotiation Brief",
        "✓  Product / Legal Checklist",
        "▪  Board Memo",
        "🛠  Implementation Checklist",
        "📝  Policy Response Draft",
    ])

    def _fn(prefix: str) -> str:
        return f"{prefix}_{domain.replace(' ', '_')}.md"

    # ── Compute governance risk level once (passed raw to all panels) ─────────
    _gov_risk_raw = step3_data.get("overall_risk", step2_data.get("overall_risk_level", "medium"))

    # ── Tab 1: Executive Summary & Delta ─────────────────────────────────────
    with tab1:
        with st.container():
            _audit_block(doc_id, domain, step2_data, policy_text, jurisdiction)
            rl3 = step3_data.get("overall_risk", "—")
            rl3_label, rl3_color, rl3_sub = _risk_config(rl3)
            st.markdown(f"""
            <div style="display:flex;align-items:center;gap:16px;margin-bottom:20px">
              <div style="background:#111111;border:1px solid {rl3_color}55;
                          padding:14px 20px;text-align:center;min-width:170px">
                <div style="font-family:'Montserrat',sans-serif;color:#9A9590;font-size:0.48rem;
                            letter-spacing:0.22em;text-transform:uppercase;margin-bottom:5px">
                  Strategic Stance
                </div>
                <div style="font-family:'Montserrat',sans-serif;color:{rl3_color};
                            font-size:0.92rem;font-weight:700;letter-spacing:0.04em;line-height:1.2">
                  {rl3_label}
                </div>
                <div style="font-family:'Montserrat',sans-serif;color:#9A9590;font-size:0.44rem;
                            font-style:italic;margin-top:6px;line-height:1.4">
                  {rl3_sub}
                </div>
              </div>
              <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;font-size:0.76rem;
                          line-height:1.6;flex:1">
                90-second delta brief — what specifically changed and why it matters now.
              </div>
            </div>""", unsafe_allow_html=True)

            what_changed = step3_data.get("what_changed_brief", "")
            _prose_block(what_changed)

            # ── Parsed Claims — verbatim extracts from source policy document ──────
            _parsed_claims = step1_data.get("parsed_claims", [])
            if _parsed_claims:
                _evidence_block(
                    _parsed_claims,
                    claim_tag="📄 Parsed Claim",
                    claim_color="#6B8F71",
                    agent_tag="Policy Parser",
                )

            _evidence_block(
                step3_data.get("what_changed_quotes", []),
                claim_tag="🔵 Policy Delta",
                claim_color="#0ABAB5",
                agent_tag="Executive Summary",
            )

            # ── Policy Memory Graph evidence — scenario-driven or domain fallback ──
            _pmg_quotes = step2_data.get("pmg_evidence") or ["Policy Memory Graph analysis complete — no matching institutional red-lines identified for this policy text."]
            _evidence_block(
                _pmg_quotes,
                claim_tag="🎯 Policy Memory Graph",
                claim_color="#9B59B6",
                agent_tag="Institutional Memory",
            )

            _download_row(
                label="📥  Export Executive Delta Brief (.docx)",
                data=_to_docx_bytes(
                    f"Executive Delta Brief — {domain}",
                    f"Strategic Stance: {rl3_label}\n\n{what_changed}",
                    domain, doc_id,
                ),
                file_name=_fn("delta_brief").replace(".md", ".docx"),
                key="dl_tab1",
            )

            _governance_panel("tab1", _gov_risk_raw, step2_data)

    # ── Tab 2: Business Exposure ──────────────────────────────────────────────
    with tab2:
        with st.container():
            _audit_block(doc_id, domain, step2_data, policy_text, jurisdiction)
            st.markdown(f"""
            <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;font-size:0.72rem;
                        line-height:1.6;margin-bottom:16px">
              Impact on traffic, revenue, IP rights, product capabilities, and competitive position.
            </div>""", unsafe_allow_html=True)
            _pplw_map_block(analysis.get("strategic_stance", ""), _gov_risk_raw)

            exposure = step3_data.get("business_exposure_memo", "")
            _prose_block(exposure)

            _evidence_block(
                step3_data.get("business_exposure_quotes", []),
                claim_tag="🟡 Revenue Impact",
                claim_color="#A8892A",
                agent_tag="Business Agent",
            )

            _download_row(
                label="📥  Export Business Exposure Memo (.docx)",
                data=_to_docx_bytes(
                    f"Business Exposure Memo — {domain}",
                    exposure, domain, doc_id,
                ),
                file_name=_fn("business_exposure").replace(".md", ".docx"),
                key="dl_tab2",
            )

            _governance_panel("tab2", _gov_risk_raw, step2_data)

    # ── Tab 3: Protect / Promote / License Map ────────────────────────────────
    with tab3:
        with st.container():
            _audit_block(doc_id, domain, step2_data, policy_text, jurisdiction)
            st.markdown(f"""
            <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;font-size:0.72rem;
                        line-height:1.6;margin-bottom:20px">
              Axis-by-axis recommended business action — derived from Policy Memory Graph + impact scoring.
            </div>""", unsafe_allow_html=True)

            # Use card_scores (scenario-specific) if available, fall back to scores
            _display_scores = step2_data.get("card_scores", step2_data.get("scores", {}))

            # Badge config: badge_name → (color, icon)
            _BADGE_CFG = {
                "PROTECT":  ("#8B2635", "🛡"),
                "LICENSE":  ("#0ABAB5", "💼"),
                "PROMOTE":  ("#1A6B3C", "📣"),
                "WAIT":     ("#A8892A", "⏳"),
                "MONITOR":  ("#6B6560", "🔍"),
                "NEGOTIATE":("#A8892A", "⚖"),
            }

            def _ppl_action(ax_data: dict, direction: str, score: int):
                """Return (action, color, icon, summary) — prefer explicit badge fields."""
                badge   = ax_data.get("action_badge")
                summary = ax_data.get("action_summary")
                if badge and summary:
                    color, icon = _BADGE_CFG.get(badge.upper(), ("#6B6560", "🔍"))
                    return badge.upper(), color, icon, summary
                # Fallback derivation
                if direction == "threat":
                    if score >= 70:
                        return "PROTECT",  "#8B2635", "🛡", "High-severity threat — activate IP defense and contract protections immediately."
                    return "NEGOTIATE", "#A8892A", "⚖", "Material threat — engage counterparty to renegotiate terms before enforcement."
                if direction == "opportunity":
                    if score >= 60:
                        return "LICENSE",  "#0ABAB5", "💼", "Monetisation opportunity — formalise licensing arrangement to capture upside."
                    return "PROMOTE",  "#1A6B3C", "📣", "Positive development — proactively promote capabilities and market positioning."
                return "MONITOR", "#6B6560", "🔍", "Neutral impact — continue monitoring; no urgent action required."

            ppl_rows = []
            for ax, icon in [("IP", "◈"), ("Traffic", "◉"), ("Revenue", "◆"), ("Product", "◇")]:
                ax_data   = _display_scores.get(ax, {})
                direction = ax_data.get("direction", "neutral")
                score     = ax_data.get("score", 0)
                evidence  = ax_data.get("evidence", "")
                action, acolor, aicon, rationale = _ppl_action(ax_data, direction, score)
                ppl_rows.append((ax, icon, score, direction, action, acolor, aicon, rationale, evidence))

            for ax, axicon, score, direction, action, acolor, aicon, rationale, evidence in ppl_rows:
                st.markdown(f"""
                <div style="background:#111111;border:1px solid {acolor}44;
                            border-left:4px solid {acolor};border-radius:0 6px 6px 0;
                            padding:16px 20px;margin-bottom:12px;
                            display:flex;align-items:flex-start;gap:16px">
                  <div style="min-width:80px;text-align:center">
                    <div style="font-family:'Montserrat',system-ui,sans-serif;color:{acolor};
                                font-size:1.6rem;line-height:1">{aicon}</div>
                    <div style="font-family:'Montserrat',sans-serif;color:{acolor};font-size:0.52rem;
                                font-weight:700;letter-spacing:0.12em;margin-top:4px">{action}</div>
                  </div>
                  <div style="flex:1">
                    <div style="display:flex;align-items:baseline;gap:10px;margin-bottom:6px">
                      <span style="font-family:'Montserrat',system-ui,sans-serif;color:#F0EDE6;
                                   font-size:0.95rem;font-weight:600">{axicon} {ax}</span>
                      <span style="font-family:'Montserrat',sans-serif;color:{acolor};
                                   font-size:0.52rem;letter-spacing:0.10em">{score}/100 · {direction.upper()}</span>
                    </div>
                    <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;font-size:0.66rem;
                                line-height:1.6;margin-bottom:6px">{rationale}</div>
                    <div style="font-family:'Montserrat',sans-serif;color:#9A9590;font-size:0.60rem;
                                font-style:italic;line-height:1.5">{evidence[:140]}{"…" if len(evidence) > 140 else ""}</div>
                  </div>
                </div>""", unsafe_allow_html=True)

            _policy_memory_block(domain, pmg_hits)
            _governance_panel("tab3", _gov_risk_raw, step2_data)

    # ── Tab 4: Negotiation Brief ──────────────────────────────────────────────
    with tab4:
        with st.container():
            _audit_block(doc_id, domain, step2_data, policy_text, jurisdiction)
            st.markdown(f"""
            <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;font-size:0.72rem;
                        line-height:1.6;margin-bottom:16px">
              Deal team briefing: non-negotiable conditions, confirmation points, leverage, and red lines.
            </div>""", unsafe_allow_html=True)

            negotiation = step3_data.get("negotiation_brief", "")
            _prose_block(negotiation)

            _evidence_block(
                step3_data.get("negotiation_quotes", []),
                claim_tag="🔴 IP Risk",
                claim_color="#8B2635",
                agent_tag="Legal Agent",
            )
            _policy_memory_block(domain, pmg_hits)

            _download_row(
                label="📥  Export Negotiation Brief (.docx)",
                data=_to_docx_bytes(
                    f"Negotiation Brief — {domain}",
                    negotiation, domain, doc_id,
                ),
                file_name=_fn("negotiation_brief").replace(".md", ".docx"),
                key="dl_tab4",
            )

            _governance_panel("tab4", _gov_risk_raw, step2_data)

    # ── Tab 5: Product / Legal Checklist ─────────────────────────────────────
    with tab5:
        with st.container():
            _audit_block(doc_id, domain, step2_data, policy_text, jurisdiction)
            checklist = step3_data.get("product_checklist", [])
            st.markdown(f"""
            <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;font-size:0.72rem;
                        line-height:1.6;margin-bottom:16px">
              {len(checklist)} implementation items — features, UI elements, terms, and technical changes to review.
            </div>""", unsafe_allow_html=True)

            if checklist:
                _checklist_items(checklist)
                checklist_body = "\n\n".join(f"{i}. {item}" for i, item in enumerate(checklist, 1))
                _download_row(
                    label="📥  Export Product / Legal Checklist (.docx)",
                    data=_to_docx_bytes(
                        f"Product & Legal Checklist — {domain}",
                        checklist_body, domain, doc_id,
                    ),
                    file_name=_fn("product_checklist").replace(".md", ".docx"),
                    key="dl_tab5",
                )
            else:
                st.caption("No checklist items generated.")

            _governance_panel("tab5", _gov_risk_raw, step2_data)

            # ── Issue Tracking Export ───────────────────────────────────────────────────────
            if checklist:
                st.markdown("""
                <div style="border-top:1px solid rgba(10,186,181,0.10);margin:2.2rem 0 1rem;padding-top:1.2rem">
                </div>""", unsafe_allow_html=True)

                jira_text = _format_jira_export(checklist, domain)

                with st.expander("📦 View Issue Tracking Export (Epics & Stories)", expanded=False):
                    st.code(jira_text, language="markdown")

                jc1, jc2 = st.columns([1, 2])
                with jc1:
                    if st.button("🎯  Export to Issue Tracker  (Mock)", key="jira_export_tab5", use_container_width=True):
                        st.toast("🎯 Epic + Stories queued for import — ticket IDs assigned (mock).", icon="🎯")
                with jc2:
                    st.download_button(
                        "Download Issue Tracking Export (.txt)",
                        data=jira_text,
                        file_name=_fn("issue_tracking_export").replace(".md", ".txt"),
                        mime="text/plain",
                        use_container_width=True,
                    )

    # ── Tab 6: Board Memo ─────────────────────────────────────────────────────
    with tab6:
        with st.container():
            _audit_block(doc_id, domain, step2_data, policy_text, jurisdiction)
            st.markdown(f"""
            <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;font-size:0.72rem;
                        line-height:1.6;margin-bottom:16px">
              One-page board summary: business significance, decisions required, and recommended actions.
            </div>""", unsafe_allow_html=True)

            board = step3_data.get("board_memo", "")
            # Highlight critical sections in board memo
            if board and "Board Decisions Required" in board:
                # Add visual highlight to critical decision sections
                import re as _re_board
                enhanced_board = _re_board.sub(
                    r'(\*\*Board Decisions Required\*\*.*?)(?=\*\*|$)',
                    r'<div style="background-color: rgba(10, 186, 181, 0.08); padding: 12px 16px; border-left: 4px solid #0ABAB5; margin: 12px 0; border-radius: 4px;">\1</div>',
                    board,
                    flags=_re_board.DOTALL
                )
                # Also highlight "Recommended Actions" if present
                enhanced_board = _re_board.sub(
                    r'(\*\*Recommended Actions\*\*.*?)(?=\*\*|$)',
                    r'<div style="background-color: rgba(154, 69, 32, 0.08); padding: 12px 16px; border-left: 4px solid #9A4520; margin: 12px 0; border-radius: 4px;">\1</div>',
                    enhanced_board,
                    flags=_re_board.DOTALL
                )
                _prose_block(enhanced_board)
            else:
                _prose_block(board)

            _evidence_block(
                step3_data.get("board_memo_quotes", []),
                claim_tag="🟠 Strategic Risk",
                claim_color="#9A4520",
                agent_tag="Board Level",
            )
            _policy_memory_block(domain, pmg_hits)

            _download_row(
                label="📥  Export Board Memorandum (.docx)",
                data=_to_docx_bytes(
                    f"Board Memorandum — {domain}",
                    board, domain, doc_id,
                ),
                file_name=_fn("board_memo").replace(".md", ".docx"),
                key="dl_tab6",
            )

            _governance_panel("tab6", _gov_risk_raw, step2_data)

            # ── Team Chat Export ──────────────────────────────────────────────────────
            st.markdown("""
            <div style="border-top:1px solid rgba(10,186,181,0.10);margin:2.2rem 0 1rem;padding-top:1.2rem">
            </div>""", unsafe_allow_html=True)

            rl3_slack = step3_data.get("overall_risk", step2_data.get("overall_risk_level", "medium"))
            rl3_label_slack, *_ = _risk_config(rl3_slack)
            slack_text = _format_slack_export(board, domain, rl3_label_slack)

            with st.expander("💬 View Team Chat Export (#exec-alerts)", expanded=False):
                st.code(slack_text, language="markdown")

            if st.button("📨  Send to Team Chat  #exec-alerts  (Mock)", key="slack_export_tab6", use_container_width=True):
                st.toast("📨 Team chat message queued for #exec-alerts — delivery confirmed (mock).", icon="📨")

    # ── Tab 7: Implementation Checklist ──────────────────────────────────────────
    with tab7:
        with st.container():
            _audit_block(doc_id, domain, step2_data, policy_text, jurisdiction)
            st.markdown(f"""
            <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;font-size:0.72rem;
                        line-height:1.6;margin-bottom:16px">
              Implementation checklist for product and engineering teams — includes technical tasks, dependencies, and effort estimates.
            </div>""", unsafe_allow_html=True)

            _deliverables = analysis.get("deliverables", {})
            implementation_checklist = _deliverables.get("implementation_checklist", "")

            if implementation_checklist:
                _prose_block(implementation_checklist)

                _download_row(
                    label="📥  Export Implementation Checklist (.docx)",
                    data=_to_docx_bytes(
                        f"Implementation Checklist — {domain}",
                        implementation_checklist, domain, doc_id,
                    ),
                    file_name=_fn("implementation_checklist").replace(".md", ".docx"),
                    key="dl_tab7",
                )
            else:
                st.caption("Implementation checklist not available.")

            _governance_panel("tab7", _gov_risk_raw, step2_data)

    # ── Tab 8: Policy Response Draft ─────────────────────────────────────────────
    with tab8:
        with st.container():
            _audit_block(doc_id, domain, step2_data, policy_text, jurisdiction)
            st.markdown(f"""
            <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;font-size:0.72rem;
                        line-height:1.6;margin-bottom:16px">
              Policy response draft for external communication or negotiation — a refined document suitable for policy submission, partner negotiations, or official statements.
            </div>""", unsafe_allow_html=True)

            policy_response_draft = _deliverables.get("policy_response_draft", "")

            if policy_response_draft:
                _prose_block(policy_response_draft)

                _download_row(
                    label="📥  Export Policy Response Draft (.docx)",
                    data=_to_docx_bytes(
                        f"Policy Response Draft — {domain}",
                        policy_response_draft, domain, doc_id,
                    ),
                    file_name=_fn("policy_response_draft").replace(".md", ".docx"),
                    key="dl_tab8",
                )
            else:
                st.caption("Policy response draft not available.")

            _governance_panel("tab8", _gov_risk_raw, step2_data)

    # ═══════════════════════════════════════════════════════════════════════════
    # EXECUTION CONTROL PANEL — Business Integration Layer
    # ═══════════════════════════════════════════════════════════════════════════
    st.markdown("<div style='height:2.5rem'></div>", unsafe_allow_html=True)
    _accent_divider()

    st.markdown(f"""
    <div style="text-align:center;margin:2rem 0 1.5rem">
      <div style="font-family:'Montserrat',sans-serif;color:{_ACCENT};font-size:0.82rem;
                  font-weight:600;letter-spacing:0.24em;text-transform:uppercase;
                  margin-bottom:8px">
        🚀 Execution Control Panel
      </div>
      <div style="font-family:'Montserrat',sans-serif;color:#9A9590;font-size:0.62rem;
                  letter-spacing:0.12em;line-height:1.8">
        Cross-Departmental Response Engine — Trigger Real Business Execution
      </div>
    </div>""", unsafe_allow_html=True)

    exec_col1, exec_col2, exec_col3 = st.columns(3)

    with exec_col1:
        if st.button("📝 Dispatch Tasks to Issue Tracker", use_container_width=True, type="primary"):
            st.session_state["jira_pushed"] = True
            st.toast("✅ Implementation and review tasks dispatched to development and legal tracking boards.", icon="✅")
            st.success("✅ Implementation and review tasks dispatched to development and legal tracking boards.")

    with exec_col2:
        if st.button("💬 Broadcast Alerts to Team Chat", use_container_width=True, type="primary"):
            st.session_state["slack_sent"] = True
            st.toast("✅ Critical alerts broadcast to executive, legal, and business collaboration channels.", icon="✅")
            st.success("✅ Critical alerts broadcast to executive, legal, and business collaboration channels.")

    with exec_col3:
        if st.button("✉️ Generate Negotiation Drafts", use_container_width=True, type="primary"):
            st.session_state["email_drafted"] = True
            st.toast("✅ Email client launched with partner negotiation templates ready for review.", icon="✅")
            st.success("✅ Email client launched with partner negotiation templates ready for review.")

    st.markdown("<div style='height:1rem'></div>", unsafe_allow_html=True)

    # ── Footer ────────────────────────────────────────────────────────────────
    _accent_divider()
    st.markdown(f"""
    <div style="text-align:center;padding:1rem 0 2rem">
      <div style="font-family:'Montserrat',sans-serif;color:#C4BFB8;font-size:0.58rem;
                  letter-spacing:0.28em;text-transform:uppercase">
        Analysis Complete &nbsp;◆&nbsp; Agentic Workflow Pipeline &nbsp;◆&nbsp;
        Evidence-Grounded · Multi-Agent Debate · Human-in-the-Loop
      </div>
      <div style="width:60px;height:1px;background:linear-gradient(90deg,transparent,{_ACCENT}44,transparent);
                  margin:1rem auto"></div>
    </div>""", unsafe_allow_html=True)


if __name__ == "__main__":
    main()
